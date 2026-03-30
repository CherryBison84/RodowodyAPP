"""
Pełne okno programu na pulpicie: wszystkie zakładki, wykresy populacji, analizy, raporty i ustawienia.
"""

from __future__ import annotations

import csv
import os
import subprocess
import sys
from collections.abc import Callable
from typing import Any

import os
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox

try:
    import ttkbootstrap as ttk

    _USE_TTKBOOTSTRAP = True
except ImportError:  # pragma: no cover — instalacja minimalna bez pip
    from tkinter import ttk

    _USE_TTKBOOTSTRAP = False

from app.analytics.inbreeding_wright import (
    batch_offspring_inbreeding_F_from_parent_pairs,
    wright_inbreeding_F,
    wright_kinship_phi_and_relationship_R,
    wright_offspring_inbreeding_F_from_parents,
)
from app.data.dataset_loader import (
    load_dataset_from_path,
    load_default_bison_report,
    standardize_bison_report_dataframe_with_column_mapping,
    load_dataset_from_url,
    load_raw_dataframe_from_url,
)
from app.pedigree.ancestor_pedigree import (
    Person,
    build_people_map as _build_people_map,
    ensure_people_for_nodes,
    get_ancestor_levels_and_edges,
    get_ancestor_levels_unbounded,
)
from app.visualizations.ancestor_plot import plot_ancestor_pedigree
from app.analytics.line_membership import (
    compute_all_line_memberships,
    get_line_membership,
)
from app.analytics.kinship_decomposition import close_kinship_note, explain_pair_kinship
from app.analytics.mean_kinship import mean_kinship_pairwise
from app.analytics.population_genetics import (
    TEST_ID,
    FounderContributionComputer,
    compute_gi_and_family_data,
    compute_population_genetics_stats,
)
from app.config import app_icon_pil_best, get_config, resolve_app_icon_ico
from app.data.validator import validate_loaded_dataset
from app.ui.tk.breeding_helpers import suggest_pairs_with_constraints
from app.ui.tk.report_helpers import (
    build_individual_report_figures,
    build_population_report_figures,
    write_text_pages_to_pdf,
)
from app.ui.tk.theme import setup_theme
from app.ui.help_dialog import show_help_window
from app.ui import help_content as hc
from app.ui.typography import apply_matplotlib_fonts, tk_font
from app.ui.streamlit.streamlit_plotting import fig_column_missing_heatmap

APP_CFG = get_config()

# Liczba założycieli na wykresie p_i (Metryki populacji / Analizy populacji).
POP_FOUNDERS_PI_TOP_N = int(APP_CFG.report_founders_top_n)


def _apply_tk_window_icon(root: tk.Misc, ico_path: Path) -> None:
    """
    Ikona okna: `wm_iconphoto` z największej klatki ICO (macOS / Linux; czytelna też przy małych 16×16).
    Windows: dodatkowo natywne `iconbitmap(.ico)` dla paska zadań.
    """
    photo = None
    try:
        from PIL import Image, ImageTk

        pil_img = app_icon_pil_best()
        if pil_img is not None:
            img = pil_img
            mw = max(img.size)
            if mw < 48:
                k = max(2, (64 + mw - 1) // mw)
                img = img.resize((img.size[0] * k, img.size[1] * k), Image.Resampling.NEAREST)
            elif mw > 256:
                img = img.copy()
                img.thumbnail((256, 256), Image.Resampling.LANCZOS)
            photo = ImageTk.PhotoImage(img)
            root.wm_iconphoto(True, photo)
            setattr(root, "_wisent_icon_photo_keepalive", photo)
    except Exception:
        photo = None

    if sys.platform == "win32" or (photo is None and sys.platform != "darwin"):
        try:
            root.iconbitmap(str(ico_path.resolve()))
        except tk.TclError:
            pass


def _open_exported_file(path: str | Path) -> None:
    """Otwiera zapisany plik w domyślnej aplikacji systemowej (macOS / Windows / Linux)."""
    p = Path(path)
    if not p.is_file():
        return
    sp = str(p.resolve())
    try:
        if sys.platform == "darwin":
            subprocess.Popen(
                ["open", sp],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
        elif sys.platform.startswith("win"):
            os.startfile(sp)  # type: ignore[attr-defined]
        elif sys.platform.startswith("linux"):
            subprocess.Popen(
                ["xdg-open", sp],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
    except Exception:
        pass


def _clear_frame(frame: ttk.Frame) -> None:
    for w in frame.winfo_children():
        w.destroy()


def _pop_stat_row(
    parent: ttk.Frame,
    row: int,
    label: str,
    value_var: tk.StringVar,
    *,
    colors: Any,
    label_wrap: int = 400,
    value_wrap: int = 0,
) -> None:
    """Wiersz metryki: opis (stonowany) + wartość (pogrubiona, kolor akcentu wykresów)."""
    label_kw: dict[str, Any] = {
        "text": label,
        "font": tk_font(11),
        "foreground": colors.TEXT,
        "justify": "left",
        "anchor": "w",
    }
    if label_wrap > 0:
        label_kw["wraplength"] = label_wrap
    ttk.Label(parent, **label_kw).grid(row=row, column=0, sticky="nw", padx=(0, 12), pady=4)
    v_kwargs: dict[str, Any] = {
        "textvariable": value_var,
        "font": tk_font(12, bold=True),
        "foreground": colors.TEXT,
    }
    if value_wrap > 0:
        v_kwargs["wraplength"] = value_wrap
        v_kwargs["anchor"] = "w"
        v_kwargs["justify"] = "left"
        vst = "nw"
    else:
        v_kwargs["anchor"] = "e"
        v_kwargs["justify"] = "right"
        vst = "ne"
    ttk.Label(parent, **v_kwargs).grid(row=row, column=1, sticky=vst, pady=4)


def _save_figure_as_jpeg(fig, *, default_basename: str = "wykres") -> None:
    """
    Zapisuje dostarczony obiekt Matplotlib (Figure) do pliku JPEG.
    """
    try:
        filename = filedialog.asksaveasfilename(
            title="Zapis wykresu (jpeg)",
            defaultextension=".jpg",
            initialfile=f"{default_basename}.jpg",
            filetypes=[("JPEG", "*.jpg *.jpeg"), ("All files", "*.*")],
        )
    except Exception:
        filename = ""
    if not filename:
        return
    try:
        fig.savefig(filename, format="jpeg", dpi=200, bbox_inches="tight")
    except Exception as e:
        messagebox.showerror("Błąd", f"Nie mogę zapisać wykresu: {e}")
        return
    _open_exported_file(filename)


def render_birth_decade_charts(
    df_use: Any,
    *,
    state: dict[str, Any],
    colors: Any,
    save_figure_fn: Callable[..., None],
    pop_birth_sex_plot_area: ttk.Frame,
    pop_birth_line_plot_area: ttk.Frame,
    pop_birth_ratio_plot_area: ttk.Frame,
    pop_comp_sex_plot_area: ttk.Frame,
    pop_comp_line_plot_area: ttk.Frame,
) -> None:
    try:
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        from datetime import datetime

        from pandas import isna  # type: ignore[import-not-found]

        now_year = datetime.now().year
        min_dec = (1881 // 10) * 10  # 1880
        max_dec = (now_year // 10) * 10
        decades = list(range(min_dec, max_dec + 1, 10))
        decade_labels = [f"{d}-{d+9}" for d in decades]

        if df_use is None or getattr(df_use, "empty", True):
            return
        if "birth_year" not in df_use.columns:
            return

        def _parse_birth_year(v: object) -> int | None:
            if v is None:
                return None
            try:
                if isinstance(v, float) and v != v:
                    return None
            except Exception:
                pass
            if isna(v):
                return None
            try:
                y_int = int(float(v))
            except Exception:
                return None
            if y_int < 1881 or y_int > now_year:
                return None
            return y_int

        birth_int = df_use["birth_year"].apply(_parse_birth_year)
        dfc = df_use.copy()
        dfc["_birth_int"] = birth_int
        dfc = dfc.dropna(subset=["_birth_int"])
        if dfc.empty:
            return
        dfc["_birth_int"] = dfc["_birth_int"].astype(int)
        dfc["decade"] = (dfc["_birth_int"] // 10) * 10

        # --- Sex split ---
        def _norm_sex(v: object) -> str | None:
            if v is None:
                return None
            s = str(v).strip().upper()
            if s == "M":
                return "M"
            if s == "F":
                return "F"
            return None

        sex_norm = dfc["sex"].apply(_norm_sex) if "sex" in dfc.columns else None
        m_counts: dict[int, int] = {}
        f_counts: dict[int, int] = {}
        if sex_norm is not None:
            vc_m = (
                dfc[sex_norm == "M"].groupby("decade").size().to_dict()
                if not dfc.empty
                else {}
            )
            vc_f = (
                dfc[sex_norm == "F"].groupby("decade").size().to_dict()
                if not dfc.empty
                else {}
            )
            m_counts = vc_m
            f_counts = vc_f

        _clear_frame(pop_birth_sex_plot_area)
        fig1 = plt.Figure(figsize=(8.6, 3.4), dpi=100)
        ax1 = fig1.add_subplot(1, 1, 1)
        x = list(range(len(decades)))
        w = 0.38
        m_vals = [m_counts.get(d, 0) for d in decades]
        f_vals = [f_counts.get(d, 0) for d in decades]
        ax1.bar([i - w / 2 for i in x], m_vals, width=w, color="#9ecbff", edgecolor=colors.ACCENT, label="M")
        ax1.bar([i + w / 2 for i in x], f_vals, width=w, color="#ffb4c1", edgecolor=colors.ACCENT, label="F")
        ax1.set_title("Urodzenia w dekadach (płeć)")
        ax1.set_xlabel("dekada")
        ax1.set_ylabel("liczba urodzeń")
        ax1.set_xticks(x)
        ax1.set_xticklabels(decade_labels, rotation=45, ha="right", fontsize=8)
        ax1.legend(fontsize=8)
        fig1.tight_layout()
        ttk.Button(
            pop_birth_sex_plot_area,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig1: save_figure_fn(f, default_basename="pop_urodzenia_plec"),
        ).pack(anchor="w", pady=(0, 6))
        canvas1 = FigureCanvasTkAgg(fig1, master=pop_birth_sex_plot_area)
        canvas1.draw()
        canvas1.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # --- Line split (LB vs LC) ---
        def _norm_line(v: object) -> str | None:
            if v is None:
                return None
            s = str(v).strip().upper()
            if s == "LB":
                return "LB"
            if s == "LC":
                return "LC"
            return None

        line_norm = dfc["line"].apply(_norm_line) if "line" in dfc.columns else None
        lb_counts: dict[int, int] = {}
        lc_counts: dict[int, int] = {}
        if line_norm is not None:
            vc_lb = (
                dfc[line_norm == "LB"].groupby("decade").size().to_dict()
                if not dfc.empty
                else {}
            )
            vc_lc = (
                dfc[line_norm == "LC"].groupby("decade").size().to_dict()
                if not dfc.empty
                else {}
            )
            lb_counts = vc_lb
            lc_counts = vc_lc

        _clear_frame(pop_birth_line_plot_area)
        fig2 = plt.Figure(figsize=(8.6, 3.4), dpi=100)
        ax2 = fig2.add_subplot(1, 1, 1)
        lb_vals = [lb_counts.get(d, 0) for d in decades]
        lc_vals = [lc_counts.get(d, 0) for d in decades]
        ax2.bar([i - w / 2 for i in x], lc_vals, width=w, color="#2e8b57", edgecolor=colors.ACCENT, label="LC")
        ax2.bar([i + w / 2 for i in x], lb_vals, width=w, color="#d64545", edgecolor=colors.ACCENT, label="LB")
        ax2.set_title("Urodzenia w dekadach (LC vs LB)")
        ax2.set_xlabel("dekada")
        ax2.set_ylabel("liczba urodzeń")
        ax2.set_xticks(x)
        ax2.set_xticklabels(decade_labels, rotation=45, ha="right", fontsize=8)
        ax2.legend(fontsize=8)
        fig2.tight_layout()
        ttk.Button(
            pop_birth_line_plot_area,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig2: save_figure_fn(f, default_basename="pop_urodzenia_linie"),
        ).pack(anchor="w", pady=(0, 6))
        canvas2 = FigureCanvasTkAgg(fig2, master=pop_birth_line_plot_area)
        canvas2.draw()
        canvas2.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # --- Ratio F/M od 1900 (po dekadach) ---
        ratio_decades = [d for d in decades if d >= 1900]
        ratio_labels = [f"{d}-{d+9}" for d in ratio_decades]

        ratio_m = m_counts
        ratio_f = f_counts

        ratio_vals: list[float] = []
        for d in ratio_decades:
            m = ratio_m.get(d, 0)
            f = ratio_f.get(d, 0)
            if m <= 0:
                ratio_vals.append(float("nan"))
            else:
                ratio_vals.append(float(f) / float(m))

        _clear_frame(pop_birth_ratio_plot_area)
        fig3 = plt.Figure(figsize=(8.6, 3.4), dpi=100)
        ax3 = fig3.add_subplot(1, 1, 1)
        xs3 = list(range(len(ratio_decades)))
        ax3.plot(xs3, ratio_vals, marker="o", linewidth=2, color=colors.MUTED)
        ax3.axhline(1.0, color=colors.ACCENT, linewidth=1, alpha=0.8)
        ax3.set_title("Female/Male ratio w dekadach (F/M) od 1900")
        ax3.set_xlabel("dekada")
        ax3.set_ylabel("F/M")
        ax3.set_xticks(xs3)
        ax3.set_xticklabels(ratio_labels, rotation=45, ha="right", fontsize=8)
        fig3.tight_layout()
        ttk.Button(
            pop_birth_ratio_plot_area,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig3: save_figure_fn(f, default_basename="pop_urodzenia_ratio_FM"),
        ).pack(anchor="w", pady=(0, 6))
        canvas3 = FigureCanvasTkAgg(fig3, master=pop_birth_ratio_plot_area)
        canvas3.draw()
        canvas3.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # --- Kompletnosc rodowodu (MG/CG/EG) wg płci oraz linii ---
        people = state.get("people")
        if people:
            # bierzemy tylko rekordy, które istnieją w `people`
            dfc = df_use.copy()
            dfc["id"] = dfc["id"].astype(str)
            dfc = dfc[dfc["id"].isin(set(people.keys()))].reset_index(drop=True)

            # liczymy MG/CG/EG per osobnik (memo)
            pid_list = dfc["id"].tolist() if "id" in dfc.columns else []
            unique_pids = sorted(set(pid_list))

            comp_memo: dict[str, tuple[int, int, float]] = {}

            for pid in unique_pids:
                levels = get_ancestor_levels_unbounded(person_id=pid, people=people)
                by_gen: dict[int, int] = {}
                for _, lvl in levels.items():
                    if lvl is None or lvl <= 0:
                        continue
                    by_gen[lvl] = by_gen.get(lvl, 0) + 1

                if not by_gen:
                    comp_memo[pid] = (0, 0, 0.0)
                    continue

                MG = int(max(by_gen.keys()))
                CG = 0
                EG = 0.0
                for g, a_g in by_gen.items():
                    pcl_g = float(a_g) / float(2**g)
                    EG += pcl_g
                    if pcl_g >= 0.999999:
                        CG += 1
                comp_memo[pid] = (MG, CG, EG)

            def _norm_sex_comp(v: object) -> str:
                if v is None:
                    return "NA"
                s = str(v).strip().upper()
                return s if s in {"M", "F"} else "NA"

            def _norm_line_comp(v: object) -> str:
                if v is None:
                    return "NA"
                s = str(v).strip().upper()
                return s if s in {"LB", "LC"} else "NA"

            dfc["sex_norm"] = dfc["sex"].apply(_norm_sex_comp) if "sex" in dfc.columns else "NA"
            dfc["line_norm"] = dfc["line"].apply(_norm_line_comp) if "line" in dfc.columns else "NA"

            def _group_means(group_col: str, categories: list[str]) -> dict[str, tuple[float, float, float]]:
                out: dict[str, tuple[float, float, float]] = {}
                for cat in categories:
                    pids = dfc[dfc[group_col] == cat]["id"].tolist()
                    if not pids:
                        out[cat] = (0.0, 0.0, 0.0)
                        continue
                    MGs = [float(comp_memo[pid][0]) for pid in pids if pid in comp_memo]
                    CGs = [float(comp_memo[pid][1]) for pid in pids if pid in comp_memo]
                    EGs = [float(comp_memo[pid][2]) for pid in pids if pid in comp_memo]
                    if not MGs:
                        out[cat] = (0.0, 0.0, 0.0)
                        continue
                    out[cat] = (
                        float(sum(MGs)) / float(len(MGs)),
                        float(sum(CGs)) / float(len(CGs)),
                        float(sum(EGs)) / float(len(EGs)),
                    )
                return out

            sex_means = _group_means("sex_norm", ["M", "F"])
            line_means = _group_means("line_norm", ["LB", "LC", "NA"])

            # --- wykres płci ---
            _clear_frame(pop_comp_sex_plot_area)
            cats = ["M", "F"]
            MGv = [sex_means[c][0] for c in cats]
            CGv = [sex_means[c][1] for c in cats]
            EGv = [sex_means[c][2] for c in cats]

            fig_c = plt.Figure(figsize=(8.6, 3.4), dpi=100)
            ax_c = fig_c.add_subplot(1, 1, 1)
            xs = list(range(len(cats)))
            ww = 0.26
            ax_c.bar([i - ww for i in xs], MGv, width=ww, color=colors.BUTTON_BG2, edgecolor=colors.ACCENT, label="MG (max)")
            ax_c.bar([i for i in xs], CGv, width=ww, color=colors.BUTTON_BG, edgecolor=colors.ACCENT, label="CG (kompletne)")
            ax_c.bar([i + ww for i in xs], EGv, width=ww, color="#6b5b4d", edgecolor=colors.ACCENT, label="EG (równoważne)")
            ax_c.set_title("Kompletność: MG/CG/EG wg płci")
            ax_c.set_xlabel("płeć")
            ax_c.set_ylabel("wartość (średnia)")
            ax_c.set_xticks(xs)
            ax_c.set_xticklabels(cats, fontsize=9)
            ax_c.legend(fontsize=8)
            fig_c.tight_layout()
            ttk.Button(
                pop_comp_sex_plot_area,
                text="Zapis wykresu (jpeg)",
                command=lambda f=fig_c: save_figure_fn(f, default_basename="pop_kompletnosc_plec"),
            ).pack(anchor="w", pady=(0, 6))
            canvas_c = FigureCanvasTkAgg(fig_c, master=pop_comp_sex_plot_area)
            canvas_c.draw()
            canvas_c.get_tk_widget().pack(fill=tk.BOTH, expand=True)

            # --- wykres linii ---
            _clear_frame(pop_comp_line_plot_area)
            cats2 = ["LB", "LC", "NA"]
            MGv2 = [line_means[c][0] for c in cats2]
            CGv2 = [line_means[c][1] for c in cats2]
            EGv2 = [line_means[c][2] for c in cats2]

            fig_l = plt.Figure(figsize=(8.6, 3.4), dpi=100)
            ax_l = fig_l.add_subplot(1, 1, 1)
            xs2 = list(range(len(cats2)))
            ww2 = 0.22
            ax_l.bar([i - ww2 for i in xs2], MGv2, width=ww2, color=colors.BUTTON_BG2, edgecolor=colors.ACCENT, label="MG (max)")
            ax_l.bar([i for i in xs2], CGv2, width=ww2, color=colors.BUTTON_BG, edgecolor=colors.ACCENT, label="CG (kompletne)")
            ax_l.bar([i + ww2 for i in xs2], EGv2, width=ww2, color="#6b5b4d", edgecolor=colors.ACCENT, label="EG (równoważne)")
            ax_l.set_title("Kompletność: MG/CG/EG wg linii")
            ax_l.set_xlabel("linia")
            ax_l.set_ylabel("wartość (średnia)")
            ax_l.set_xticks(xs2)
            ax_l.set_xticklabels(cats2, fontsize=9)
            ax_l.legend(fontsize=8)
            fig_l.tight_layout()
            ttk.Button(
                pop_comp_line_plot_area,
                text="Zapis wykresu (jpeg)",
                command=lambda f=fig_l: save_figure_fn(f, default_basename="pop_kompletnosc_linie"),
            ).pack(anchor="w", pady=(0, 6))
            canvas_l = FigureCanvasTkAgg(fig_l, master=pop_comp_line_plot_area)
            canvas_l.draw()
            canvas_l.get_tk_widget().pack(fill=tk.BOTH, expand=True)
    except Exception:
        return


def render_inbreeding_year_trends(
    df_use: Any,
    *,
    state: dict[str, Any],
    pop_depth_inb_var: tk.StringVar,
    pop_unbounded_inb_var: tk.BooleanVar | tk.StringVar,
    pop_inb_year_sex_plot_area: ttk.Frame,
    pop_inb_year_line_plot_area: ttk.Frame,
    pop_ria_overall_var: tk.StringVar,
    pop_ne_var: tk.StringVar,
    save_figure_fn: Callable[..., None],
) -> None:
    """
    Average F oraz Rate of Inbred Animals (RIA, %) w czasie (TP),
    osobno dla:
    - płci (M/F)
    - linii (LB/LC)
    """
    try:
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        from datetime import datetime

        people = state.get("people")
        if not people:
            return

        if df_use is None or getattr(df_use, "empty", True):
            return
        if "birth_year" not in df_use.columns:
            return

        now_year = datetime.now().year

        def _parse_birth_year(v: object) -> int | None:
            if v is None:
                return None
            try:
                if isinstance(v, float) and v != v:
                    return None
            except Exception:
                pass
            try:
                y_int = int(float(v))
            except Exception:
                return None
            if y_int < 1900 or y_int > now_year:
                return None
            return y_int

        dfc = df_use.copy()
        dfc["id"] = dfc["id"].astype(str)
        dfc["birth_year_int"] = dfc["birth_year"].apply(_parse_birth_year)
        dfc = dfc.dropna(subset=["birth_year_int"]).reset_index(drop=True)
        if dfc.empty:
            return
        dfc["birth_year_int"] = dfc["birth_year_int"].astype(int)

        # Limit pokoleń dla obliczeń F:
        try:
            depth = int(str(pop_depth_inb_var.get()).strip())
        except Exception:
            depth = 4
        if depth < 0:
            depth = 0
        depth = min(depth, 30)
        max_generations_back = None if bool(pop_unbounded_inb_var.get()) else depth

        unique_ids = sorted(set(dfc["id"].tolist()))
        f_map: dict[str, float] = {}
        for pid in unique_ids:
            if pid not in people:
                continue
            try:
                f_map[pid] = float(
                    wright_inbreeding_F(
                        person_id=pid,
                        people=people,  # type: ignore[arg-type]
                        max_generations_back=max_generations_back,
                    ).F
                )
            except Exception:
                f_map[pid] = float("nan")

        # Przechowujemy wynik F na potrzeby dodatkowych wykresów (bez ponownego liczenia).
        state["population_f_map"] = f_map
        state["population_f_max_generations_back"] = max_generations_back

        dfc["_F"] = dfc["id"].apply(lambda pid: f_map.get(str(pid), float("nan")))
        dfc = dfc.dropna(subset=["_F"]).reset_index(drop=True)
        if dfc.empty:
            return

        years = sorted(set(dfc["birth_year_int"].tolist()))

        eps_inbred = 1e-15  # F>0 (w sensie numerycznym)

        try:
            f_vals = dfc["_F"].tolist() if "_F" in dfc.columns else []
            if f_vals:
                ria_overall = 100.0 * float(sum(1 for v in f_vals if v > eps_inbred)) / float(len(f_vals))
                pop_ria_overall_var.set(f"{ria_overall:.1f}%")
        except Exception:
            pass

        # --- Wykres wg płci ---
        _clear_frame(pop_inb_year_sex_plot_area)
        cats_sex = ["M", "F"]
        colors_sex = {"M": "#9ecbff", "F": "#ffb4c1"}

        fig = plt.Figure(figsize=(9, 6), dpi=100)
        ax_avg = fig.add_subplot(2, 1, 1)
        ax_ria = fig.add_subplot(2, 1, 2)

        for cat in cats_sex:
            dfc_cat = dfc.copy()
            sex_col = dfc_cat["sex"] if "sex" in dfc_cat.columns else None
            if sex_col is None:
                continue
            mask_cat = dfc_cat["sex"].astype(str).str.strip().str.upper() == cat
            avgF: list[float] = []
            ria: list[float] = []
            for y in years:
                g = dfc_cat[(dfc_cat["birth_year_int"] == y) & mask_cat]
                if g.empty:
                    avgF.append(float("nan"))
                    ria.append(float("nan"))
                    continue
                vals = g["_F"].tolist()
                avgF.append(float(sum(vals)) / float(len(vals)))
                ria.append(100.0 * float(sum(1 for v in vals if v > eps_inbred)) / float(len(vals)))
            ax_avg.plot(years, avgF, marker="o", markersize=2, linewidth=2, color=colors_sex[cat], label=f"{cat}")
            ax_ria.plot(years, ria, marker="o", markersize=2, linewidth=2, color=colors_sex[cat], label=f"{cat}")

        ax_avg.set_title("Average Inbreeding Coefficient (F) w TP — wg płci")
        ax_avg.set_xlabel("rok urodzenia")
        ax_avg.set_ylabel("średnie F")
        ax_avg.grid(True, alpha=0.25)
        ax_avg.legend(fontsize=8)

        ax_ria.set_title("Rate of Inbred Animals (RIA) w TP — wg płci")
        ax_ria.set_xlabel("rok urodzenia")
        ax_ria.set_ylabel("RIA (%) — F>0")
        ax_ria.grid(True, alpha=0.25)
        ax_ria.legend(fontsize=8)

        if len(years) > 15:
            step = 5
            ticks = [y for i, y in enumerate(years) if i % step == 0]
            ax_avg.set_xticks(ticks)
            ax_ria.set_xticks(ticks)
        fig.tight_layout()

        ttk.Button(
            pop_inb_year_sex_plot_area,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig: save_figure_fn(f, default_basename="pop_inbred_trend_plec"),
        ).pack(anchor="w", pady=(0, 6))

        canvas = FigureCanvasTkAgg(fig, master=pop_inb_year_sex_plot_area)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # --- Wykres wg linii ---
        _clear_frame(pop_inb_year_line_plot_area)
        cats_line = ["LB", "LC", "NA"]
        colors_line = {"LB": "#d64545", "LC": "#2e8b57", "NA": "#d6d0c4"}

        fig2 = plt.Figure(figsize=(9, 6), dpi=100)
        ax_avg2 = fig2.add_subplot(2, 1, 1)
        ax_ria2 = fig2.add_subplot(2, 1, 2)

        line_col = dfc["line"] if "line" in dfc.columns else None
        if line_col is None:
            return

        line_norm = dfc["line"].astype(str).str.strip().str.upper()
        for cat in cats_line:
            mask_cat = line_norm == cat
            avgF2: list[float] = []
            ria2: list[float] = []
            for y in years:
                g = dfc[(dfc["birth_year_int"] == y) & mask_cat]
                if g.empty:
                    avgF2.append(float("nan"))
                    ria2.append(float("nan"))
                    continue
                vals = g["_F"].tolist()
                avgF2.append(float(sum(vals)) / float(len(vals)))
                ria2.append(100.0 * float(sum(1 for v in vals if v > eps_inbred)) / float(len(vals)))
            ax_avg2.plot(years, avgF2, marker="o", markersize=2, linewidth=2, color=colors_line[cat], label=f"{cat}")
            ax_ria2.plot(years, ria2, marker="o", markersize=2, linewidth=2, color=colors_line[cat], label=f"{cat}")

        ax_avg2.set_title("Average Inbreeding Coefficient (F) w TP — wg linii")
        ax_avg2.set_xlabel("rok urodzenia")
        ax_avg2.set_ylabel("średnie F")
        ax_avg2.grid(True, alpha=0.25)
        ax_avg2.legend(fontsize=8)

        ax_ria2.set_title("Rate of Inbred Animals (RIA) w TP — wg linii")
        ax_ria2.set_xlabel("rok urodzenia")
        ax_ria2.set_ylabel("RIA (%) — F>0")
        ax_ria2.grid(True, alpha=0.25)
        ax_ria2.legend(fontsize=8)

        if len(years) > 15:
            step = 5
            ticks2 = [y for i, y in enumerate(years) if i % step == 0]
            ax_avg2.set_xticks(ticks2)
            ax_ria2.set_xticks(ticks2)
        fig2.tight_layout()

        ttk.Button(
            pop_inb_year_line_plot_area,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig2: save_figure_fn(f, default_basename="pop_inbred_trend_linie"),
        ).pack(anchor="w", pady=(0, 6))

        canvas2 = FigureCanvasTkAgg(fig2, master=pop_inb_year_line_plot_area)
        canvas2.draw()
        canvas2.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # --- N_e (efektywna wielkość populacji) z trendu F ---
        try:
            import numpy as np

            avgF_all: list[float] = []
            for y in years:
                g_all = dfc[dfc["birth_year_int"] == y]
                if g_all.empty:
                    avgF_all.append(float("nan"))
                else:
                    vals_all = g_all["_F"].tolist()
                    avgF_all.append(float(sum(vals_all)) / float(len(vals_all)))

            xs: list[float] = []
            ys: list[float] = []
            for y, v in zip(years, avgF_all):
                if v == v and y == y:  # nan-safe
                    xs.append(float(y))
                    ys.append(float(v))

            if len(xs) >= 2:
                slope_per_year = float(np.polyfit(xs, ys, 1)[0])  # dF/dyear
                gi_mean = state.get("population_gi_mean")
                if gi_mean is not None and gi_mean > 0 and slope_per_year == slope_per_year:
                    deltaF_per_gen = slope_per_year * float(gi_mean)
                    if deltaF_per_gen > 0:
                        ne = 1.0 / (2.0 * deltaF_per_gen)
                        pop_ne_var.set(f"{ne:.1f}")
                    else:
                        pop_ne_var.set("brak wzrostu F (ΔF<=0)")
        except Exception:
            pass

    except Exception:
        return


def render_f_distribution_charts(
    *,
    df_use: Any,
    state: dict[str, Any],
    colors: Any,
    save_figure_fn: Callable[..., None],
    pop_f_hist_plot_area: ttk.Frame,
    pop_f_scatter_plot_area: ttk.Frame,
) -> None:
    """Dodatkowe wykresy: rozkład F w populacji oraz F vs rok urodzenia."""
    try:
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        import numpy as np
    except Exception:
        return

    try:
        people = state.get("people")
        f_map: dict[str, float] = state.get("population_f_map") or {}
        if not people or not f_map:
            return
        if df_use is None or getattr(df_use, "empty", True):
            return
        if "birth_year" not in df_use.columns:
            return

        def _parse_birth_year(v: object, *, lo: int = 1900, hi: int | None = None) -> int | None:
            if hi is None:
                from datetime import datetime

                hi = datetime.now().year
            if v is None:
                return None
            try:
                if isinstance(v, float) and v != v:
                    return None
            except Exception:
                pass
            try:
                y_int = int(float(v))
            except Exception:
                return None
            if y_int < lo or y_int > hi:
                return None
            return y_int

        # Budujemy listy wartości dla osób z f_map.
        f_rows: list[tuple[int, str, str, float]] = []
        id_list = [str(x) for x in df_use["id"].astype(str).tolist()] if "id" in df_use.columns else []
        id_set = set(id_list)

        # Szybka mapa: id -> birth_year (unikamy df_use.loc w pętli).
        birth_map: dict[str, object] = {}
        try:
            if "id" in df_use.columns and "birth_year" in df_use.columns:
                tmp_ids = df_use["id"].astype(str)
                birth_series = df_use["birth_year"]
                birth_map = {rid: bv for rid, bv in zip(tmp_ids.tolist(), birth_series.tolist())}
        except Exception:
            birth_map = {}

        for pid, fv in f_map.items():
            if pid not in id_set:
                continue
            if fv != fv:  # nan-safe
                continue
            if pid not in people:
                continue
            p = people.get(pid)
            sex = (getattr(p, "sex", None) or "").strip().upper()
            line = (getattr(p, "line", None) or "").strip().upper()
            bval = birth_map.get(pid)
            bint = _parse_birth_year(bval)
            if bint is None:
                continue
            f_rows.append((bint, sex or "NA", line or "NA", float(fv)))

        if len(f_rows) < 2:
            return

        # --- Histogram F ---
        _clear_frame(pop_f_hist_plot_area)
        f_vals = [r[3] for r in f_rows]
        # Dla czytelności: log jeśli zakres jest duży; bezpiecznie na wartości <=0.
        f_pos = [v for v in f_vals if v > 0]
        use_log = (max(f_vals) / max(f_pos) > 50) if f_pos else False

        fig_h = plt.Figure(figsize=(8.6, 3.4), dpi=100)
        ax_h = fig_h.add_subplot(1, 1, 1)
        bins = min(30, max(10, int(np.sqrt(len(f_vals)))))
        if use_log:
            ax_h.hist([v for v in f_vals if v > 0], bins=bins, color=colors.BUTTON_BG2, edgecolor=colors.ACCENT)
            ax_h.set_yscale("log")
            ax_h.set_title("Rozkład F w populacji (osoby z F>0)")
        else:
            ax_h.hist(f_vals, bins=bins, color=colors.BUTTON_BG2, edgecolor=colors.ACCENT)
            ax_h.set_title("Rozkład F w populacji (histogram)")
        ax_h.set_xlabel("F (Wright, osobnik)")
        ax_h.set_ylabel("liczba osobników")
        fig_h.tight_layout()
        ttk.Button(
            pop_f_hist_plot_area,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig_h: save_figure_fn(f, default_basename="pop_F_hist"),
        ).pack(anchor="w", pady=(0, 6))
        canvas_h = FigureCanvasTkAgg(fig_h, master=pop_f_hist_plot_area)
        canvas_h.draw()
        canvas_h.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # --- Scatter: F vs rok urodzenia ---
        _clear_frame(pop_f_scatter_plot_area)
        fig_s = plt.Figure(figsize=(8.6, 3.4), dpi=100)
        ax_s = fig_s.add_subplot(1, 1, 1)
        for sex_key, col in {"M": "#9ecbff", "F": "#ffb4c1", "NA": colors.BUTTON_BG2}.items():
            xs = [b for b, sx, _ln, fv in f_rows if sx == sex_key]
            ys = [fv for _b, sx, _ln, fv in f_rows if sx == sex_key]
            if not xs:
                continue
            ax_s.scatter(xs, ys, s=10, alpha=0.65, color=col, label=sex_key)

        ax_s.set_title("F w populacji vs rok urodzenia (punktowy)")
        ax_s.set_xlabel("rok urodzenia")
        ax_s.set_ylabel("F (Wright)")
        ax_s.grid(True, alpha=0.25)
        ax_s.legend(fontsize=8)
        fig_s.tight_layout()
        ttk.Button(
            pop_f_scatter_plot_area,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig_s: save_figure_fn(f, default_basename="pop_F_scatter"),
        ).pack(anchor="w", pady=(0, 6))
        canvas_s = FigureCanvasTkAgg(fig_s, master=pop_f_scatter_plot_area)
        canvas_s.draw()
        canvas_s.get_tk_widget().pack(fill=tk.BOTH, expand=True)
    except Exception:
        return


def render_inbreeding_year_trends_smoothed(
    df_use: Any,
    *,
    state: dict[str, Any],
    pop_depth_inb_var: tk.StringVar,
    pop_unbounded_inb_var: tk.BooleanVar | tk.StringVar,
    pop_inb_smooth_plot_area: ttk.Frame,
    save_figure_fn: Callable[..., None],
    smooth_window: int = 7,
) -> None:
    """Wykres: średnie F i RIA (%) w czasie z wygładzeniem (jak na przykładach)."""
    try:
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        import numpy as np
        from datetime import datetime
    except Exception:
        return

    try:
        people = state.get("people")
        if not people:
            return
        if df_use is None or getattr(df_use, "empty", True):
            return
        if "birth_year" not in df_use.columns:
            return

        # Pobieramy wcześniej policzone F (oszczędza CPU). Jeśli brak, policzymy lokalnie.
        f_map: dict[str, float] = state.get("population_f_map") or {}
        max_generations_back = None if bool(pop_unbounded_inb_var.get()) else None
        try:
            depth = int(str(pop_depth_inb_var.get()).strip())
        except Exception:
            depth = 4
        if not bool(pop_unbounded_inb_var.get()):
            depth = max(0, min(30, depth))
            max_generations_back = depth
        state_max_gen = state.get("population_f_max_generations_back")
        if not f_map or (state_max_gen is not None and max_generations_back is not None and state_max_gen != max_generations_back):
            # F może być z innego limitu pokoleń — nie mieszajmy. Przeliczmy na szybko dla wykresu.
            unique_ids = sorted(set(df_use["id"].astype(str).tolist()))
            f_map = {}
            for pid in unique_ids:
                if pid not in people:
                    continue
                try:
                    f_map[pid] = float(
                        wright_inbreeding_F(
                            person_id=pid,
                            people=people,  # type: ignore[arg-type]
                            max_generations_back=max_generations_back,
                        ).F
                    )
                except Exception:
                    f_map[pid] = float("nan")

        def _parse_birth_year(v: object) -> int | None:
            if v is None:
                return None
            try:
                if isinstance(v, float) and v != v:
                    return None
            except Exception:
                pass
            try:
                y_int = int(float(v))
            except Exception:
                return None
            now_year = datetime.now().year
            if y_int < 1900 or y_int > now_year:
                return None
            return y_int

        dfc = df_use.copy()
        dfc["id"] = dfc["id"].astype(str)
        dfc["birth_year_int"] = dfc["birth_year"].apply(_parse_birth_year)
        dfc = dfc.dropna(subset=["birth_year_int"]).reset_index(drop=True)
        if dfc.empty:
            return
        dfc["birth_year_int"] = dfc["birth_year_int"].astype(int)
        dfc["_F"] = dfc["id"].apply(lambda pid: f_map.get(str(pid), float("nan")))
        dfc = dfc.dropna(subset=["_F"]).reset_index(drop=True)
        if dfc.empty:
            return

        years = sorted(set(dfc["birth_year_int"].tolist()))
        eps_inbred = 1e-15  # F>0 (w sensie numerycznym)

        avgF: list[float] = []
        ria: list[float] = []
        for y in years:
            g = dfc[dfc["birth_year_int"] == y]
            if g.empty:
                avgF.append(float("nan"))
                ria.append(float("nan"))
                continue
            vals = g["_F"].tolist()
            avgF.append(float(sum(vals)) / float(len(vals)))
            ria.append(100.0 * float(sum(1 for v in vals if v > eps_inbred)) / float(len(vals)))

        def _smooth(arr: list[float], w: int) -> list[float]:
            a = np.array(arr, dtype=float)
            valid = ~np.isnan(a)
            if valid.sum() < 3:
                return arr
            a0 = np.where(valid, a, 0.0)
            kernel = np.ones(w, dtype=float)
            s = np.convolve(a0, kernel, mode="same")
            c = np.convolve(valid.astype(float), kernel, mode="same")
            out = np.where(c > 0, s / c, np.nan)
            return out.tolist()

        smooth_w = max(3, int(smooth_window))
        sm_avgF = _smooth(avgF, smooth_w)
        sm_ria = _smooth(ria, smooth_w)

        _clear_frame(pop_inb_smooth_plot_area)
        fig = plt.Figure(figsize=(10.0, 6.2), dpi=100)
        ax_avg = fig.add_subplot(2, 1, 1)
        ax_ria = fig.add_subplot(2, 1, 2)

        # Surowe (czarne, półprzezroczyste) + trend wygładzony (czerwony).
        ax_avg.plot(years, avgF, color="black", alpha=0.22, linewidth=1)
        ax_avg.plot(years, sm_avgF, color=colors.ACCENT, linewidth=2.5)
        ax_avg.set_title("Average F (raw + smooth)")
        ax_avg.set_ylabel("F")
        ax_avg.grid(True, alpha=0.25)

        ax_ria.plot(years, ria, color="black", alpha=0.22, linewidth=1)
        ax_ria.plot(years, sm_ria, color=colors.ACCENT, linewidth=2.5)
        ax_ria.set_title("RIA (%) (raw + smooth)")
        ax_ria.set_xlabel("Year of Birth")
        ax_ria.set_ylabel("RIA (%) — F>0")
        ax_ria.grid(True, alpha=0.25)

        # Dopasowania jak na przykładach: litery paneli + korelacja F vs RIA.
        try:
            a = np.array(sm_avgF, dtype=float)
            b = np.array(sm_ria, dtype=float)
            mask = ~np.isnan(a) & ~np.isnan(b)
            r_val = float(np.corrcoef(a[mask], b[mask])[0, 1]) if int(mask.sum()) >= 3 else None
        except Exception:
            r_val = None
        ax_avg.text(0.98, 0.08, "A", transform=ax_avg.transAxes, ha="right", va="bottom", fontsize=18, color="black")
        ax_ria.text(0.98, 0.08, "B", transform=ax_ria.transAxes, ha="right", va="bottom", fontsize=18, color="black")
        if r_val is not None:
            fig.text(0.01, 0.01, f"F i RIA — korelacja (r) = {r_val:.3f}", fontsize=9)

        fig.tight_layout()
        ttk.Button(
            pop_inb_smooth_plot_area,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig: save_figure_fn(f, default_basename="pop_F_RIA_smooth"),
        ).pack(anchor="w", pady=(0, 6))
        canvas = FigureCanvasTkAgg(fig, master=pop_inb_smooth_plot_area)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
    except Exception:
        return


def render_gi_four_paths_smoothed(
    *,
    gi_data: dict[str, Any],
    colors: Any,
    save_figure_fn: Callable[..., None],
    pop_gi_4paths_plot_area: ttk.Frame,
    smooth_window: int = 3,
) -> None:
    """GI w 4 podoknach (Sire–Son, Sire–Daughter, Dam–Son, Dam–Daughter) + wygładzenie."""
    try:
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        import numpy as np
    except Exception:
        return

    try:
        gi_decades: dict[str, dict[int, list[float]]] = gi_data.get("gi_decades") or {}
        all_decades = sorted(
            set().union(*[set(gi_decades.get(k, {}).keys()) for k in gi_decades.keys()])
        )
        if not all_decades:
            return

        def _means_for(path_key: str) -> list[float]:
            out: list[float] = []
            for d in all_decades:
                xs = gi_decades.get(path_key, {}).get(d, [])
                out.append(float(sum(xs)) / float(len(xs)) if xs else float("nan"))
            return out

        series = {
            "FS": ("Sire-Son", "M→syn", _means_for("FS"), colors.BUTTON_BG2),
            "FD": ("Sire-Daughter", "M→córka", _means_for("FD"), colors.BUTTON_BG),
            "MS": ("Dam-Son", "F→syn", _means_for("MS"), colors.MUTED),
            "MD": ("Dam-Daughter", "F→córka", _means_for("MD"), colors.ACCENT),
        }

        def _smooth(arr: list[float], w: int) -> list[float]:
            a = np.array(arr, dtype=float)
            valid = ~np.isnan(a)
            if valid.sum() < 3:
                return arr
            a0 = np.where(valid, a, 0.0)
            kernel = np.ones(w, dtype=float)
            s = np.convolve(a0, kernel, mode="same")
            c = np.convolve(valid.astype(float), kernel, mode="same")
            out = np.where(c > 0, s / c, np.nan)
            return out.tolist()

        sw = max(3, int(smooth_window))

        _clear_frame(pop_gi_4paths_plot_area)
        fig = plt.Figure(figsize=(10.6, 7.2), dpi=100)
        ax_ss = fig.add_subplot(2, 2, 1)
        ax_sd = fig.add_subplot(2, 2, 2)
        ax_ds = fig.add_subplot(2, 2, 3)
        ax_dd = fig.add_subplot(2, 2, 4)
        axes = [ax_ss, ax_sd, ax_ds, ax_dd]
        keys = ["FS", "FD", "MS", "MD"]

        for ax, k in zip(axes, keys):
            title = series[k][0]
            raw_vals = series[k][2]
            sm_vals = _smooth(raw_vals, sw)
            ax.plot(all_decades, raw_vals, color=series[k][3], alpha=0.25, linewidth=1.2)
            ax.plot(all_decades, sm_vals, color="black", linewidth=2.2)
            ax.set_title(title)
            ax.set_xlabel("Year of Birth")
            ax.set_ylabel("GI")
            ax.grid(True, alpha=0.25)

        fig.suptitle("Generation Intervals over Birth Decades (4 paths)", y=0.99, fontsize=12)
        fig.tight_layout(rect=[0, 0, 1, 0.97])

        ttk.Button(
            pop_gi_4paths_plot_area,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig: save_figure_fn(f, default_basename="pop_GI_4paths_smooth"),
        ).pack(anchor="w", pady=(0, 6))
        canvas = FigureCanvasTkAgg(fig, master=pop_gi_4paths_plot_area)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
    except Exception:
        return


def render_founders_pi_chart(
    *,
    state: dict[str, Any],
    colors: Any,
    people: dict[str, Any] | None,
    save_figure_fn: Callable[..., None],
    pop_founders_plot_area: ttk.Frame,
) -> None:
    """Top założycieli (p_i) w zakładce populacji."""
    try:
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

        _clear_frame(pop_founders_plot_area)

        contribs = state.get("population_founder_contributions") or {}
        founder_items: list[tuple[Any, Any]] = []
        if isinstance(contribs, dict):
            try:
                founder_items = sorted(contribs.items(), key=lambda kv: kv[1], reverse=True)
            except Exception:
                founder_items = []

        top_k = min(POP_FOUNDERS_PI_TOP_N, len(founder_items))
        fig_fe = plt.Figure(figsize=(11.0, 4.2), dpi=100)
        ax_fe = fig_fe.add_subplot(1, 1, 1)

        if top_k > 0:
            top_items = founder_items[:top_k]
            vals = [float(v) for _, v in top_items]
            ids = [str(fid) for fid, _ in top_items]

            ax_fe.bar(range(len(top_items)), vals, color=colors.BUTTON_BG2, edgecolor=colors.ACCENT)
            ax_fe.set_title(f"Top {top_k} założycieli (p_i)")
            ax_fe.set_xlabel("założyciel (ID + imię)")
            ax_fe.set_ylabel("p_i (udział)")
            ax_fe.set_xticks(range(len(top_items)))

            labels: list[str] = []
            for fid in ids:
                p = people.get(fid) if people else None
                nm = getattr(p, "name", None) if p else None
                labels.append(f"{fid} ({nm})" if nm else fid)
            ax_fe.set_xticklabels(labels, rotation=75, ha="right", fontsize=7)
        else:
            ax_fe.text(0.5, 0.5, "Brak danych founder contributions", ha="center", va="center")
            ax_fe.axis("off")

        fig_fe.tight_layout()
        ttk.Button(
            pop_founders_plot_area,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig_fe: save_figure_fn(f, default_basename="pop_founders_pi"),
        ).pack(anchor="w", pady=(0, 6))
        canvas_fe = FigureCanvasTkAgg(fig_fe, master=pop_founders_plot_area)
        canvas_fe.draw()
        canvas_fe.get_tk_widget().pack(fill=tk.BOTH, expand=True)
    except Exception:
        pass


def render_gi_and_family_charts(
    *,
    colors: Any,
    save_figure_fn: Callable[..., None],
    people: dict[str, Any] | None,
    gi_data: dict[str, Any],
    pop_gi_plot_area: ttk.Frame,
    pop_gi_trend_plot_area: ttk.Frame,
    pop_family_plot_area: ttk.Frame,
) -> None:
    """Średnie GI, trend GI po dekadach, histogram wielkości rodzin pełnego rodzeństwa."""
    if not people:
        return

    try:
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        from collections import Counter

        gi_decades: dict[str, dict[int, list[float]]] = gi_data.get("gi_decades") or {
            "FS": {},
            "FD": {},
            "MS": {},
            "MD": {},
        }

        # GI bar chart
        _clear_frame(pop_gi_plot_area)
        gi_vals = [
            gi_data.get("gi_fs"),
            gi_data.get("gi_fd"),
            gi_data.get("gi_ms"),
            gi_data.get("gi_md"),
        ]
        # Bez U+2192 (→) — brak glifu w wielu fontach sans (np. Helvetica Neue); matplotlib ostrzega przy rysowaniu.
        gi_labels = ["Ojciec–Syn", "Ojciec–Córka", "Matka–Syn", "Matka–Córka"]
        fig_gi = plt.Figure(figsize=(8.6, 3.4), dpi=100)
        ax_gi = fig_gi.add_subplot(1, 1, 1)
        x = list(range(len(gi_labels)))
        bar_colors = [colors.BUTTON_BG2, colors.BUTTON_BG2, colors.BUTTON_BG, colors.BUTTON_BG]
        shown = [(i, v) for i, v in enumerate(gi_vals) if v is not None]
        if shown:
            means = [v if v is not None else 0.0 for v in gi_vals]
            ax_gi.bar(x, means, color=bar_colors, edgecolor=colors.ACCENT)
            ax_gi.set_xticks(x)
            ax_gi.set_xticklabels(gi_labels, rotation=20, ha="right", fontsize=8)
        else:
            ax_gi.text(0.5, 0.5, "Brak danych GI", ha="center", va="center")
            ax_gi.axis("off")
        ax_gi.set_title("Odstęp międzypokoleniowy (GI)")
        ax_gi.set_ylabel("GI (lata)")
        fig_gi.tight_layout()
        ttk.Button(
            pop_gi_plot_area,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig_gi: save_figure_fn(f, default_basename="pop_gi_mean"),
        ).pack(anchor="w", pady=(0, 6))
        canvas_gi = FigureCanvasTkAgg(fig_gi, master=pop_gi_plot_area)
        canvas_gi.draw()
        canvas_gi.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # --- GI trend (w dekadach) ---
        _clear_frame(pop_gi_trend_plot_area)
        all_decades = sorted(set().union(*[set(gi_decades[k].keys()) for k in gi_decades.keys()]))
        if all_decades:
            decade_labels = [f"{d}-{d+9}" for d in all_decades]

            def _decade_mean(path_key: str, d: int) -> float | None:
                xs = gi_decades.get(path_key, {}).get(d, [])
                if not xs:
                    return None
                return float(sum(xs)) / float(len(xs))

            gi_fs = [_decade_mean("FS", d) for d in all_decades]
            gi_fd = [_decade_mean("FD", d) for d in all_decades]
            gi_ms = [_decade_mean("MS", d) for d in all_decades]
            gi_md = [_decade_mean("MD", d) for d in all_decades]

            fig_t = plt.Figure(figsize=(8.6, 3.4), dpi=100)
            ax_t = fig_t.add_subplot(1, 1, 1)
            x_t = list(range(len(all_decades)))

            def _to_plot(arr: list[float | None]) -> list[float]:
                return [float(v) if v is not None else float("nan") for v in arr]

            ax_t.plot(x_t, _to_plot(gi_fs), marker="o", linewidth=2, color="#9ecbff", label="Ojciec–Syn")
            ax_t.plot(x_t, _to_plot(gi_fd), marker="o", linewidth=2, color="#ffb4c1", label="Ojciec–Córka")
            ax_t.plot(x_t, _to_plot(gi_ms), marker="o", linewidth=2, color="#2e8b57", label="Matka–Syn")
            ax_t.plot(x_t, _to_plot(gi_md), marker="o", linewidth=2, color="#d64545", label="Matka–Córka")

            ax_t.set_title("GI (trend) — Ojciec/Mak x płeć potomstwa (dekady)")
            ax_t.set_xlabel("dekada urodzenia potomstwa")
            ax_t.set_ylabel("średni GI (lata)")
            ax_t.grid(True, alpha=0.25)
            ax_t.legend(fontsize=8)

            if len(x_t) > 15:
                step = 2
            else:
                step = 1
            ticks = [i for i in x_t if i % step == 0]
            ax_t.set_xticks(ticks)
            ax_t.set_xticklabels(
                [decade_labels[i] for i in ticks],
                rotation=35,
                ha="right",
                fontsize=8,
            )
            fig_t.tight_layout()
            ttk.Button(
                pop_gi_trend_plot_area,
                text="Zapis wykresu (jpeg)",
                command=lambda f=fig_t: save_figure_fn(f, default_basename="pop_gi_trend"),
            ).pack(anchor="w", pady=(0, 6))
            canvas_t = FigureCanvasTkAgg(fig_t, master=pop_gi_trend_plot_area)
            canvas_t.draw()
            canvas_t.get_tk_widget().pack(fill=tk.BOTH, expand=True)
        else:
            fig_t = plt.Figure(figsize=(8.6, 3.4), dpi=100)
            ax_t = fig_t.add_subplot(1, 1, 1)
            ax_t.text(0.5, 0.5, "Brak danych GI w dekadach", ha="center", va="center")
            ax_t.axis("off")
            fig_t.tight_layout()
            ttk.Button(
                pop_gi_trend_plot_area,
                text="Zapis wykresu (jpeg)",
                command=lambda f=fig_t: save_figure_fn(f, default_basename="pop_gi_trend"),
            ).pack(anchor="w", pady=(0, 6))
            canvas_t = FigureCanvasTkAgg(fig_t, master=pop_gi_trend_plot_area)
            canvas_t.draw()
            canvas_t.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # Family size histogram
        _clear_frame(pop_family_plot_area)
        fam_sizes = gi_data.get("family_sizes") or []
        fig_f = plt.Figure(figsize=(8.6, 3.4), dpi=100)
        ax_f = fig_f.add_subplot(1, 1, 1)
        if fam_sizes:
            c = Counter(int(s) for s in fam_sizes)
            max_show = 10
            labels: list[str] = []
            counts: list[int] = []
            for s in range(1, max_show + 1):
                labels.append(str(s))
                counts.append(int(c.get(s, 0)))
            labels.append(f"{max_show}+")
            counts.append(int(sum(v for k, v in c.items() if k > max_show)))

            x2 = list(range(len(labels)))
            ax_f.bar(x2, counts, color=colors.BUTTON_BG2, edgecolor=colors.ACCENT)
            ax_f.set_xticks(x2)
            ax_f.set_xticklabels(labels, fontsize=8)
            ax_f.set_title("Rozkład wielkości rodzin pełnego rodzeństwa")
            ax_f.set_xlabel("wielkość rodziny (liczba rodzeństwa pełnego)")
            ax_f.set_ylabel("liczba rodzin")
        else:
            ax_f.text(0.5, 0.5, "Brak danych rodzin", ha="center", va="center")
            ax_f.axis("off")
        fig_f.tight_layout()
        ttk.Button(
            pop_family_plot_area,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig_f: save_figure_fn(f, default_basename="pop_family_sizes"),
        ).pack(anchor="w", pady=(0, 6))
        canvas_f = FigureCanvasTkAgg(fig_f, master=pop_family_plot_area)
        canvas_f.draw()
        canvas_f.get_tk_widget().pack(fill=tk.BOTH, expand=True)
    except Exception:
        pass


def run_tk_pro() -> None:
    apply_matplotlib_fonts()
    if _USE_TTKBOOTSTRAP:
        # Domyślnie „sandstone” (ziemia/ściółka); pełna paleta z Theme; nadpisywanie: WISENT_GUI_THEME=minty|litera|…
        theme_name = (os.environ.get("WISENT_GUI_THEME") or "sandstone").strip() or "sandstone"
        root = ttk.Window(themename=theme_name)
        root.title("WisentPedigree Pro+")
        root.geometry("1280x860")
    else:
        root = tk.Tk()
        root.title("WisentPedigree Pro+")
        root.geometry("1280x860")

    _ico = resolve_app_icon_ico()
    if _ico is not None:
        _apply_tk_window_icon(root, _ico)

    colors = setup_theme(root)

    # Lokalny styl dla fragmentu „Wskaźniki genetyczne i demograficzne”
    # — ma być czarny i spójny z resztą UI.
    try:
        st = root.style
        st.configure("PopMeta.TLabelframe", background=colors.PANEL_BG)
        st.configure(
            "PopMeta.TLabelframe.Label",
            background=colors.PANEL_BG,
            foreground=colors.TEXT,
            font=tk_font(12, bold=True),
        )
    except Exception:
        pass

    # -------------------------
    # Logo + header
    # -------------------------
    header = ttk.Frame(root, padding=(12, 10), style="TFrame")
    header.pack(side=tk.TOP, fill=tk.X)

    header_actions = ttk.Frame(header)
    header_actions.pack(side=tk.RIGHT, padx=(12, 0))

    def on_save_methods_guide_pdf() -> None:
        from app.ui.methods_guide_pdf import write_methods_guide_pdf

        filename = filedialog.asksaveasfilename(
            title="Zapisz przewodnik metod (PDF)",
            defaultextension=".pdf",
            initialfile="WisentPedigree_Pro_przewodnik_metod_2026.pdf",
            filetypes=[("PDF", "*.pdf"), ("Wszystkie pliki", "*.*")],
        )
        if not filename:
            return
        try:
            write_methods_guide_pdf(filename)
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się utworzyć PDF: {e}")
            return
        _set_status(f"Zapisano przewodnik metod: {filename}")
        _open_exported_file(filename)

    ttk.Button(
        header_actions,
        text="Przewodnik metod (PDF)",
        command=on_save_methods_guide_pdf,
    ).pack(side=tk.LEFT, padx=(0, 8))
    ttk.Button(
        header_actions,
        text="Pomoc (parametry i wykresy)",
        command=lambda: show_help_window(root, "WisentPedigree Pro+ — pomoc", hc.FULL_HELP_DOCUMENT),
    ).pack(side=tk.LEFT)

    logo_path = Path(__file__).resolve().parents[2] / "logo.png"
    logo_img = None
    if logo_path.exists():
        try:
            logo_img = tk.PhotoImage(file=str(logo_path))
            # quick downscale (Tk PhotoImage can be heavy; subsample is OK).
            try:
                w = logo_img.width()
            except Exception:
                w = 0
            # Docelowa szerokość logo w nagłówku (Tk PhotoImage sub-sampling).
            target_w = 150
            if w and w > target_w:
                factor = max(2, int(round(w / target_w)))
                logo_img = logo_img.subsample(factor, factor)
        except Exception:
            logo_img = None

    if logo_img is not None:
        ttk.Label(header, image=logo_img).pack(side=tk.LEFT, padx=(0, 12))

    ttk.Label(header, text="WisentPedigree Pro+", font=tk_font(18, bold=True)).pack(side=tk.LEFT)
    subtitle = ttk.Label(
        header,
        text="Import → walidacja → rejestr → analiza osobnika / par → populacja → raport (DOCX/PDF)",
        foreground=colors.MUTED,
    )
    subtitle.pack(side=tk.LEFT, padx=(16, 0))

    # -------------------------
    # Main: Notebook tabs
    # -------------------------
    status_var = tk.StringVar(value="Gotowe.")
    status_bar = tk.Frame(root, bd=1, relief=tk.SUNKEN, bg=colors.APP_BG)
    status_bar.pack(side=tk.BOTTOM, fill=tk.X)
    status_left = tk.Label(
        status_bar,
        textvariable=status_var,
        anchor="w",
        bg=colors.APP_BG,
        fg=colors.TEXT,
        padx=6,
    )
    status_left.pack(side=tk.LEFT, fill=tk.X, expand=True)
    status_right = tk.Label(
        status_bar,
        text="Autor: Magdalena Perlińska-Teresiak • 2026",
        anchor="e",
        bg=colors.APP_BG,
        fg=colors.MUTED,
        padx=8,
    )
    status_right.pack(side=tk.RIGHT)

    notebook = ttk.Notebook(root)
    notebook.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

    tab_loading = ttk.Frame(notebook, padding=14)
    tab_validation = ttk.Frame(notebook, padding=14)
    tab_persons = ttk.Frame(notebook, padding=14)
    tab_pedigree = ttk.Frame(notebook, padding=14)
    tab_analysis = ttk.Frame(notebook, padding=14)
    tab_population = ttk.Frame(notebook, padding=14)
    tab_reports = ttk.Frame(notebook, padding=14)
    tab_settings = ttk.Frame(notebook, padding=14)

    notebook.add(tab_loading, text="Import danych")
    notebook.add(tab_validation, text="Walidacja bazy")
    notebook.add(tab_persons, text="Rejestr osobników")
    notebook.add(tab_pedigree, text="Graf pedigree")
    notebook.add(tab_analysis, text="Analiza: osobnik / para / ranking")
    notebook.add(tab_population, text="Metryki populacji")
    notebook.add(tab_reports, text="Raportowanie")
    notebook.add(tab_settings, text="Konfiguracja")

    # Wewnętrzny notebook zakładki Analityka hodowlana.
    analyses_nb = ttk.Notebook(tab_analysis)
    analyses_nb.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

    tab_inb = ttk.Frame(analyses_nb, padding=10)
    tab_kpair = ttk.Frame(analyses_nb, padding=10)
    tab_mating = ttk.Frame(analyses_nb, padding=10)
    tab_breeding = ttk.Frame(analyses_nb, padding=10)
    analyses_nb.add(tab_inb, text="Osobnik — F (Wright)")
    analyses_nb.add(tab_kpair, text="Para — Φ i wyjaśnienie")
    analyses_nb.add(tab_mating, text="Ranking kojarzeń")
    analyses_nb.add(tab_breeding, text="Plan hodowli")

    # -------------------------
    # Reports tab
    # -------------------------
    rep_top = ttk.Frame(tab_reports)
    rep_top.pack(side=tk.TOP, fill=tk.X, pady=(14, 0))

    rep_person_id_var = tk.StringVar(value="")
    ttk.Label(rep_top, text="ID (Number):", foreground=colors.MUTED).grid(row=0, column=0, sticky="w", padx=(0, 10))
    rep_person_id_entry = ttk.Entry(rep_top, textvariable=rep_person_id_var, width=16)
    rep_person_id_entry.grid(row=0, column=1, sticky="w", padx=(0, 18))

    rep_ind_unbounded_var = tk.BooleanVar(value=True)
    rep_ind_unbounded_cb = ttk.Checkbutton(
        rep_top,
        text="Osobnik: bez limitu (do founderów)",
        variable=rep_ind_unbounded_var,
    )
    rep_ind_unbounded_cb.grid(row=0, column=2, sticky="w")

    rep_ind_depth_var = tk.StringVar(value="4")
    rep_ind_depth_row = ttk.Frame(rep_top)
    rep_ind_depth_row.grid(row=1, column=0, columnspan=3, sticky="w", pady=(8, 0))
    ttk.Label(rep_ind_depth_row, text="Osobnik: max pokoleń (gdy limit):").pack(side=tk.LEFT)
    rep_ind_depth_entry = ttk.Entry(rep_ind_depth_row, textvariable=rep_ind_depth_var, width=10)
    rep_ind_depth_entry.pack(side=tk.LEFT, padx=(8, 0))

    rep_pop_unbounded_var = tk.BooleanVar(value=False)
    rep_pop_unbounded_cb = ttk.Checkbutton(
        rep_top,
        text="Populacja: bez limitu",
        variable=rep_pop_unbounded_var,
    )
    rep_pop_unbounded_cb.grid(row=2, column=2, sticky="w", pady=(10, 0))

    rep_pop_depth_var = tk.StringVar(value="4")
    rep_pop_depth_row = ttk.Frame(rep_top)
    rep_pop_depth_row.grid(row=3, column=0, columnspan=3, sticky="w", pady=(6, 0))
    ttk.Label(rep_pop_depth_row, text="Populacja: max pokoleń (gdy limit):").pack(side=tk.LEFT)
    rep_pop_depth_entry = ttk.Entry(rep_pop_depth_row, textvariable=rep_pop_depth_var, width=10)
    rep_pop_depth_entry.pack(side=tk.LEFT, padx=(8, 0))

    rep_include_ind_var = tk.BooleanVar(value=True)
    rep_include_pop_var = tk.BooleanVar(value=True)
    rep_include_validation_var = tk.BooleanVar(value=True)
    rep_include_frame = ttk.LabelFrame(rep_top, text="Co uwzględnić w raporcie")
    rep_include_frame.grid(row=1, column=3, rowspan=3, sticky="nsew", padx=(24, 0), pady=(0, 0))
    rep_include_ind_cb = ttk.Checkbutton(
        rep_include_frame,
        text="Osobnik",
        variable=rep_include_ind_var,
    )
    rep_include_ind_cb.pack(anchor="w", padx=10, pady=(10, 0))
    rep_include_pop_cb = ttk.Checkbutton(
        rep_include_frame,
        text="Populacja",
        variable=rep_include_pop_var,
    )
    rep_include_pop_cb.pack(anchor="w", padx=10, pady=(6, 10))

    rep_include_validation_cb = ttk.Checkbutton(
        rep_include_frame,
        text="Walidacja bazy (cross-check)",
        variable=rep_include_validation_var,
    )
    rep_include_validation_cb.pack(anchor="w", padx=10, pady=(0, 10))

    def _sync_report_depth_state() -> None:
        rep_ind_depth_entry.configure(state="disabled" if bool(rep_ind_unbounded_var.get()) else "normal")
        rep_pop_depth_entry.configure(state="disabled" if bool(rep_pop_unbounded_var.get()) else "normal")

    _sync_report_depth_state()
    rep_ind_unbounded_var.trace_add("write", lambda *_args: _sync_report_depth_state())
    rep_pop_unbounded_var.trace_add("write", lambda *_args: _sync_report_depth_state())

    rep_btns = ttk.Frame(tab_reports)
    rep_btns.pack(side=tk.TOP, fill=tk.X, pady=(10, 0))

    # Ustawienia domyślne dla raportów (konfigurowane w zakładce Ustawienia).
    rep_default_include_plots_var = tk.BooleanVar(value=True)
    rep_include_plots_export_var = tk.BooleanVar(value=bool(rep_default_include_plots_var.get()))

    rep_output_text = tk.Text(tab_reports, height=18, wrap="word")
    rep_output_text.pack(side=tk.TOP, fill=tk.BOTH, expand=True, pady=(12, 0))
    rep_output_text.configure(state="disabled", bg=colors.ENTRY_BG, fg=colors.TEXT, insertbackground=colors.TEXT)

    rep_report_preview_var = tk.StringVar(value="")

    def _set_rep_status(msg: str) -> None:
        rep_report_preview_var.set(msg)

    ttk.Label(tab_reports, textvariable=rep_report_preview_var, foreground=colors.MUTED).pack(anchor="w", pady=(0, 6))

    def _fmt_line_membership(mem: object, *, label: str = "") -> str:
        if mem is None:
            return f"{label}NA"
        sire_id = getattr(mem, "sire_founder_id", None)
        sire_name = getattr(mem, "sire_founder_name", None)
        dam_id = getattr(mem, "dam_founder_id", None)
        dam_name = getattr(mem, "dam_founder_name", None)
        sire_steps = getattr(mem, "sire_steps", 0)
        dam_steps = getattr(mem, "dam_steps", 0)
        sire = f"{sire_id} ({sire_name})" if sire_id else "NA"
        dam = f"{dam_id} ({dam_name})" if dam_id else "NA"
        return f"{label}Sireline: {sire} [steps={sire_steps}]\n{label}Damline: {dam} [steps={dam_steps}]"

    def on_generate_report() -> None:
        df_std = state.get("df_std")
        people = state.get("people")
        if not people or df_std is None:
            messagebox.showinfo("Info", "Najpierw wczytaj bazę.")
            return

        if not (rep_include_ind_var.get() or rep_include_pop_var.get() or rep_include_validation_var.get()):
            messagebox.showinfo(
                "Info",
                "Wybierz przynajmniej jedną sekcję raportu: Osobnik, Populacja albo Walidacja bazy.",
            )
            return

        lines: list[str] = []
        lines.append("WisentPedigree Pro+ — Raport")
        lines.append(f"Plik: {state.get('raw_filename') or 'wczytana baza'}")
        lines.append("")

        if rep_include_validation_var.get():
            vrep = state.get("validation_report")
            if vrep is None:
                # Fallback: jeśli z jakiegoś powodu nie było jeszcze liczone.
                try:
                    from datetime import datetime

                    vrep = validate_loaded_dataset(
                        df_std=df_std,  # type: ignore[arg-type]
                        people=people,  # type: ignore[arg-type]
                        current_year=datetime.now().year,
                    )
                except Exception:
                    vrep = None

            lines.append("Walidacja bazy (cross-check spójności):")
            if vrep is None:
                lines.append("- brak danych walidacyjnych")
            else:
                lines.extend(vrep.to_text().splitlines())
            lines.append("")

        # -------------------------
        # Individual report
        # -------------------------
        if rep_include_ind_var.get():
            person_id = str(rep_person_id_var.get()).strip()
            if not person_id:
                messagebox.showerror("Błąd", "Podaj ID (Number) dla raportu osobnika.")
                return
            if person_id not in people:
                messagebox.showerror("Błąd", "Wybrane ID nie istnieje w wczytanych danych.")
                return

            if bool(rep_ind_unbounded_var.get()):
                f_res = wright_inbreeding_F(person_id=person_id, people=people, max_generations_back=None)
                depth_txt = "bez limitu (do founderów)"
                max_back = None
            else:
                try:
                    depth = int(str(rep_ind_depth_var.get()).strip())
                except Exception:
                    messagebox.showerror("Błąd", "Max pokoleń dla osobnika musi być liczbą całkowitą.")
                    return
                depth = max(0, depth)
                f_res = wright_inbreeding_F(person_id=person_id, people=people, max_generations_back=depth)
                depth_txt = f"limit = {depth}"
                max_back = depth

            lines.append(f"1) Raport osobnika — ID {person_id}")
            lines.append(f"- Inbred (Wright F): {f_res.F:.6f} ({depth_txt})")
            lines.append(f"- Ojciec: {f_res.father_id} — {f_res.father_name}")
            lines.append(f"- Matka : {f_res.mother_id} — {f_res.mother_name}")
            lines.append(f"- Miejsce urodzenia: {getattr(people.get(person_id), 'birth_location', None) or '—'}")

            # Completeness (MG/EG/PCI) — analogicznie jak w UI "Inbred (F)".
            MG = 0
            EG = 0.0
            PCI = 0.0
            by_gen: dict[int, int] = {}
            try:
                levels = get_ancestor_levels_unbounded(person_id=person_id, people=people)
                for _aid, lvl in levels.items():
                    try:
                        g = int(lvl)
                    except Exception:
                        continue
                    if g <= 0:
                        continue
                    by_gen[g] = by_gen.get(g, 0) + 1

                if by_gen:
                    MG = int(max(by_gen.keys()))
                    pci_sum = 0.0
                    for g in range(1, MG + 1):
                        a_g = int(by_gen.get(g, 0))
                        pcl_g = float(a_g) / float(2**g)
                        EG += pcl_g
                        pci_sum += pcl_g
                    PCI = pci_sum / float(MG) if MG > 0 else 0.0
            except Exception:
                pass

            lines.append(f"- Kompletność rodowodu (ANC, bez limitu): MG={MG}, EG={EG:.4f}, PCI={PCI:.4f}")

            subj_mem = get_line_membership(person_id, people)
            father_mem = get_line_membership(f_res.father_id, people) if f_res.father_id and f_res.father_id in people else None
            mother_mem = get_line_membership(f_res.mother_id, people) if f_res.mother_id and f_res.mother_id in people else None
            def _norm_line(v: object) -> str:
                if v is None:
                    return "NA"
                if isinstance(v, float) and v != v:
                    return "NA"
                s = str(v).strip().upper()
                return s if s in {"LB", "LC"} else "NA"

            own_line = _norm_line(getattr(people.get(person_id), "line", None))
            father_line = _norm_line(getattr(people.get(f_res.father_id) if f_res.father_id else None, "line", None))
            mother_line = _norm_line(getattr(people.get(f_res.mother_id) if f_res.mother_id else None, "line", None))
            lines.append("- Linie (sire/dam):")
            lines.append(f"- Linie LB/LC: oceniany={own_line}  ojciec={father_line}  matka={mother_line}")
            lines.append(_fmt_line_membership(subj_mem, label="  "))
            lines.append(_fmt_line_membership(father_mem, label="  Ojciec: ") if father_mem else "  Ojciec: NA")
            lines.append(_fmt_line_membership(mother_mem, label="  Matka: ") if mother_mem else "  Matka: NA")
            lines.append("")

            lines.append("Metodyka (skrócona):")
            lines.append("- F(i) = Phi(sire(i), dam(i)); Phi liczymy rekurencyjnie po rodowodzie.")
            lines.append("- Brak rodzica traktowany jest jak founder-stop (wkład do Phi dla tej ścieżki = 0).")
            lines.append("")

        # -------------------------
        # Population report
        # -------------------------
        if rep_include_pop_var.get():
            try:
                max_generations_back = None if bool(rep_pop_unbounded_var.get()) else int(str(rep_pop_depth_var.get()).strip())
            except Exception:
                messagebox.showerror("Błąd", "Max pokoleń dla populacji musi być liczbą całkowitą.")
                return
            if max_generations_back is not None:
                max_generations_back = max(0, max_generations_back)

            lines.append("2) Raport populacji (genetyka i kompletność)")
            try:
                stats = compute_population_genetics_stats(
                    df_std=df_std,  # type: ignore[arg-type]
                    people=people,  # type: ignore[arg-type]
                    max_generations_back=max_generations_back,
                    calc_f=True,
                    calc_completeness=True,
                    calc_founders=True,
                    calc_lines=True,
                )

                lines.append(f"- Liczba osobników (po odfiltrowaniu ID testowego): n={stats.n}")
                lines.append(f"- Założyciele (brak ojca lub matki): {stats.n_founders_any_missing_parent}")
                lines.append("")
                lines.append("- Inbred (Wright F) — statystyki:")
                lines.append(f"  mean={stats.inbreeding.mean_F:.6f}  median={stats.inbreeding.median_F:.6f}")
                lines.append(f"  min={stats.inbreeding.min_F:.6f}  max={stats.inbreeding.max_F:.6f}")
                lines.append(f"  zeros (F=0): {stats.inbreeding.zeros}/{stats.inbreeding.n}")
                lines.append("")
                lines.append("- Kompletność rodowodu:")
                lines.append(f"  mean EG={stats.completeness.mean_EG:.4f}")
                lines.append(f"  mean PCI={stats.completeness.mean_PCI:.4f}")
                lines.append("")
                lines.append("- Założyciele (founder contributions):")
                lines.append(f"  f_e={stats.founders.f_e:.4f}")
                lines.append(f"  f_a={stats.founders.f_a:.4f}")
                if stats.founders.f_a > 0:
                    lines.append(f"  Bottleneck f_e/f_a={stats.founders.f_e / stats.founders.f_a:.3f}")
                lines.append(f"  Dryf (aproksymacja, f_ge=f_e): 1.000")
                lines.append("")
                # Mean kinship (po parach i!=j) dla tej samej populacji i limitu pokoleń co F.
                try:
                    id_list = [str(x) for x in df_std["id"].tolist()]
                    mk_phi, mk_r, mk_note = mean_kinship_pairwise(
                        people,
                        id_list,
                        max_generations_back=max_generations_back,
                    )
                    if mk_phi is not None and mk_r is not None:
                        lines.append("- Mean kinship (pary i≠j):")
                        lines.append(f"  śr. Φ={mk_phi:.6f}  śr. R={mk_r:.6f}")
                        if mk_note:
                            lines.append(f"  uwaga: {mk_note}")
                        lines.append("")
                except Exception:
                    pass

                # GI i rodziny pełnego rodzeństwa (jak w panelu Metryki populacji).
                try:
                    gi_data = compute_gi_and_family_data(df_std, people)
                    lines.append("- Generation Interval (GI):")
                    if gi_data.get("gi_all") is not None:
                        lines.append(f"  GI średnio={float(gi_data['gi_all']):.2f} lat")
                    if gi_data.get("gi_fs") is not None:
                        lines.append(f"  Ojciec→syn={float(gi_data['gi_fs']):.2f} lat")
                    if gi_data.get("gi_fd") is not None:
                        lines.append(f"  Ojciec→córka={float(gi_data['gi_fd']):.2f} lat")
                    if gi_data.get("gi_ms") is not None:
                        lines.append(f"  Matka→syn={float(gi_data['gi_ms']):.2f} lat")
                    if gi_data.get("gi_md") is not None:
                        lines.append(f"  Matka→córka={float(gi_data['gi_md']):.2f} lat")
                    fs_fam = gi_data.get("family_sizes") or []
                    if fs_fam:
                        lines.append(
                            f"  Rodziny pełnego rodzeństwa: n={len(fs_fam)}, średnia wielkość={float(sum(fs_fam))/float(len(fs_fam)):.2f}"
                        )
                    lines.append("")
                except Exception:
                    pass

                lines.append("- Linie (kolumna line):")
                lines.append(f"  LB={stats.line_counts.get('LB', 0)}  LC={stats.line_counts.get('LC', 0)}  NA={stats.line_counts.get('NA', 0)}")
                # Podsumowanie miejsc urodzenia (top) — przydatne po dodaniu filtra birth_location.
                try:
                    if "birth_location" in df_std.columns:
                        loc_s = df_std["birth_location"].copy()
                        loc_s = loc_s.astype(str).str.strip()
                        loc_s = loc_s.replace({"": "NA", "nan": "NA", "None": "NA", "NaN": "NA"})
                        counts = loc_s.value_counts(dropna=False)
                        top_items = counts.head(8).to_dict()
                        lines.append("- Miejsce urodzenia (top):")
                        lines.append("  " + " | ".join([f"{k}={int(v)}" for k, v in top_items.items()]))
                except Exception:
                    pass
            except Exception as e:
                lines.append(f"- Błąd liczenia metryk populacyjnych: {e}")
            lines.append("")

        content = "\n".join(lines).strip() + "\n"
        rep_output_text.configure(state="normal")
        rep_output_text.delete("1.0", tk.END)
        rep_output_text.insert("1.0", content)
        rep_output_text.configure(state="disabled")
        _set_rep_status("Raport wygenerowany.")

    def on_save_report_docx() -> None:
        try:
            text_content = rep_output_text.get("1.0", tk.END).strip()
        except Exception:
            text_content = ""
        if not text_content:
            messagebox.showinfo("Info", "Najpierw wygeneruj raport.")
            return

        include_plots = bool(rep_include_plots_export_var.get())

        filename = filedialog.asksaveasfilename(
            title="Zapisz raport (DOCX)",
            defaultextension=".docx",
            initialfile="wisent_pedigree_report.docx",
            filetypes=[("DOCX", "*.docx"), ("All files", "*.*")],
        )
        if not filename:
            return

        try:
            from docx import Document  # type: ignore[import-not-found]

            doc = Document()
            doc.add_heading("WisentPedigree Pro+ — Raport", level=1)
            doc.add_paragraph("")
            for ln in text_content.splitlines():
                doc.add_paragraph(ln)
            doc.add_paragraph("")
            if include_plots:
                doc.add_paragraph(
                    "Wybrano opcję dołączenia wykresów. "
                    "Pełne osadzanie wykresów jest dostępne w eksporcie PDF."
                )
            else:
                doc.add_paragraph("Wykresy nie zostały dołączone (wg wyboru użytkownika).")
            doc.save(filename)
            _set_rep_status(f"Zapisano: {Path(filename).name}")
            _open_exported_file(filename)
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się zapisać DOCX: {e}")

    def on_save_report_pdf() -> None:
        try:
            text_content = rep_output_text.get("1.0", tk.END).strip()
        except Exception:
            text_content = ""

        if not text_content:
            messagebox.showinfo("Info", "Najpierw wygeneruj raport.")
            return

        try:
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_pdf import PdfPages
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie mogę użyć Matplotlib do PDF: {e}")
            return

        filename = filedialog.asksaveasfilename(
            title="Zapisz raport (PDF)",
            defaultextension=".pdf",
            initialfile="wisent_pedigree_report.pdf",
            filetypes=[("PDF", "*.pdf"), ("All files", "*.*")],
        )
        if not filename:
            return

        include_plots = bool(rep_include_plots_export_var.get())

        people = state.get("people")
        df_std = state.get("df_std")
        if not people or df_std is None:
            messagebox.showinfo("Info", "Najpierw wczytaj bazę.")
            return

        try:
            with PdfPages(filename) as pdf:
                write_text_pages_to_pdf(pdf, text=text_content)

                if include_plots and rep_include_ind_var.get():
                    figs_ind = build_individual_report_figures(
                        person_id=str(rep_person_id_var.get()).strip(),
                        people=people,  # type: ignore[arg-type]
                        ind_unbounded=bool(rep_ind_unbounded_var.get()),
                        ind_depth=int(str(rep_ind_depth_var.get()).strip() or "4"),
                        colors=colors,
                    )
                    for f in figs_ind:
                        pdf.savefig(f)
                        plt.close(f)

                if include_plots and rep_include_pop_var.get():
                    figs_pop = build_population_report_figures(
                        people=people,  # type: ignore[arg-type]
                        df_std=df_std,
                        pop_unbounded=bool(rep_pop_unbounded_var.get()),
                        pop_depth=int(str(rep_pop_depth_var.get()).strip() or "4"),
                        colors=colors,
                    )
                    for f in figs_pop:
                        pdf.savefig(f)
                        plt.close(f)

            _set_rep_status(f"Zapisano: {Path(filename).name}")
            _open_exported_file(filename)
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się zapisać PDF: {e}")

    rep_generate_btn = ttk.Button(rep_btns, text="Generuj raport", command=on_generate_report)
    rep_generate_btn.pack(side=tk.LEFT)
    rep_include_plots_export_cb = ttk.Checkbutton(
        rep_btns,
        text="Dołącz wykresy do eksportu",
        variable=rep_include_plots_export_var,
    )
    rep_include_plots_export_cb.pack(side=tk.LEFT, padx=(12, 0))
    rep_save_btn = ttk.Button(rep_btns, text="Zapisz raport (DOCX)", command=on_save_report_docx)
    rep_save_btn.pack(side=tk.LEFT, padx=(12, 0))
    rep_save_pdf_btn = ttk.Button(rep_btns, text="Zapisz raport (PDF)", command=on_save_report_pdf)
    rep_save_pdf_btn.pack(side=tk.LEFT, padx=(12, 0))
    ttk.Button(rep_btns, text="Pomoc", command=lambda: show_help_window(root, "Raportowanie", hc.SECTION_REPORTS)).pack(
        side=tk.LEFT, padx=(16, 0)
    )

    # -------------------------
    # Breeding plan tab
    # -------------------------
    ttk.Label(
        tab_breeding,
        text="Ta sekcja jest w przygotowaniu i wymaga dodatkowego przemyślenia logiki hodowlanej.\n"
        "W kolejnej iteracji dodamy finalny workflow doboru par i cele zarządzania różnorodnością.",
        foreground=colors.MUTED,
        wraplength=1000,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    ttk.Button(
        tab_breeding,
        text="Pomoc: plan hodowlany",
        command=lambda: show_help_window(root, "Plan hodowli", hc.SECTION_BREEDING),
    ).pack(anchor="w", pady=(0, 6))

    plan_main = ttk.Frame(tab_breeding)
    plan_main.pack(fill=tk.BOTH, expand=True, pady=(4, 0))

    plan_left = ttk.Frame(plan_main)
    plan_left.pack(side=tk.LEFT, fill=tk.Y, expand=False, pady=(0, 0), padx=(0, 14))

    plan_right = ttk.Frame(plan_main)
    plan_right.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

    plan_right_pane = ttk.Panedwindow(plan_right, orient=tk.VERTICAL)
    plan_right_pane.pack(fill=tk.BOTH, expand=True)

    plan_tree_upper = ttk.Frame(plan_right_pane)
    plan_detail_lower = ttk.LabelFrame(
        plan_right_pane,
        text="Zaznaczona para — statystyki i rodowód potomka",
    )
    plan_right_pane.add(plan_tree_upper, weight=2)
    plan_right_pane.add(plan_detail_lower, weight=3)

    plan_detail_body = ttk.Frame(plan_detail_lower)
    plan_detail_body.pack(fill=tk.BOTH, expand=True, padx=8, pady=8)

    plan_pair_stats_text = tk.Text(
        plan_detail_body,
        height=11,
        wrap="word",
        state="disabled",
        bg=colors.ENTRY_BG,
        fg=colors.TEXT,
        insertbackground=colors.TEXT,
        font=tk_font(9),
    )
    plan_pair_stats_text.pack(fill=tk.BOTH, expand=True, pady=(0, 6))
    plan_pair_btn_row = ttk.Frame(plan_detail_body)
    plan_pair_btn_row.pack(fill=tk.X, pady=(4, 0))
    plan_sel_pair: dict[str, str | None] = {"dam": None, "sire": None}

    # Risk calculation settings
    plan_risk_unbounded_var = tk.BooleanVar(value=False)
    plan_risk_depth_var = tk.StringVar(value="4")

    plan_risk_cb = ttk.Checkbutton(
        plan_left,
        text="Ryzyko bez limitu (do founderów)",
        variable=plan_risk_unbounded_var,
        state="disabled",
    )
    plan_risk_cb.pack(anchor="w", pady=(0, 6))

    plan_risk_depth_row = ttk.Frame(plan_left)
    plan_risk_depth_row.pack(anchor="w", pady=(0, 10), fill=tk.X)
    ttk.Label(plan_risk_depth_row, text="Max pokoleń (gdy limit):", foreground=colors.MUTED).pack(side=tk.LEFT)
    plan_risk_depth_entry = ttk.Entry(plan_risk_depth_row, textvariable=plan_risk_depth_var, width=8, state="disabled")
    plan_risk_depth_entry.pack(side=tk.LEFT, padx=(8, 0))

    def _sync_plan_risk_depth_state() -> None:
        depth_state = "disabled" if bool(plan_risk_unbounded_var.get()) else "normal"
        plan_risk_depth_entry.configure(state=depth_state)

    _sync_plan_risk_depth_state()
    plan_risk_unbounded_var.trace_add("write", lambda *_args: _sync_plan_risk_depth_state())

    # Single pair calculation inputs
    plan_pair_frame = ttk.LabelFrame(plan_left, text="Policz ryzyko dla pary")
    plan_pair_frame.pack(side=tk.TOP, fill=tk.X, pady=(0, 14))
    plan_pair_body = ttk.Frame(plan_pair_frame)
    plan_pair_body.pack(fill=tk.X, padx=10, pady=10)

    ttk.Label(plan_pair_body, text="Samica (F) — ID: ", foreground=colors.MUTED).grid(row=0, column=0, sticky="w")
    plan_dam_id_var = tk.StringVar(value="")
    plan_dam_id_entry = ttk.Entry(plan_pair_body, textvariable=plan_dam_id_var, width=16, state="disabled")
    plan_dam_id_entry.grid(row=0, column=1, sticky="w", padx=(8, 0))

    ttk.Label(plan_pair_body, text="Samiec (M) — ID: ", foreground=colors.MUTED).grid(row=1, column=0, sticky="w", pady=(6, 0))
    plan_sire_id_var = tk.StringVar(value="")
    plan_sire_id_entry = ttk.Entry(plan_pair_body, textvariable=plan_sire_id_var, width=16, state="disabled")
    plan_sire_id_entry.grid(row=1, column=1, sticky="w", padx=(8, 0), pady=(6, 0))

    plan_calc_pair_result_var = tk.StringVar(value="Ryzyko F( potomek ) = -")
    ttk.Label(plan_pair_body, textvariable=plan_calc_pair_result_var, foreground=colors.TEXT).grid(
        row=2, column=0, columnspan=2, sticky="w", pady=(10, 0)
    )

    def _set_plan_status(msg: str) -> None:
        status_var.set(msg)
        try:
            root.update_idletasks()
        except Exception:
            pass

    def _norm_line(v: object) -> str:
        if v is None:
            return "NA"
        if isinstance(v, float) and v != v:
            return "NA"
        s = str(v).strip().upper()
        return s if s in {"LB", "LC"} else "NA"

    def on_calc_plan_pair() -> None:
        people = state.get("people")
        df_std = state.get("df_std")
        if not people or df_std is None:
            messagebox.showinfo("Info", "Najpierw wczytaj bazę.")
            return

        dam_id = str(plan_dam_id_var.get()).strip()
        sire_id = str(plan_sire_id_var.get()).strip()
        if not dam_id or not sire_id:
            messagebox.showerror("Błąd", "Podaj ID samicy i samca.")
            return
        if dam_id not in people or sire_id not in people:
            messagebox.showerror("Błąd", "Wybrane ID nie istnieje w wczytanych danych.")
            return

        try:
            max_back = None if bool(plan_risk_unbounded_var.get()) else int(str(plan_risk_depth_var.get()).strip())
        except Exception:
            messagebox.showerror("Błąd", "Max pokoleń musi być liczbą całkowitą.")
            return
        if max_back is not None and max_back < 0:
            messagebox.showerror("Błąd", "Max pokoleń nie może być ujemne.")
            return

        pairs = [(sire_id, dam_id)]
        try:
            F_off = batch_offspring_inbreeding_F_from_parent_pairs(
                pairs,
                people,
                max_generations_back=max_back,
            )[0]
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie mogę policzyć F potomka: {e}")
            return

        dam_line = _norm_line(getattr(people.get(dam_id), "line", None))
        sire_line = _norm_line(getattr(people.get(sire_id), "line", None))
        plan_calc_pair_result_var.set(f"Ryzyko F( potomek ) = {F_off:.6f} (dam line={dam_line}, sire line={sire_line})")
        _set_plan_status("Gotowe: policzono ryzyko pary.")

    plan_calc_pair_btn = ttk.Button(plan_pair_body, text="Policz ryzyko", command=on_calc_plan_pair, state="disabled")
    plan_calc_pair_btn.grid(row=3, column=0, columnspan=2, sticky="w", pady=(10, 0))

    # Ranking / suggestion
    plan_rank_frame = ttk.LabelFrame(plan_left, text="Podpowiedz pary (TOP-N)")
    plan_rank_frame.pack(side=tk.TOP, fill=tk.X, pady=(0, 14))
    plan_rank_body = ttk.Frame(plan_rank_frame)
    plan_rank_body.pack(fill=tk.X, padx=10, pady=10)

    ttk.Label(plan_rank_body, text="Linia (dla obu rodziców):", foreground=colors.MUTED).pack(anchor="w")
    plan_line_filter_var = tk.StringVar(value="Bez filtra")
    plan_line_filter_combo = ttk.Combobox(
        plan_rank_body,
        textvariable=plan_line_filter_var,
        state="disabled",
        values=["Bez filtra", "LB", "LC", "LB+LC", "NA"],
        width=16,
    )
    plan_line_filter_combo.pack(anchor="w", pady=(6, 0))

    ttk.Label(plan_rank_body, text="Wiek (lat) — min:", foreground=colors.MUTED).pack(anchor="w", pady=(10, 0))
    plan_min_age_var = tk.StringVar(value="0")
    plan_min_age_entry = ttk.Entry(plan_rank_body, textvariable=plan_min_age_var, width=8, state="disabled")
    plan_min_age_entry.pack(anchor="w", pady=(2, 0))

    ttk.Label(plan_rank_body, text="Wiek (lat) — max:", foreground=colors.MUTED).pack(anchor="w", pady=(8, 0))
    plan_max_age_var = tk.StringVar(value="80")
    plan_max_age_entry = ttk.Entry(plan_rank_body, textvariable=plan_max_age_var, width=8, state="disabled")
    plan_max_age_entry.pack(anchor="w", pady=(2, 0))

    ttk.Label(plan_rank_body, text="Limit kandydatów (dla wydajności):", foreground=colors.MUTED).pack(anchor="w", pady=(10, 0))
    plan_candidate_limit_var = tk.StringVar(value="25")
    plan_candidate_limit_entry = ttk.Entry(plan_rank_body, textvariable=plan_candidate_limit_var, width=8, state="disabled")
    plan_candidate_limit_entry.pack(anchor="w", pady=(2, 0))

    ttk.Label(plan_rank_body, text="TOP N par:", foreground=colors.MUTED).pack(anchor="w", pady=(10, 0))
    plan_top_n_var = tk.StringVar(value="20")
    plan_top_n_entry = ttk.Entry(plan_rank_body, textvariable=plan_top_n_var, width=8, state="disabled")
    plan_top_n_entry.pack(anchor="w", pady=(2, 0))

    # Goals / diversity constraints
    plan_diversity_frame = ttk.LabelFrame(plan_rank_body, text="Cele różnorodności")
    plan_diversity_frame.pack(anchor="w", fill=tk.X, pady=(12, 0))
    plan_diversity_body = ttk.Frame(plan_diversity_frame)
    plan_diversity_body.pack(fill=tk.X, padx=10, pady=10)

    plan_goal_mean_enabled_var = tk.BooleanVar(value=False)
    plan_goal_mean_enabled_cb = ttk.Checkbutton(
        plan_diversity_body,
        text="Cel: średnie F potomka <=",
        variable=plan_goal_mean_enabled_var,
        state="disabled",
    )
    plan_goal_mean_enabled_cb.pack(anchor="w")
    plan_goal_mean_F_var = tk.StringVar(value="0.05")
    plan_goal_mean_F_entry = ttk.Entry(plan_diversity_body, textvariable=plan_goal_mean_F_var, width=8, state="disabled")
    plan_goal_mean_F_entry.pack(anchor="w", pady=(2, 0), padx=(18, 0))

    plan_goal_max_enabled_var = tk.BooleanVar(value=False)
    plan_goal_max_enabled_cb = ttk.Checkbutton(
        plan_diversity_body,
        text="Cel: tylko pary z F potomka <=",
        variable=plan_goal_max_enabled_var,
        state="disabled",
    )
    plan_goal_max_enabled_cb.pack(anchor="w", pady=(8, 0))
    plan_goal_max_F_var = tk.StringVar(value="0.10")
    plan_goal_max_F_entry = ttk.Entry(plan_diversity_body, textvariable=plan_goal_max_F_var, width=8, state="disabled")
    plan_goal_max_F_entry.pack(anchor="w", pady=(2, 0), padx=(18, 0))

    ttk.Label(plan_diversity_body, text="Limit użyć samicy (max):", foreground=colors.MUTED).pack(anchor="w", pady=(10, 0))
    plan_max_uses_dam_var = tk.StringVar(value="3")
    plan_max_uses_dam_entry = ttk.Entry(plan_diversity_body, textvariable=plan_max_uses_dam_var, width=8, state="disabled")
    plan_max_uses_dam_entry.pack(anchor="w", pady=(2, 0), padx=(18, 0))

    ttk.Label(plan_diversity_body, text="Limit użyć samca (max):", foreground=colors.MUTED).pack(anchor="w", pady=(8, 0))
    plan_max_uses_sire_var = tk.StringVar(value="3")
    plan_max_uses_sire_entry = ttk.Entry(plan_diversity_body, textvariable=plan_max_uses_sire_var, width=8, state="disabled")
    plan_max_uses_sire_entry.pack(anchor="w", pady=(2, 0), padx=(18, 0))

    plan_suggest_result_var = tk.StringVar(value="")
    ttk.Label(plan_rank_body, textvariable=plan_suggest_result_var, foreground=colors.TEXT).pack(anchor="w", pady=(10, 0))

    def on_suggest_pairs() -> None:
        people = state.get("people")
        df_std = state.get("df_std")
        if not people or df_std is None:
            messagebox.showinfo("Info", "Najpierw wczytaj bazę.")
            return

        try:
            top_n = int(str(plan_top_n_var.get()).strip())
        except Exception:
            top_n = 20
        if top_n <= 0:
            top_n = 20

        try:
            max_back = None if bool(plan_risk_unbounded_var.get()) else int(str(plan_risk_depth_var.get()).strip())
        except Exception:
            messagebox.showerror("Błąd", "Max pokoleń musi być liczbą całkowitą.")
            return
        if max_back is not None and max_back < 0:
            messagebox.showerror("Błąd", "Max pokoleń nie może być ujemne.")
            return

        plan_suggest_result_var.set("Liczenie…")
        _set_plan_status("Liczenie rankingów par…")
        try:
            # Diversity constraints
            try:
                max_dam_uses = int(str(plan_max_uses_dam_var.get()).strip())
            except Exception:
                max_dam_uses = 3
            try:
                max_sire_uses = int(str(plan_max_uses_sire_var.get()).strip())
            except Exception:
                max_sire_uses = 3
            max_dam_uses = max(1, max_dam_uses)
            max_sire_uses = max(1, max_sire_uses)

            goal_mean_enabled = bool(plan_goal_mean_enabled_var.get())
            goal_max_enabled = bool(plan_goal_max_enabled_var.get())
            try:
                goal_mean_F = float(str(plan_goal_mean_F_var.get()).strip())
            except Exception:
                goal_mean_F = 0.05
            try:
                goal_max_F = float(str(plan_goal_max_F_var.get()).strip())
            except Exception:
                goal_max_F = 0.10

            try:
                amin = int(str(plan_min_age_var.get()).strip())
                amax = int(str(plan_max_age_var.get()).strip())
            except Exception:
                amin, amax = 0, 80
            try:
                cand_limit = int(str(plan_candidate_limit_var.get()).strip())
            except Exception:
                cand_limit = 25

            pair_result = suggest_pairs_with_constraints(
                df_std=df_std,  # type: ignore[arg-type]
                people=people,  # type: ignore[arg-type]
                min_age=amin,
                max_age=amax,
                line_mode=str(plan_line_filter_var.get()),
                candidate_limit=cand_limit,
                top_n=top_n,
                max_generations_back=max_back,
                max_dam_uses=max_dam_uses,
                max_sire_uses=max_sire_uses,
                goal_max_enabled=goal_max_enabled,
                goal_max_F=goal_max_F,
            )

            accepted = pair_result.suggestions

            _clear_plan_tree()
            for row in accepted:
                plan_tree.insert(
                    "",
                    "end",
                    values=(
                        row.dam_id,
                        row.dam_line,
                        row.dam_age,
                        row.sire_id,
                        row.sire_line,
                        row.sire_age,
                        f"{row.offspring_F:.6f}",
                    ),
                )

            if accepted:
                mean_F = pair_result.mean_F
                max_F = pair_result.max_F
                warn_mean = goal_mean_enabled and mean_F > goal_mean_F
                warn_txt = " UWAGA: średnie F przekracza cel." if warn_mean else ""
                plan_suggest_result_var.set(
                    f"Gotowe: TOP {len(accepted)}/{top_n} par. Śr. F={mean_F:.6f}, max F={max_F:.6f}.{warn_txt}"
                )
            else:
                plan_suggest_result_var.set("Brak par spełniających ograniczenia (TOP=0).")
            _set_plan_status("Gotowe: podpowiedź par.")
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się policzyć rankingów par: {e}")
            plan_suggest_result_var.set("")
            _set_plan_status("Błąd: ranking par.")

    plan_suggest_btn = ttk.Button(plan_rank_body, text="Podpowiedz pary", command=on_suggest_pairs, state="disabled")
    plan_suggest_btn.pack(anchor="w", pady=(12, 0))

    # Results tree
    plan_tree_columns = ("dam_id", "dam_line", "dam_age", "sire_id", "sire_line", "sire_age", "off_F")
    plan_tree = ttk.Treeview(
        plan_tree_upper,
        columns=plan_tree_columns,
        show="headings",
        height=10,
        style="Treeview",
    )
    plan_tree.heading("dam_id", text="Samica (ID)")
    plan_tree.heading("dam_line", text="Samica (LB/LC)")
    plan_tree.heading("dam_age", text="Samica (wiek)")
    plan_tree.heading("sire_id", text="Samiec (ID)")
    plan_tree.heading("sire_line", text="Samiec (LB/LC)")
    plan_tree.heading("sire_age", text="Samiec (wiek)")
    plan_tree.heading("off_F", text="Ryzyko: F potomka")

    for col in plan_tree_columns:
        plan_tree.column(col, width=120, anchor="w")

    plan_tree_vsb = ttk.Scrollbar(plan_tree_upper, orient="vertical", command=plan_tree.yview)
    plan_tree.configure(yscrollcommand=plan_tree_vsb.set)
    plan_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    plan_tree_vsb.pack(side=tk.RIGHT, fill=tk.Y)

    PLAN_HYPO_ID = "__PLAN_HYPO_OFFSPRING__"

    plan_pedigree_win_btn = ttk.Button(
        plan_pair_btn_row,
        text="Pokaż rodowód potomka w nowym oknie",
        state="disabled",
    )
    plan_pedigree_win_btn.pack(side=tk.LEFT)

    def _clear_plan_tree() -> None:
        for item in plan_tree.get_children():
            plan_tree.delete(item)
        _clear_plan_pair_detail()

    def _clear_plan_pair_detail() -> None:
        plan_sel_pair["dam"] = None
        plan_sel_pair["sire"] = None
        plan_pedigree_win_btn.configure(state="disabled")
        plan_pair_stats_text.configure(state="normal")
        plan_pair_stats_text.delete("1.0", tk.END)
        plan_pair_stats_text.insert(
            "1.0",
            "Zaznacz wiersz w tabeli par (klik), aby zobaczyć statystyki.\n"
            "Rodowód hipotetycznego potomka otworzysz przyciskiem poniżej.\n",
        )
        plan_pair_stats_text.configure(state="disabled")

    def _plan_f_max_back() -> int | None:
        try:
            return None if bool(plan_risk_unbounded_var.get()) else int(str(plan_risk_depth_var.get()).strip())
        except Exception:
            return 4

    def _plan_pedigree_plot_depth() -> int:
        """Ograniczona głębokość grafu (czytelny podgląd w panelu)."""
        mb = _plan_f_max_back()
        if mb is None:
            return 4
        return max(1, min(int(mb), 4))

    def _pci_bundle_for_plan(pid: str, pmap: dict[str, Person]) -> tuple[int, float, float]:
        MG, EG, PCI = 0, 0.0, 0.0
        try:
            levels = get_ancestor_levels_unbounded(person_id=pid, people=pmap)
            by_gen: dict[int, int] = {}
            for _aid, lvl in levels.items():
                if lvl is None:
                    continue
                try:
                    g = int(lvl)
                except Exception:
                    continue
                if g <= 0:
                    continue
                by_gen[g] = by_gen.get(g, 0) + 1
            if by_gen:
                MG = int(max(by_gen.keys()))
                pci_sum = 0.0
                for g in range(1, MG + 1):
                    a_g = int(by_gen.get(g, 0))
                    pcl_g = float(a_g) / float(2**g)
                    EG += pcl_g
                    pci_sum += pcl_g
                PCI = pci_sum / float(MG) if MG > 0 else 0.0
        except Exception:
            pass
        return MG, EG, PCI

    def _open_plan_offspring_pedigree_window() -> None:
        people = state.get("people")
        dam_id = plan_sel_pair.get("dam")
        sire_id = plan_sel_pair.get("sire")
        if not people or not dam_id or not sire_id:
            messagebox.showinfo("Info", "Najpierw zaznacz parę w tabeli wyników.")
            return
        if dam_id not in people or sire_id not in people:
            messagebox.showerror("Błąd", "Wybrane ID nie występują w wczytanych danych.")
            return

        plot_depth = _plan_pedigree_plot_depth()
        people_h = dict(people)
        people_h[PLAN_HYPO_ID] = Person(
            id=PLAN_HYPO_ID,
            name="potomek (hipotetyczny)",
            sex=None,
            line=None,
            father_id=sire_id,
            mother_id=dam_id,
            birth_year=None,
        )
        try:
            levels, edges = get_ancestor_levels_and_edges(
                person_id=PLAN_HYPO_ID,
                depth=plot_depth,
                people=people_h,
            )
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

            fig_p = plot_ancestor_pedigree(
                person_id=PLAN_HYPO_ID,
                levels=levels,
                edges=edges,
                people=people_h,
                readable_mode=True,
                enable_click_highlight=False,
                on_node_click=None,
            )
        except Exception as exc:
            messagebox.showerror("Błąd", f"Nie udało się narysować rodowodu: {exc}")
            return

        win = tk.Toplevel(root)
        win.title(f"Rodowód potomka (hipotetyczny) — {sire_id} × {dam_id}")
        win.geometry("1040x720")
        win.minsize(640, 480)
        win.configure(bg=colors.APP_BG)

        top_bar = ttk.Frame(win, padding=(8, 8))
        top_bar.pack(side=tk.TOP, fill=tk.X)
        ttk.Label(
            top_bar,
            text=f"Hipotetyczny potomek: samiec {sire_id} × samica {dam_id}  |  do {plot_depth} pokoleń wstecz",
            foreground=colors.MUTED,
        ).pack(side=tk.LEFT)
        ttk.Button(
            top_bar,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig_p: _save_figure_as_jpeg(f, default_basename=f"plan_para_{dam_id}_{sire_id}"),
        ).pack(side=tk.RIGHT, padx=(8, 0))
        ttk.Button(top_bar, text="Zamknij", command=win.destroy).pack(side=tk.RIGHT)

        plot_host = ttk.Frame(win)
        plot_host.pack(side=tk.TOP, fill=tk.BOTH, expand=True, padx=8, pady=(0, 8))
        canvas_p = FigureCanvasTkAgg(fig_p, master=plot_host)
        canvas_p.draw()
        canvas_p.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    plan_pedigree_win_btn.configure(command=_open_plan_offspring_pedigree_window)

    def _on_plan_tree_select(_evt: object | None = None) -> None:
        people = state.get("people")
        if not people:
            _clear_plan_pair_detail()
            return
        sel = plan_tree.selection()
        if not sel:
            _clear_plan_pair_detail()
            return
        vals = plan_tree.item(sel[0], "values")
        if len(vals) < 7:
            _clear_plan_pair_detail()
            return
        dam_id = str(vals[0]).strip()
        dam_line = str(vals[1]).strip()
        dam_age = str(vals[2]).strip()
        sire_id = str(vals[3]).strip()
        sire_line = str(vals[4]).strip()
        sire_age = str(vals[5]).strip()
        off_f_tbl = str(vals[6]).strip()
        if not dam_id or not sire_id:
            _clear_plan_pair_detail()
            return
        if dam_id not in people or sire_id not in people:
            _clear_plan_pair_detail()
            return

        max_back = _plan_f_max_back()
        plot_depth = _plan_pedigree_plot_depth()
        try:
            F_off = wright_offspring_inbreeding_F_from_parents(
                father_id=sire_id,
                mother_id=dam_id,
                people=people,
                max_generations_back=max_back,
            )
        except Exception:
            F_off = float("nan")
        try:
            F_dam = wright_inbreeding_F(person_id=dam_id, people=people, max_generations_back=max_back).F
        except Exception:
            F_dam = float("nan")
        try:
            F_sire = wright_inbreeding_F(person_id=sire_id, people=people, max_generations_back=max_back).F
        except Exception:
            F_sire = float("nan")

        MG_d, EG_d, PCI_d = _pci_bundle_for_plan(dam_id, people)
        MG_s, EG_s, PCI_s = _pci_bundle_for_plan(sire_id, people)

        d_nm = getattr(people.get(dam_id), "name", None) or "-"
        s_nm = getattr(people.get(sire_id), "name", None) or "-"
        mb_note = "bez limitu (do founderów)" if max_back is None else f"max {max_back} pokoleń"

        lines = [
            f"Samica: {dam_id} ({d_nm})  |  linia: {dam_line}  |  wiek (z rankingu): {dam_age}",
            f"Samiec: {sire_id} ({s_nm})  |  linia: {sire_line}  |  wiek (z rankingu): {sire_age}",
            "",
            f"Parametry F jak w panelu ryzyka ({mb_note}).",
            f"F potomka Φ(samiec, samica) = {F_off:.6f}  (w tabeli: {off_f_tbl})",
            f"F samicy = {F_dam:.6f}  |  F samca = {F_sire:.6f}",
            "",
            f"Kompletność rodowodu (ANC, bez limitu) — samica: MG={MG_d}, EG={EG_d:.4f}, PCI={PCI_d:.4f}",
            f"Kompletność rodowodu (ANC, bez limitu) — samiec: MG={MG_s}, EG={EG_s:.4f}, PCI={PCI_s:.4f}",
            "",
            f"Graf rodowodu potomka: do {plot_depth} pokoleń wstecz — otwórz przyciskiem „Pokaż rodowód potomka w nowym oknie”.",
        ]
        plan_sel_pair["dam"] = dam_id
        plan_sel_pair["sire"] = sire_id
        plan_pedigree_win_btn.configure(state="normal")

        plan_pair_stats_text.configure(state="normal")
        plan_pair_stats_text.delete("1.0", tk.END)
        plan_pair_stats_text.insert("1.0", "\n".join(lines))
        plan_pair_stats_text.configure(state="disabled")

    plan_tree.bind("<<TreeviewSelect>>", _on_plan_tree_select, add="+")
    _clear_plan_pair_detail()

    # -------------------------
    # Settings (domyślne zachowanie UI)
    # -------------------------
    ttk.Button(
        tab_settings,
        text="Pomoc: znaczenie ustawień",
        command=lambda: show_help_window(root, "Konfiguracja", hc.SECTION_SETTINGS),
    ).pack(anchor="w", pady=(0, 6))

    settings_divider = ttk.Separator(tab_settings, orient="horizontal")
    settings_divider.pack(fill=tk.X, pady=(2, 10))

    # Rodowód (graf przodków)
    ttk.Label(tab_settings, text="Rodowód (graf przodków)", font=tk_font(12, bold=True)).pack(anchor="w")

    settings_anc_readable_var = tk.BooleanVar(value=True)
    # Domyślnie pokazujemy pole "Max pokoleń" (czyli limit jest aktywny).
    settings_anc_unbounded_var = tk.BooleanVar(value=False)
    settings_anc_click_var = tk.BooleanVar(value=True)
    settings_anc_depth_var = tk.StringVar(value="4")

    ttk.Checkbutton(
        tab_settings,
        text="Domyślnie: tryb czytelny (mniej etykiet)",
        variable=settings_anc_readable_var,
    ).pack(anchor="w", pady=(6, 0))

    ttk.Checkbutton(
        tab_settings,
        text="Domyślnie: bez limitu (do founderów)",
        variable=settings_anc_unbounded_var,
    ).pack(anchor="w", pady=(6, 0))

    ttk.Checkbutton(
        tab_settings,
        text="Domyślnie: klik w węzeł (podświetlenie + aktualizacja linii)",
        variable=settings_anc_click_var,
    ).pack(anchor="w", pady=(6, 0))

    anc_depth_row = ttk.Frame(tab_settings)
    anc_depth_row.pack(anchor="w", pady=(6, 0))
    ttk.Label(anc_depth_row, text="Domyślny max pokoleń (gdy limit):").pack(side=tk.LEFT)
    settings_anc_depth_entry = ttk.Entry(anc_depth_row, textvariable=settings_anc_depth_var, width=8)
    settings_anc_depth_entry.pack(side=tk.LEFT, padx=(8, 0))

    def _sync_settings_anc_depth_state() -> None:
        st = "disabled" if bool(settings_anc_unbounded_var.get()) else "normal"
        settings_anc_depth_entry.configure(state=st)

    _sync_settings_anc_depth_state()
    settings_anc_unbounded_var.trace_add("write", lambda *_args: _sync_settings_anc_depth_state())

    settings_divider2 = ttk.Separator(tab_settings, orient="horizontal")
    settings_divider2.pack(fill=tk.X, pady=(14, 10))

    # Analizy: inbred (wybrany osobnik)
    ttk.Label(tab_settings, text="Analiza indywidualna (Inbred F)", font=tk_font(12, bold=True)).pack(anchor="w")
    settings_inb_unbounded_var = tk.BooleanVar(value=True)
    settings_inb_depth_var = tk.StringVar(value="4")

    ttk.Checkbutton(
        tab_settings,
        text="Domyślnie: bez limitu (do founderów)",
        variable=settings_inb_unbounded_var,
    ).pack(anchor="w", pady=(6, 0))

    inb_depth_row = ttk.Frame(tab_settings)
    inb_depth_row.pack(anchor="w", pady=(6, 0))
    ttk.Label(inb_depth_row, text="Domyślny max pokoleń (gdy limit):").pack(side=tk.LEFT)
    settings_inb_depth_entry = ttk.Entry(inb_depth_row, textvariable=settings_inb_depth_var, width=8)
    settings_inb_depth_entry.pack(side=tk.LEFT, padx=(8, 0))

    def _sync_settings_inb_depth_state() -> None:
        st = "disabled" if bool(settings_inb_unbounded_var.get()) else "normal"
        settings_inb_depth_entry.configure(state=st)

    _sync_settings_inb_depth_state()
    settings_inb_unbounded_var.trace_add("write", lambda *_args: _sync_settings_inb_depth_state())

    settings_divider3 = ttk.Separator(tab_settings, orient="horizontal")
    settings_divider3.pack(fill=tk.X, pady=(14, 10))

    # Populacja: trendy F
    ttk.Label(tab_settings, text="Populacja (F, trendy w czasie)", font=tk_font(12, bold=True)).pack(anchor="w")
    settings_pop_unbounded_var = tk.BooleanVar(value=False)
    settings_pop_depth_var = tk.StringVar(value="4")

    ttk.Checkbutton(
        tab_settings,
        text="Domyślnie: bez limitu (do founderów)",
        variable=settings_pop_unbounded_var,
    ).pack(anchor="w", pady=(6, 0))

    pop_depth_row = ttk.Frame(tab_settings)
    pop_depth_row.pack(anchor="w", pady=(6, 0))
    ttk.Label(pop_depth_row, text="Domyślny max pokoleń (gdy limit):").pack(side=tk.LEFT)
    settings_pop_depth_entry = ttk.Entry(pop_depth_row, textvariable=settings_pop_depth_var, width=8)
    settings_pop_depth_entry.pack(side=tk.LEFT, padx=(8, 0))

    def _sync_settings_pop_depth_state() -> None:
        st = "disabled" if bool(settings_pop_unbounded_var.get()) else "normal"
        settings_pop_depth_entry.configure(state=st)

    _sync_settings_pop_depth_state()
    settings_pop_unbounded_var.trace_add("write", lambda *_args: _sync_settings_pop_depth_state())

    settings_divider4 = ttk.Separator(tab_settings, orient="horizontal")
    settings_divider4.pack(fill=tk.X, pady=(14, 10))

    # Raporty
    ttk.Label(tab_settings, text="Raporty", font=tk_font(12, bold=True)).pack(anchor="w")
    ttk.Checkbutton(
        tab_settings,
        text="Domyślnie: dołączaj wykresy w eksporcie raportu (DOCX/PDF)",
        variable=rep_default_include_plots_var,
    ).pack(anchor="w", pady=(6, 0))

    settings_divider5 = ttk.Separator(tab_settings, orient="horizontal")
    settings_divider5.pack(fill=tk.X, pady=(14, 10))

    # Plan hodowlany
    ttk.Label(tab_settings, text="Plan hodowlany", font=tk_font(12, bold=True)).pack(anchor="w")

    settings_plan_unbounded_var = tk.BooleanVar(value=False)
    settings_plan_depth_var = tk.StringVar(value="4")
    settings_plan_min_age_var = tk.StringVar(value="0")
    settings_plan_max_age_var = tk.StringVar(value="80")
    settings_plan_candidate_limit_var = tk.StringVar(value="25")
    settings_plan_top_n_var = tk.StringVar(value="20")
    settings_plan_max_dam_uses_var = tk.StringVar(value="3")
    settings_plan_max_sire_uses_var = tk.StringVar(value="3")
    settings_plan_goal_mean_enabled_var = tk.BooleanVar(value=False)
    settings_plan_goal_mean_f_var = tk.StringVar(value="0.05")
    settings_plan_goal_max_enabled_var = tk.BooleanVar(value=False)
    settings_plan_goal_max_f_var = tk.StringVar(value="0.10")

    ttk.Checkbutton(
        tab_settings,
        text="Domyślnie: ryzyko bez limitu (do founderów)",
        variable=settings_plan_unbounded_var,
    ).pack(anchor="w", pady=(6, 0))

    plan_depth_settings_row = ttk.Frame(tab_settings)
    plan_depth_settings_row.pack(anchor="w", pady=(6, 0))
    ttk.Label(plan_depth_settings_row, text="Domyślny max pokoleń (Plan hodowlany):").pack(side=tk.LEFT)
    settings_plan_depth_entry = ttk.Entry(plan_depth_settings_row, textvariable=settings_plan_depth_var, width=8)
    settings_plan_depth_entry.pack(side=tk.LEFT, padx=(8, 0))

    def _sync_settings_plan_depth_state() -> None:
        st = "disabled" if bool(settings_plan_unbounded_var.get()) else "normal"
        settings_plan_depth_entry.configure(state=st)

    _sync_settings_plan_depth_state()
    settings_plan_unbounded_var.trace_add("write", lambda *_args: _sync_settings_plan_depth_state())

    plan_age_settings_row = ttk.Frame(tab_settings)
    plan_age_settings_row.pack(anchor="w", pady=(6, 0))
    ttk.Label(plan_age_settings_row, text="Wiek min/max (Plan):").pack(side=tk.LEFT)
    ttk.Entry(plan_age_settings_row, textvariable=settings_plan_min_age_var, width=6).pack(side=tk.LEFT, padx=(8, 4))
    ttk.Entry(plan_age_settings_row, textvariable=settings_plan_max_age_var, width=6).pack(side=tk.LEFT, padx=(0, 0))

    plan_limits_settings_row = ttk.Frame(tab_settings)
    plan_limits_settings_row.pack(anchor="w", pady=(6, 0))
    ttk.Label(plan_limits_settings_row, text="Limit kandydatów / TOP N:").pack(side=tk.LEFT)
    ttk.Entry(plan_limits_settings_row, textvariable=settings_plan_candidate_limit_var, width=6).pack(side=tk.LEFT, padx=(8, 4))
    ttk.Entry(plan_limits_settings_row, textvariable=settings_plan_top_n_var, width=6).pack(side=tk.LEFT, padx=(0, 0))

    plan_uses_settings_row = ttk.Frame(tab_settings)
    plan_uses_settings_row.pack(anchor="w", pady=(6, 0))
    ttk.Label(plan_uses_settings_row, text="Max użyć (samica / samiec):").pack(side=tk.LEFT)
    ttk.Entry(plan_uses_settings_row, textvariable=settings_plan_max_dam_uses_var, width=6).pack(side=tk.LEFT, padx=(8, 4))
    ttk.Entry(plan_uses_settings_row, textvariable=settings_plan_max_sire_uses_var, width=6).pack(side=tk.LEFT, padx=(0, 0))

    ttk.Checkbutton(
        tab_settings,
        text="Cel: włącz próg średniego F",
        variable=settings_plan_goal_mean_enabled_var,
    ).pack(anchor="w", pady=(6, 0))
    goal_mean_row = ttk.Frame(tab_settings)
    goal_mean_row.pack(anchor="w", pady=(2, 0))
    ttk.Label(goal_mean_row, text="Domyślne średnie F <= ").pack(side=tk.LEFT)
    ttk.Entry(goal_mean_row, textvariable=settings_plan_goal_mean_f_var, width=8).pack(side=tk.LEFT, padx=(8, 0))

    ttk.Checkbutton(
        tab_settings,
        text="Cel: włącz próg maksymalnego F",
        variable=settings_plan_goal_max_enabled_var,
    ).pack(anchor="w", pady=(6, 0))
    goal_max_row = ttk.Frame(tab_settings)
    goal_max_row.pack(anchor="w", pady=(2, 0))
    ttk.Label(goal_max_row, text="Domyślne max F <= ").pack(side=tk.LEFT)
    ttk.Entry(goal_max_row, textvariable=settings_plan_goal_max_f_var, width=8).pack(side=tk.LEFT, padx=(8, 0))

    # -------------------------
    # Shared app state
    # -------------------------
    state: dict[str, object] = {
        "df_std": None,
        "people": None,
        "line_memberships": {},
        "df_raw": None,
        "raw_filename": None,
    }
    # Czy kontrolki po wczytaniu bazy są aktywne (nie wnioskować ze stanu innych widgetów — macOS/ttk bywa zawodne).
    controls_enabled_ref: list[bool] = [False]

    # -------------------------
    # Populacja tab (podstawowe metryki)
    # -------------------------
    pop_nb = ttk.Notebook(tab_population)
    pop_nb.pack(side=tk.TOP, fill=tk.BOTH, expand=True, pady=(14, 0))

    tab_pop_basic = ttk.Frame(pop_nb)
    tab_pop_meta = ttk.Frame(pop_nb)
    tab_pop_charts = ttk.Frame(pop_nb)

    pop_nb.add(tab_pop_basic, text="Podstawowe statystyki")
    pop_nb.add(tab_pop_meta, text="Wskaźniki genetyczne i demograficzne")
    pop_nb.add(tab_pop_charts, text="Wykresy")

    pop_frame = tab_pop_basic

    pop_total_var = tk.StringVar(value="-")
    pop_founders_count_var = tk.StringVar(value="-")
    pop_known_parents_var = tk.StringVar(value="-")
    pop_known_father_var = tk.StringVar(value="-")
    pop_known_mother_var = tk.StringVar(value="-")
    pop_lines_var = tk.StringVar(value="-")

    # Parametry liczenia inbredu (F) dla metryk populacyjnych (bez wpływu na "Analizy").
    pop_depth_inb_var = tk.StringVar(value="4")
    pop_unbounded_inb_var = tk.BooleanVar(value=False)

    # Synchronizacja domyślnych ustawień dla wykresów F w zakładce "Populacja".
    pop_depth_inb_var.set(settings_pop_depth_var.get())
    pop_unbounded_inb_var.set(bool(settings_pop_unbounded_var.get()))
    settings_pop_unbounded_var.trace_add(
        "write", lambda *_args: pop_unbounded_inb_var.set(bool(settings_pop_unbounded_var.get()))
    )
    settings_pop_depth_var.trace_add(
        "write", lambda *_args: pop_depth_inb_var.set(settings_pop_depth_var.get())
    )

    # Genetyka populacyjna (TP/RP — zależnie od implementacji).
    pop_f_e_var = tk.StringVar(value="—")
    pop_f_a_var = tk.StringVar(value="—")
    pop_bottleneck_var = tk.StringVar(value="—")
    pop_ne_var = tk.StringVar(value="—")
    pop_drift_var = tk.StringVar(value="—")
    pop_ria_overall_var = tk.StringVar(value="—")
    pop_mk_phi_var = tk.StringVar(value="—")
    pop_mk_r_var = tk.StringVar(value="—")
    pop_mk_note_var = tk.StringVar(value="")

    # Demografia / rodowód.
    pop_gi_mean_var = tk.StringVar(value="—")
    pop_gi_father_son_var = tk.StringVar(value="—")
    pop_gi_father_daughter_var = tk.StringVar(value="—")
    pop_gi_mother_son_var = tk.StringVar(value="—")
    pop_gi_mother_daughter_var = tk.StringVar(value="—")

    pop_family_count_var = tk.StringVar(value="—")
    pop_family_mean_size_var = tk.StringVar(value="—")

    pop_title_row = ttk.Frame(pop_frame)
    pop_title_row.pack(side=tk.TOP, fill=tk.X)
    ttk.Label(pop_title_row, text="Podstawowe statystyki:", foreground=colors.TEXT, font=tk_font(12, bold=True)).pack(
        side=tk.LEFT, anchor="w"
    )
    ttk.Button(
        pop_title_row,
        text="Pomoc: metryki i wykresy",
        command=lambda: show_help_window(
            root,
            "Populacja — interpretacja",
            hc.SECTION_POPULATION
            + "\n\n---\n\n"
            + hc.SECTION_REFERENCES
            + "\n\n---\n\n"
            + hc.all_charts_text(),
        ),
    ).pack(side=tk.RIGHT)
    pop_basic_grid = ttk.Frame(pop_frame)
    pop_basic_grid.pack(side=tk.TOP, fill=tk.X, pady=(6, 0))
    pop_basic_grid.columnconfigure(0, weight=1)
    pop_basic_grid.columnconfigure(1, weight=1)

    ttk.Label(pop_basic_grid, textvariable=pop_total_var, foreground=colors.TEXT).grid(row=0, column=0, sticky="w")
    ttk.Label(pop_basic_grid, textvariable=pop_founders_count_var, foreground=colors.TEXT).grid(row=1, column=0, sticky="w")
    ttk.Label(pop_basic_grid, textvariable=pop_known_parents_var, foreground=colors.TEXT).grid(row=2, column=0, sticky="w")

    ttk.Label(pop_basic_grid, textvariable=pop_known_father_var, foreground=colors.TEXT).grid(row=0, column=1, sticky="w")
    ttk.Label(pop_basic_grid, textvariable=pop_known_mother_var, foreground=colors.TEXT).grid(row=1, column=1, sticky="w")
    ttk.Label(pop_basic_grid, textvariable=pop_lines_var, foreground=colors.TEXT).grid(row=2, column=1, sticky="w", pady=(10, 0))

    # --- Parametry liczenia F (dla wykresów populacyjnych) ---
    pop_inb_param_frame = ttk.Frame(pop_frame)
    pop_inb_param_frame.pack(side=tk.TOP, fill=tk.X, pady=(6, 0))
    ttk.Label(
        pop_inb_param_frame,
        text="Parametry F (dla wykresów populacyjnych):",
        foreground=colors.TEXT,
        font=tk_font(12, bold=True),
    ).pack(anchor="w")

    pop_unbounded_inb_cb = ttk.Checkbutton(
        pop_inb_param_frame,
        text="Bez ograniczenia (do founderów)",
        variable=pop_unbounded_inb_var,
    )
    pop_unbounded_inb_cb.pack(anchor="w", pady=(6, 0))

    pop_depth_inb_row = ttk.Frame(pop_inb_param_frame)
    pop_depth_inb_row.pack(anchor="w", pady=(6, 0))
    ttk.Label(pop_depth_inb_row, text="Max pokoleń:").pack(side=tk.LEFT)
    pop_depth_inb_entry = ttk.Entry(pop_depth_inb_row, textvariable=pop_depth_inb_var, width=10)
    pop_depth_inb_entry.pack(side=tk.LEFT, padx=(8, 0))

    def _sync_pop_inb_depth_state() -> None:
        st = "disabled" if bool(pop_unbounded_inb_var.get()) else "normal"
        pop_depth_inb_entry.configure(state=st)

    _sync_pop_inb_depth_state()
    pop_unbounded_inb_var.trace_add("write", lambda *_args: _sync_pop_inb_depth_state())

    # --- Genetyka i demografia populacji (skrót) ---
    pop_frame = tab_pop_meta
    ttk.Label(
        pop_frame,
        text="Wskaźniki genetyczne i demograficzne",
        foreground=colors.TEXT,
        font=tk_font(13, bold=True),
    ).pack(anchor="w", pady=(12, 2))
    ttk.Label(
        pop_frame,
        text="Podsumowanie metryk populacji — wartości liczbowe z bieżącej bazy (jak przy wykresach F).",
        foreground=colors.MUTED,
        font=tk_font(9),
        wraplength=920,
        justify="left",
    ).pack(anchor="w", pady=(0, 6))

    pop_meta_stack = ttk.Frame(pop_frame)
    pop_meta_stack.pack(side=tk.TOP, fill=tk.X, pady=(4, 0))

    lf_founders = ttk.Labelframe(
        pop_meta_stack,
        text=" Założyciele, bottleneck, dryf, N_e, RIA ",
        padding=(12, 10),
        style="PopMeta.TLabelframe",
    )
    lf_founders.pack(fill=tk.X, pady=(0, 8))
    founders_row = ttk.Frame(lf_founders)
    founders_row.pack(fill=tk.X)
    gf_left = ttk.Frame(founders_row)
    gf_left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
    gf_right = ttk.Frame(founders_row)
    gf_right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    for _gf in (gf_left, gf_right):
        _gf.columnconfigure(0, weight=2)
        _gf.columnconfigure(1, weight=1)
    _pop_stat_row(gf_left, 0, "f_e (efektywna liczba założycieli)", pop_f_e_var, colors=colors)
    _pop_stat_row(gf_left, 1, "Bottleneck f_e/f_a", pop_bottleneck_var, colors=colors)
    _pop_stat_row(gf_left, 2, "Dryf f_e/f_ge (f_ge=f_e, aproksymacja)", pop_drift_var, colors=colors)
    _pop_stat_row(gf_right, 0, "f_a (efektywna liczba przodków)", pop_f_a_var, colors=colors)
    _pop_stat_row(
        gf_right,
        1,
        "N_e (efektywna wielkość populacji, z trendu średniego F)",
        pop_ne_var,
        colors=colors,
        value_wrap=280,
    )
    _pop_stat_row(gf_right, 2, "RIA ogółem (F>0)", pop_ria_overall_var, colors=colors)

    lf_mk = ttk.Labelframe(
        pop_meta_stack,
        text=" Mean kinship (Φ, R po parach i≠j) ",
        padding=(12, 10),
        style="PopMeta.TLabelframe",
    )
    lf_mk.pack(fill=tk.X, pady=(0, 8))
    mk_vals = ttk.Frame(lf_mk)
    mk_vals.pack(fill=tk.X)
    mk_vals.columnconfigure(0, weight=1)
    mk_vals.columnconfigure(1, weight=1)
    mk_phi_col = ttk.Frame(mk_vals)
    mk_phi_col.grid(row=0, column=0, sticky="nw", padx=(0, 16))
    ttk.Label(
        mk_phi_col,
        text="śr. Φ (Malecot, pary i≠j)",
        font=tk_font(11),
        foreground=colors.TEXT,
        anchor="w",
    ).pack(anchor="w")
    ttk.Label(
        mk_phi_col,
        textvariable=pop_mk_phi_var,
        font=tk_font(13, bold=True),
        foreground=colors.TEXT,
        anchor="w",
    ).pack(anchor="w", pady=(2, 0))
    mk_r_col = ttk.Frame(mk_vals)
    mk_r_col.grid(row=0, column=1, sticky="nw")
    ttk.Label(
        mk_r_col,
        text="śr. R = 2Φ (Wright, autosomy)",
        font=tk_font(11),
        foreground=colors.TEXT,
        anchor="w",
    ).pack(anchor="w")
    ttk.Label(
        mk_r_col,
        textvariable=pop_mk_r_var,
        font=tk_font(13, bold=True),
        foreground=colors.TEXT,
        anchor="w",
    ).pack(anchor="w", pady=(2, 0))
    ttk.Label(
        lf_mk,
        textvariable=pop_mk_note_var,
        font=tk_font(10),
        foreground=colors.TEXT,
        wraplength=860,
        justify="left",
        anchor="w",
    ).pack(anchor="w", fill=tk.X, pady=(10, 0))

    lf_gi = ttk.Labelframe(
        pop_meta_stack,
        text=" Generation Interval (GI) i rodziny pełnego rodzeństwa ",
        padding=(12, 10),
        style="PopMeta.TLabelframe",
    )
    lf_gi.pack(fill=tk.X, pady=(0, 4))
    gi_row = ttk.Frame(lf_gi)
    gi_row.pack(fill=tk.X)
    gi_left = ttk.Frame(gi_row)
    gi_left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
    gi_right = ttk.Frame(gi_row)
    gi_right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    for _gi in (gi_left, gi_right):
        _gi.columnconfigure(0, weight=2)
        _gi.columnconfigure(1, weight=1)
    _pop_stat_row(gi_left, 0, "GI (Generation Interval, średnio), lata", pop_gi_mean_var, colors=colors)
    _pop_stat_row(gi_left, 1, "GI: Ojciec → syn", pop_gi_father_son_var, colors=colors)
    _pop_stat_row(gi_left, 2, "GI: Matka → syn", pop_gi_mother_son_var, colors=colors)
    _pop_stat_row(gi_left, 3, "Średnia wielkość rodziny (pełne rodzeństwo)", pop_family_mean_size_var, colors=colors)
    _pop_stat_row(gi_right, 0, "Liczba rodzin pełnego rodzeństwa", pop_family_count_var, colors=colors)
    _pop_stat_row(gi_right, 1, "GI: Ojciec → córka", pop_gi_father_daughter_var, colors=colors)
    _pop_stat_row(gi_right, 2, "GI: Matka → córka", pop_gi_mother_daughter_var, colors=colors)

    # -------------------------
    # Wykresy: każdy wykres = 1 zakładka
    # -------------------------
    pop_frame = tab_pop_charts
    pop_charts_frame = ttk.Frame(pop_frame)
    pop_charts_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True, pady=(14, 0))

    # Zamiast ttk.Notebook (który ma zwykle jeden wiersz zakładek i przy długich
    # nazwach nie pokazuje ich wszystkich), robimy dwa rzędy przycisków-tabów
    # przełączających jeden wspólny panel z treścią wykresu.
    pop_tabs_bar = ttk.Frame(pop_charts_frame)
    pop_tabs_bar.pack(side=tk.TOP, fill=tk.X, pady=(0, 6))
    pop_tabs_row1 = ttk.Frame(pop_tabs_bar)
    pop_tabs_row1.pack(side=tk.TOP, fill=tk.X)
    pop_tabs_row2 = ttk.Frame(pop_tabs_bar)
    pop_tabs_row2.pack(side=tk.TOP, fill=tk.X, pady=(4, 0))

    pop_plots_container = ttk.Frame(pop_charts_frame)
    pop_plots_container.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

    plot_tabs: list[tuple[str, ttk.Frame]] = []

    def _short_tab_text(title: str) -> str:
        s = str(title).strip()
        if len(s) <= 26:
            return s
        return s[:24] + "…"

    def _plot_tab(title: str) -> ttk.Frame:
        tab = ttk.Frame(pop_plots_container, padding=10)
        plot_tabs.append((title, tab))
        return tab

    # Urodzenia (płeć)
    tab_birth_sex = _plot_tab("Urodzenia według płci")
    ttk.Label(tab_birth_sex, text="Liczba osobników urodzonych w dekadach (płeć)", font=tk_font(12, bold=True)).pack(
        anchor="w"
    )
    ttk.Label(
        tab_birth_sex,
        text="Interpretacja: " + hc.CHART_BIRTH_SEX,
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_birth_sex_plot_area = ttk.Frame(tab_birth_sex)
    pop_birth_sex_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # Urodzenia (linie)
    tab_birth_line = _plot_tab("Urodzenia według linii (LB/LC)")
    ttk.Label(tab_birth_line, text="Liczba osobników urodzonych w dekadach (LB vs LC)", font=tk_font(12, bold=True)).pack(
        anchor="w"
    )
    ttk.Label(
        tab_birth_line,
        text="Interpretacja: " + hc.CHART_BIRTH_LINE,
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_birth_line_plot_area = ttk.Frame(tab_birth_line)
    pop_birth_line_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # Female/Male ratio
    tab_birth_ratio = _plot_tab("Stosunek płci (ur. ≥ 1900)")
    ttk.Label(tab_birth_ratio, text="Female/Male ratio urodzeń od 1900 roku", font=tk_font(12, bold=True)).pack(
        anchor="w"
    )
    ttk.Label(
        tab_birth_ratio,
        text="Interpretacja: " + hc.CHART_FM_RATIO,
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_birth_ratio_plot_area = ttk.Frame(tab_birth_ratio)
    pop_birth_ratio_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # GI (bar)
    tab_gi = _plot_tab("Średni interwał pokoleniowy (GI)")
    ttk.Label(tab_gi, text="Odstęp międzypokoleniowy (GI) — średni wiek rodziców", font=tk_font(12, bold=True)).pack(
        anchor="w"
    )
    ttk.Label(
        tab_gi,
        text="Interpretacja: " + hc.CHART_GI_BAR,
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_gi_plot_area = ttk.Frame(tab_gi)
    pop_gi_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # GI trend
    tab_gi_trend = _plot_tab("Trend interwału pokoleniowego")
    ttk.Label(tab_gi_trend, text="GI w czasie (trend) — dekady i 4 ścieżki", font=tk_font(12, bold=True)).pack(
        anchor="w"
    )
    ttk.Label(
        tab_gi_trend,
        text="Interpretacja: " + hc.CHART_GI_TREND,
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_gi_trend_plot_area = ttk.Frame(tab_gi_trend)
    pop_gi_trend_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # GI — 4 podokna (jak na przykładzie z 4 ścieżkami)
    tab_gi_4paths = _plot_tab("GI (4 ścieżki) — surowe + wygładzone")
    ttk.Label(tab_gi_4paths, text="GI w 4 wariantach (ojciec→syn, ojciec→córka, matka→syn, matka→córka)", font=tk_font(12, bold=True)).pack(
        anchor="w"
    )
    ttk.Label(
        tab_gi_4paths,
        text="Interpretacja: średnie GI w dekadach + wygładzenie (okrno ruchome).",
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_gi_4paths_plot_area = ttk.Frame(tab_gi_4paths)
    pop_gi_4paths_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # Rodziny pełne
    tab_family = _plot_tab("Kompletność struktury rodzin")
    ttk.Label(tab_family, text="Struktura rodzin pełnego rodzeństwa", font=tk_font(12, bold=True)).pack(anchor="w")
    ttk.Label(
        tab_family,
        text="Interpretacja: " + hc.CHART_FAMILY,
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_family_plot_area = ttk.Frame(tab_family)
    pop_family_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # Kompletność (płeć)
    tab_comp_sex = _plot_tab("Kompletność rodowodu — płeć")
    ttk.Label(tab_comp_sex, text="Kompletność rodowodu: MG / CG / EG wg płci", font=tk_font(12, bold=True)).pack(
        anchor="w"
    )
    ttk.Label(
        tab_comp_sex,
        text="Interpretacja: " + hc.CHART_COMP_SEX,
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_comp_sex_plot_area = ttk.Frame(tab_comp_sex)
    pop_comp_sex_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # Kompletność (linie)
    tab_comp_line = _plot_tab("Kompletność rodowodu — linie")
    ttk.Label(tab_comp_line, text="Kompletność rodowodu: MG / CG / EG wg linii LB / LC", font=tk_font(12, bold=True)).pack(
        anchor="w"
    )
    ttk.Label(
        tab_comp_line,
        text="Interpretacja: " + hc.CHART_COMP_LINE,
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_comp_line_plot_area = ttk.Frame(tab_comp_line)
    pop_comp_line_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # Inbred TP (płeć)
    tab_inb_year_sex = _plot_tab("Trend współczynnika F — płeć")
    ttk.Label(tab_inb_year_sex, text="Average F i RIA (%) w czasie — wg płci", font=tk_font(12, bold=True)).pack(
        anchor="w"
    )
    ttk.Label(
        tab_inb_year_sex,
        text="Interpretacja: " + hc.CHART_INBRED_TP_SEX,
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_inb_year_sex_plot_area = ttk.Frame(tab_inb_year_sex)
    pop_inb_year_sex_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # Inbred TP (linie)
    tab_inb_year_line = _plot_tab("Trend współczynnika F — linie")
    ttk.Label(tab_inb_year_line, text="Average F i RIA (%) w czasie — wg linii LB/LC", font=tk_font(12, bold=True)).pack(
        anchor="w"
    )
    ttk.Label(
        tab_inb_year_line,
        text="Interpretacja: " + hc.CHART_INBRED_TP_LINE,
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_inb_year_line_plot_area = ttk.Frame(tab_inb_year_line)
    pop_inb_year_line_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # F i RIA (TP) — surowe + wygładzone
    tab_inb_smooth = _plot_tab("F i RIA (TP) — surowe + wygładzone")
    ttk.Label(tab_inb_smooth, text="Average F oraz RIA (%) w czasie (TP ogółem)", font=tk_font(12, bold=True)).pack(
        anchor="w"
    )
    ttk.Label(
        tab_inb_smooth,
        text="Interpretacja: surowe punkty (czarne) i wygładzony trend (czerwony).",
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_inb_smooth_plot_area = ttk.Frame(tab_inb_smooth)
    pop_inb_smooth_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # Rozkład F w populacji
    tab_f_hist = _plot_tab("Rozkład F (osoby)")
    ttk.Label(tab_f_hist, text="Jak wygląda rozkład F w całej populacji", font=tk_font(12, bold=True)).pack(
        anchor="w"
    )
    ttk.Label(
        tab_f_hist,
        text="Hist. F (Wright) — wartości użyte do wykresu są zgodne z aktualnym limitem pokoleń.",
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_f_hist_plot_area = ttk.Frame(tab_f_hist)
    pop_f_hist_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # F vs rok urodzenia
    tab_f_scatter = _plot_tab("F vs rok urodzenia")
    ttk.Label(tab_f_scatter, text="F (Wright) w funkcji roku urodzenia", font=tk_font(12, bold=True)).pack(
        anchor="w"
    )
    ttk.Label(
        tab_f_scatter,
        text="Wykres punktowy: każda kropka = osobnik z liczonym F.",
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_f_scatter_plot_area = ttk.Frame(tab_f_scatter)
    pop_f_scatter_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # Founder contributions
    tab_founders = _plot_tab("Ranking wkładu założycieli (p_i)")
    ttk.Label(
        tab_founders,
        text=f"Wkład genetyczny założycieli (top {POP_FOUNDERS_PI_TOP_N} p_i)",
        font=tk_font(12, bold=True),
    ).pack(anchor="w")
    ttk.Label(
        tab_founders,
        text="Interpretacja: " + hc.CHART_FOUNDERS_PI,
        font=tk_font(8),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))
    pop_founders_plot_area = ttk.Frame(tab_founders)
    pop_founders_plot_area.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

    # --- Zbuduj dwurzędowy pasek zakładek (tytuły przycięte, jedna treść naraz) ---
    split_idx = (len(plot_tabs) + 1) // 2
    for idx, (title, tab_frame) in enumerate(plot_tabs):
        target_row = pop_tabs_row1 if idx < split_idx else pop_tabs_row2
        btn_txt = _short_tab_text(title)
        # width w znakach; dzięki krótszym etykietom unikamy overflow.
        btn = ttk.Button(
            target_row,
            text=btn_txt,
            width=26,
            command=lambda i=idx: _show_pop_plot_tab(i),
        )
        btn.pack(side=tk.LEFT, padx=(0, 6), pady=(2, 2))

    def _show_pop_plot_tab(idx: int) -> None:
        for _t, fr in plot_tabs:
            fr.pack_forget()
        plot_tabs[idx][1].pack(fill=tk.BOTH, expand=True)

    # Wybór domyślny: pierwsza zakładka.
    if plot_tabs:
        _show_pop_plot_tab(0)

    def _update_population_metrics(df_std) -> None:
        TEST_ID = "99999"

        def _clear_gen_demo_vars() -> None:
            pop_f_e_var.set("—")
            pop_f_a_var.set("—")
            pop_bottleneck_var.set("—")
            pop_ne_var.set("—")
            pop_drift_var.set("—")
            pop_ria_overall_var.set("—")
            pop_mk_phi_var.set("—")
            pop_mk_r_var.set("—")
            pop_mk_note_var.set("")
            pop_gi_mean_var.set("—")
            pop_gi_father_son_var.set("—")
            pop_gi_father_daughter_var.set("—")
            pop_gi_mother_son_var.set("—")
            pop_gi_mother_daughter_var.set("—")
            pop_family_count_var.set("—")
            pop_family_mean_size_var.set("—")

        if df_std is None or getattr(df_std, "empty", True):
            pop_total_var.set("-")
            pop_founders_count_var.set("-")
            pop_known_parents_var.set("-")
            pop_known_father_var.set("-")
            pop_known_mother_var.set("-")
            pop_lines_var.set("-")
            _clear_gen_demo_vars()
            return

        try:
            df_use = df_std.copy()
            df_use["id"] = df_use["id"].astype(str)
            df_use = df_use[df_use["id"] != TEST_ID]
        except Exception:
            df_use = df_std

        total = int(len(df_use))
        if total == 0:
            pop_total_var.set("0 osób (po odfiltrowaniu test ID)")
            pop_founders_count_var.set("-")
            pop_known_parents_var.set("-")
            pop_known_father_var.set("-")
            pop_known_mother_var.set("-")
            pop_lines_var.set("-")
            _clear_gen_demo_vars()
            return

        father_known = df_use["father_id"].notna()
        mother_known = df_use["mother_id"].notna()

        founders_mask = (~father_known) & (~mother_known)
        founders = int(founders_mask.sum())

        both_known = father_known & mother_known
        both_known_count = int(both_known.sum())
        at_least_one = int((father_known | mother_known).sum())

        def _pct(n: int) -> float:
            return round(100.0 * float(n) / float(total), 1)

        pop_total_var.set(f"- Liczba osobników: {total}")
        pop_founders_count_var.set(f"- Założyciele (brak ojca i matki): {founders} ({_pct(founders)}%)")
        pop_known_parents_var.set(
            f"- Znane oboje rodziców: {both_known_count} ({_pct(both_known_count)}%) • "
            f"- Znany przynajmniej jeden rodzic: {at_least_one} ({_pct(at_least_one)}%)"
        )
        pop_known_father_var.set(f"- Znany ojciec: {int(father_known.sum())} ({_pct(int(father_known.sum()))}%)")
        pop_known_mother_var.set(f"- Znana matka: {int(mother_known.sum())} ({_pct(int(mother_known.sum()))}%)")

        # Rozkład linii dla osobników (E w Excelu: line).
        if "line" in df_use.columns:
            line_vals = df_use["line"].astype(str).str.strip().str.upper().fillna("NA")
            lb = int((line_vals == "LB").sum())
            lc = int((line_vals == "LC").sum())
            c_other = int((~line_vals.isin(["LB", "LC"])).sum())
            pop_lines_var.set(f"- Linie: LB={lb} • LC={lc} • reszta={c_other}")
        else:
            pop_lines_var.set("- Linie: brak kolumny `line`")

        # --- Wkład założycieli / bottleneck (na podstawie p_i i founder-stop) ---
        _clear_gen_demo_vars()

        try:
            people_for_stats = state.get("people")
            if people_for_stats:
                depth_val: int = 4
                try:
                    depth_val = int(str(pop_depth_inb_var.get()).strip())
                except Exception:
                    depth_val = 4
                if depth_val < 0:
                    depth_val = 0
                depth_val = min(depth_val, 30)
                max_gen_back = None if bool(pop_unbounded_inb_var.get()) else depth_val

                id_list = [str(x) for x in df_use["id"].tolist()]
                try:
                    m_phi, m_r, mk_note = mean_kinship_pairwise(
                        people_for_stats,
                        id_list,
                        max_generations_back=max_gen_back,
                    )
                    if m_phi is not None and m_r is not None:
                        pop_mk_phi_var.set(f"{m_phi:.6f}")
                        pop_mk_r_var.set(f"{m_r:.6f}")
                        pop_mk_note_var.set(mk_note)
                except Exception:
                    pop_mk_phi_var.set("—")
                    pop_mk_r_var.set("—")
                    pop_mk_note_var.set("Nie udało się policzyć mean kinship.")

                stats_founders = compute_population_genetics_stats(
                    df_std=df_use,  # type: ignore[arg-type]
                    people=people_for_stats,  # type: ignore[arg-type]
                    max_generations_back=max_gen_back,
                    calc_f=False,
                    calc_completeness=False,
                    calc_founders=True,
                    calc_lines=False,
                )
                state["population_founder_contributions"] = stats_founders.founder_contributions
                f_e = float(stats_founders.founders.f_e)
                f_a = float(stats_founders.founders.f_a)
                pop_f_e_var.set(f"{f_e:.4f}")
                pop_f_a_var.set(f"{f_a:.4f}")
                if f_a > 0:
                    pop_bottleneck_var.set(f"{f_e / f_a:.3f}")
                # W aktualnej implementacji `f_a` jest policzone spójnie z founder-stop,
                # a brak osobnej wersji `f_ge` — przyjmujemy więc przybliżenie f_ge = f_e.
                pop_drift_var.set("1.000")
        except Exception:
            # Jeśli założyciele policzyć się nie uda, reszta metryk może działać.
            pass

        # --- GI (Generation Interval) oraz struktura rodzin ---
        people = state.get("people")
        state["population_gi_mean"] = None
        pop_gi_mean_var.set("—")
        pop_gi_father_son_var.set("—")
        pop_gi_father_daughter_var.set("—")
        pop_gi_mother_son_var.set("—")
        pop_gi_mother_daughter_var.set("—")
        pop_family_count_var.set("—")
        pop_family_mean_size_var.set("—")

        gi_data = compute_gi_and_family_data(df_use, people if people else {})
        gi_all = gi_data.get("gi_all")
        if gi_all is not None:
            state["population_gi_mean"] = gi_all
            pop_gi_mean_var.set(f"{float(gi_all):.2f} lat")
        if gi_data.get("gi_fs") is not None:
            pop_gi_father_son_var.set(f"{float(gi_data['gi_fs']):.2f} lat")
        if gi_data.get("gi_fd") is not None:
            pop_gi_father_daughter_var.set(f"{float(gi_data['gi_fd']):.2f} lat")
        if gi_data.get("gi_ms") is not None:
            pop_gi_mother_son_var.set(f"{float(gi_data['gi_ms']):.2f} lat")
        if gi_data.get("gi_md") is not None:
            pop_gi_mother_daughter_var.set(f"{float(gi_data['gi_md']):.2f} lat")

        fs_fam: list = gi_data.get("family_sizes") or []
        state["population_family_sizes"] = fs_fam
        if fs_fam:
            pop_family_count_var.set(str(len(fs_fam)))
            pop_family_mean_size_var.set(f"{float(sum(fs_fam)) / float(len(fs_fam)):.2f}")

        # Aktualizacja wykresów po wczytaniu danych.
        try:
            render_birth_decade_charts(
                df_use,
                state=state,
                colors=colors,
                save_figure_fn=_save_figure_as_jpeg,
                pop_birth_sex_plot_area=pop_birth_sex_plot_area,
                pop_birth_line_plot_area=pop_birth_line_plot_area,
                pop_birth_ratio_plot_area=pop_birth_ratio_plot_area,
                pop_comp_sex_plot_area=pop_comp_sex_plot_area,
                pop_comp_line_plot_area=pop_comp_line_plot_area,
            )
            render_inbreeding_year_trends(
                df_use,
                state=state,
                pop_depth_inb_var=pop_depth_inb_var,
                pop_unbounded_inb_var=pop_unbounded_inb_var,
                pop_inb_year_sex_plot_area=pop_inb_year_sex_plot_area,
                pop_inb_year_line_plot_area=pop_inb_year_line_plot_area,
                pop_ria_overall_var=pop_ria_overall_var,
                pop_ne_var=pop_ne_var,
                save_figure_fn=_save_figure_as_jpeg,
            )

            render_inbreeding_year_trends_smoothed(
                df_use,
                state=state,
                pop_depth_inb_var=pop_depth_inb_var,
                pop_unbounded_inb_var=pop_unbounded_inb_var,
                pop_inb_smooth_plot_area=pop_inb_smooth_plot_area,
                save_figure_fn=_save_figure_as_jpeg,
                smooth_window=int(APP_CFG.f_ria_smooth_window),
            )

            render_f_distribution_charts(
                df_use=df_use,
                state=state,
                colors=colors,
                save_figure_fn=_save_figure_as_jpeg,
                pop_f_hist_plot_area=pop_f_hist_plot_area,
                pop_f_scatter_plot_area=pop_f_scatter_plot_area,
            )

            render_founders_pi_chart(
                state=state,
                colors=colors,
                people=people,
                save_figure_fn=_save_figure_as_jpeg,
                pop_founders_plot_area=pop_founders_plot_area,
            )

            render_gi_and_family_charts(
                colors=colors,
                save_figure_fn=_save_figure_as_jpeg,
                people=people,
                gi_data=gi_data,
                pop_gi_plot_area=pop_gi_plot_area,
                pop_gi_trend_plot_area=pop_gi_trend_plot_area,
                pop_family_plot_area=pop_family_plot_area,
            )

            render_gi_four_paths_smoothed(
                gi_data=gi_data,
                colors=colors,
                save_figure_fn=_save_figure_as_jpeg,
                pop_gi_4paths_plot_area=pop_gi_4paths_plot_area,
                smooth_window=int(APP_CFG.gi_smooth_window),
            )
        except Exception:
            pass

    state["population_metrics"] = {}

    # -------------------------
    # Persons tab (unlimited preview)
    # -------------------------
    persons_controls = ttk.Frame(tab_persons)
    persons_controls.pack(side=tk.TOP, fill=tk.X)

    persons_left = ttk.Frame(persons_controls)
    persons_left.pack(side=tk.LEFT, fill=tk.X, expand=True)

    persons_search_var = tk.StringVar(value="")
    ttk.Label(persons_left, text="Szukaj ID:", foreground=colors.MUTED).pack(side=tk.LEFT, padx=(0, 6))
    persons_search_entry = ttk.Entry(persons_left, textvariable=persons_search_var, width=18, state="disabled")
    persons_search_entry.pack(side=tk.LEFT)
    find_person_btn = ttk.Button(persons_left, text="Znajdź", state="disabled")
    find_person_btn.pack(side=tk.LEFT, padx=(8, 0))

    # Dodatkowy filtr do rejestru i (opcjonalnie) pod „mating”.
    birth_location_filter_var = tk.StringVar(value="Bez filtra")
    ttk.Label(persons_left, text="Miejsce ur.:", foreground=colors.MUTED).pack(side=tk.LEFT, padx=(16, 6))
    persons_birth_location_filter_cb = ttk.Combobox(
        persons_left,
        textvariable=birth_location_filter_var,
        state="disabled",
        width=22,
        values=["Bez filtra"],
    )
    persons_birth_location_filter_cb.pack(side=tk.LEFT, padx=(0, 0))

    def _norm_birth_location_val(v: object) -> str:
        """Normalizacja pod filtr: puste/NaN → 'NA', reszta → string bez otaczających spacji."""
        if v is None:
            return "NA"
        try:
            if isinstance(v, float) and v != v:
                return "NA"
        except Exception:
            pass
        s = str(v).strip()
        if not s or s.lower() in {"nan", "none", "null"}:
            return "NA"
        return s

    ttk.Button(
        persons_controls,
        text="Pomoc: lista osobników",
        command=lambda: show_help_window(root, "Rejestr osobników", hc.SECTION_PERSONS),
    ).pack(side=tk.RIGHT)
    ttk.Label(tab_persons, text="Kliknij nagłówek kolumny, aby sortować.", foreground=colors.MUTED).pack(
        anchor="w", pady=(4, 0)
    )

    def _set_status(msg: str) -> None:
        status_var.set(msg)
        root.update_idletasks()

    # -------------------------
    # Loading tab: mapowanie kolumn
    # -------------------------
    LOAD_NONE = "<nie wybrano>"

    loading_frame = ttk.Frame(tab_loading)
    loading_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

    ttk.Label(
        loading_frame,
        text=(
            "1) Wybierz plik (CSV/XLSX).\n"
            "2) Jeśli automatyczne mapowanie po nazwach nie zadziała, przypisz kolumny do pól aplikacji.\n"
            "3) Kliknij „Zastosuj mapowanie i wczytaj”."
        ),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(6, 0))

    ttk.Button(
        loading_frame,
        text="Pomoc: import danych",
        command=lambda: show_help_window(
            root,
            "Import danych",
            hc.SECTION_LOADING + "\n\n" + hc.SECTION_VALIDATION,
        ),
    ).pack(anchor="w", pady=(4, 0))

    loading_top_btns = ttk.Frame(loading_frame)
    loading_top_btns.pack(anchor="w", pady=(14, 0), fill=tk.X)

    raw_file_var = tk.StringVar(value="Brak wczytanego pliku.")
    ttk.Label(loading_frame, textvariable=raw_file_var, foreground=colors.MUTED).pack(anchor="w", pady=(8, 0))

    mapping_note_var = tk.StringVar(value="Wczytaj plik aby rozpocząć mapowanie.")
    ttk.Label(loading_frame, textvariable=mapping_note_var, foreground=colors.TEXT, wraplength=900).pack(anchor="w", pady=(4, 10))

    validation_var = tk.StringVar(value="Walidacja bazy: -")

    ttk.Label(
        tab_validation,
        text=(
            "Podsumowanie walidacji wczytanej bazy. Po sprawdzeniu przejdź do rejestru lub grafu pedigree.\n"
            "Przepływ: import → walidacja → rejestr → analiza osobnika / par → populacja → raport."
        ),
        foreground=colors.MUTED,
        wraplength=920,
        justify="left",
    ).pack(anchor="w", pady=(0, 10))
    ttk.Button(
        tab_validation,
        text="Pomoc: walidacja bazy",
        command=lambda: show_help_window(root, "Walidacja bazy", hc.SECTION_VALIDATION),
    ).pack(anchor="w", pady=(0, 8))

    def on_export_validation_issues_csv() -> None:
        rep = state.get("validation_report")
        if rep is None:
            messagebox.showinfo("Info", "Brak raportu walidacji. Wczytaj najpierw bazę.")
            return
        filename = filedialog.asksaveasfilename(
            title="Eksport problemów walidacji (CSV)",
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv"), ("Wszystkie pliki", "*.*")],
        )
        if not filename:
            return
        try:
            with open(filename, "w", encoding="utf-8-sig", newline="") as f:
                f.write(rep.to_csv_string())
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się zapisać pliku: {e}")
            return
        n = len(rep.export_rows)
        _set_status(f"Zapisano CSV walidacji ({n} wierszy z problemami): {filename}")
        _open_exported_file(filename)

    validation_block = ttk.Frame(tab_validation)
    validation_block.pack(anchor="w", fill=tk.X, pady=(0, 10))
    ttk.Label(validation_block, textvariable=validation_var, foreground=colors.MUTED, wraplength=900, justify="left").pack(
        anchor="w"
    )
    validation_export_btn = ttk.Button(
        validation_block,
        text="Eksport problemów walidacji do CSV (id, typ, szczegóły)…",
        command=on_export_validation_issues_csv,
        state="disabled",
    )
    validation_export_btn.pack(anchor="w", pady=(6, 0))

    validation_miss_lf = ttk.LabelFrame(
        tab_validation,
        text="Braki danych — heatmapa (% wierszy z luką w kolumnie)",
    )
    validation_miss_lf.pack(anchor="w", fill=tk.BOTH, expand=True, pady=(12, 0))
    ttk.Label(
        validation_miss_lf,
        text=(
            "Dla każdej kolumny: udział wierszy z NaN, pustym polem lub zapisem „nan”. "
            "Mapa jak na pasie: w każdym polu nazwa kolumny i % braków; jasno = mało braków, ciemno (mech/kora) = więcej."
        ),
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", padx=8, pady=(8, 4))
    validation_miss_plot_host = ttk.Frame(validation_miss_lf)
    validation_miss_plot_host.pack(anchor="w", fill=tk.BOTH, expand=True, padx=8, pady=(0, 10))

    def _refresh_validation_missing_heatmap() -> None:
        _clear_frame(validation_miss_plot_host)
        df_v = state.get("df_std")
        if df_v is None or getattr(df_v, "empty", True):
            ttk.Label(
                validation_miss_plot_host,
                text="Wczytaj bazę (Import / domyślny plik), aby zobaczyć mapę braków.",
                foreground=colors.MUTED,
            ).pack(anchor="w", pady=8)
            return
        try:
            fig_m = fig_column_missing_heatmap(df_v)
        except Exception as ex:
            ttk.Label(
                validation_miss_plot_host,
                text=f"Nie udało się narysować heatmapy braków: {ex}",
                foreground=colors.MUTED,
                wraplength=880,
            ).pack(anchor="w", pady=8)
            return
        miss_btns = ttk.Frame(validation_miss_plot_host)
        miss_btns.pack(anchor="w", pady=(0, 6))
        ttk.Button(
            miss_btns,
            text="Zapis wykresu (JPEG)…",
            command=lambda f=fig_m: _save_figure_as_jpeg(f, default_basename="walidacja_braki_heatmap"),
        ).pack(side=tk.LEFT)
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

        canvas_m = FigureCanvasTkAgg(fig_m, master=validation_miss_plot_host)
        canvas_m.draw()
        canvas_m.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    internet_lf = ttk.LabelFrame(loading_frame, text="Pobieranie z internetu (EBPB)")
    internet_lf.pack(side=tk.TOP, fill=tk.BOTH, expand=False, pady=(0, 10), padx=0)

    ebpb_url_var = tk.StringVar(value="")
    ttk.Label(internet_lf, text="Wklej bezpośredni URL do pliku z 'pobierz zestawienie':").pack(anchor="w", padx=10, pady=(10, 0))
    ebpb_row = ttk.Frame(internet_lf)
    ebpb_row.pack(fill=tk.X, padx=10, pady=(8, 10))
    ebpb_entry = ttk.Entry(ebpb_row, textvariable=ebpb_url_var)
    ebpb_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
    ebpb_download_btn = ttk.Button(ebpb_row, text="Pobierz i wczytaj", command=lambda: None)
    ebpb_download_btn.pack(side=tk.LEFT, padx=(10, 0))

    map_lf = ttk.LabelFrame(loading_frame, text="Mapowanie kolumn")
    map_lf.pack(side=tk.TOP, fill=tk.BOTH, expand=False, pady=(0, 10))

    map_grid = ttk.Frame(map_lf)
    map_grid.pack(fill=tk.X, expand=True, padx=10, pady=10)

    # Wewnętrzne pola aplikacji (df_std)
    map_id_var = tk.StringVar(value=LOAD_NONE)
    map_sex_var = tk.StringVar(value=LOAD_NONE)
    map_line_var = tk.StringVar(value=LOAD_NONE)
    map_birth_year_var = tk.StringVar(value=LOAD_NONE)
    map_birth_location_var = tk.StringVar(value=LOAD_NONE)
    map_father_id_var = tk.StringVar(value=LOAD_NONE)
    map_mother_id_var = tk.StringVar(value=LOAD_NONE)
    map_name_var = tk.StringVar(value=LOAD_NONE)
    map_father_line_var = tk.StringVar(value=LOAD_NONE)
    map_mother_line_var = tk.StringVar(value=LOAD_NONE)

    def _clean_colname(s: object) -> str:
        if s is None:
            return ""
        return str(s).replace("\n", " ").strip()

    def _set_cb_options(cb: ttk.Combobox, cols: list[str]) -> None:
        options = [LOAD_NONE] + cols
        cb.configure(values=options, state="readonly")

    ttk.Label(map_grid, text="ID (id/Number)", foreground=colors.MUTED).grid(row=0, column=0, sticky="w", pady=4)
    cb_id = ttk.Combobox(map_grid, textvariable=map_id_var, width=45)
    cb_id.grid(row=0, column=1, sticky="ew", padx=(10, 0), pady=4)

    ttk.Label(map_grid, text="Płeć (sex: M/F)", foreground=colors.MUTED).grid(row=1, column=0, sticky="w", pady=4)
    cb_sex = ttk.Combobox(map_grid, textvariable=map_sex_var, width=45)
    cb_sex.grid(row=1, column=1, sticky="ew", padx=(10, 0), pady=4)

    ttk.Label(map_grid, text="Linia (line: LB/LC)", foreground=colors.MUTED).grid(row=2, column=0, sticky="w", pady=4)
    cb_line = ttk.Combobox(map_grid, textvariable=map_line_var, width=45)
    cb_line.grid(row=2, column=1, sticky="ew", padx=(10, 0), pady=4)

    ttk.Label(map_grid, text="Rok urodzenia (birth_year)", foreground=colors.MUTED).grid(row=3, column=0, sticky="w", pady=4)
    cb_birth_year = ttk.Combobox(map_grid, textvariable=map_birth_year_var, width=45)
    cb_birth_year.grid(row=3, column=1, sticky="ew", padx=(10, 0), pady=4)

    ttk.Label(map_grid, text="Miejsce urodzenia (birth_location) [opcjonalnie]", foreground=colors.MUTED).grid(
        row=4, column=0, sticky="w", pady=4
    )
    cb_birth_location = ttk.Combobox(map_grid, textvariable=map_birth_location_var, width=45)
    cb_birth_location.grid(row=4, column=1, sticky="ew", padx=(10, 0), pady=4)

    ttk.Label(map_grid, text="Ojciec (father_id)", foreground=colors.MUTED).grid(row=5, column=0, sticky="w", pady=4)
    cb_father_id = ttk.Combobox(map_grid, textvariable=map_father_id_var, width=45)
    cb_father_id.grid(row=5, column=1, sticky="ew", padx=(10, 0), pady=4)

    ttk.Label(map_grid, text="Matka (mother_id)", foreground=colors.MUTED).grid(row=6, column=0, sticky="w", pady=4)
    cb_mother_id = ttk.Combobox(map_grid, textvariable=map_mother_id_var, width=45)
    cb_mother_id.grid(row=6, column=1, sticky="ew", padx=(10, 0), pady=4)

    ttk.Label(map_grid, text="Imię (name) [opcjonalnie]", foreground=colors.MUTED).grid(row=7, column=0, sticky="w", pady=4)
    cb_name = ttk.Combobox(map_grid, textvariable=map_name_var, width=45)
    cb_name.grid(row=7, column=1, sticky="ew", padx=(10, 0), pady=4)

    ttk.Label(map_grid, text="Linia ojca (father_line) [opcjonalnie]", foreground=colors.MUTED).grid(
        row=8, column=0, sticky="w", pady=4
    )
    cb_father_line = ttk.Combobox(map_grid, textvariable=map_father_line_var, width=45)
    cb_father_line.grid(row=8, column=1, sticky="ew", padx=(10, 0), pady=4)

    ttk.Label(map_grid, text="Linia matki (mother_line) [opcjonalnie]", foreground=colors.MUTED).grid(
        row=9, column=0, sticky="w", pady=4
    )
    cb_mother_line = ttk.Combobox(map_grid, textvariable=map_mother_line_var, width=45)
    cb_mother_line.grid(row=9, column=1, sticky="ew", padx=(10, 0), pady=4)

    # kolumny muszą się zmieścić
    map_grid.columnconfigure(1, weight=1)

    def _clear_mapping_vars() -> None:
        for v in [
            map_id_var,
            map_sex_var,
            map_line_var,
            map_birth_year_var,
            map_birth_location_var,
            map_father_id_var,
            map_mother_id_var,
            map_name_var,
            map_father_line_var,
            map_mother_line_var,
        ]:
            v.set(LOAD_NONE)

    def _autoselect_mapping(cols_clean: list[str]) -> None:
        # Tylko jeśli exact match po czyszczeniu.
        colset = {c.upper(): c for c in cols_clean}

        def pick(candidates: list[str]) -> Optional[str]:
            for cand in candidates:
                key = str(cand).upper()
                if key in colset:
                    return colset[key]
            return None

        map_id_var.set(pick(["Number", "ID", "id"]) or LOAD_NONE)
        map_sex_var.set(pick(["Sex", "sex", "Gender", "gender"]) or LOAD_NONE)
        map_line_var.set(pick(["Line", "line"]) or LOAD_NONE)
        map_birth_year_var.set(pick(["Birth year", "birth_year", "BirthYear"]) or LOAD_NONE)
        map_birth_location_var.set(
            pick(["Birth location", "birth_location", "Birth place", "birth_place"]) or LOAD_NONE
        )
        map_father_id_var.set(pick(["Father", "father_id", "Father ID", "father"]) or LOAD_NONE)
        map_mother_id_var.set(pick(["Mother", "mother_id", "Mother ID", "mother"]) or LOAD_NONE)
        map_name_var.set(pick(["Name", "name", "Alt name", "alt_name"]) or LOAD_NONE)
        # Domyślne "niestandardowe" nazwy (Excel bywa puste/Unnamed:..)
        map_father_line_var.set(pick(["Unnamed: 9", "father_line"]) or LOAD_NONE)
        map_mother_line_var.set(pick(["Unnamed: 12", "mother_line"]) or LOAD_NONE)

    def _refresh_mapping_form_for_raw_df(df_raw: object) -> None:
        try:
            cols = list(getattr(df_raw, "columns"))
        except Exception:
            cols = []
        cols_clean = [_clean_colname(c) for c in cols]
        _set_cb_options(cb_id, cols_clean)
        _set_cb_options(cb_sex, cols_clean)
        _set_cb_options(cb_line, cols_clean)
        _set_cb_options(cb_birth_year, cols_clean)
        _set_cb_options(cb_birth_location, cols_clean)
        _set_cb_options(cb_father_id, cols_clean)
        _set_cb_options(cb_mother_id, cols_clean)
        _set_cb_options(cb_name, cols_clean)
        _set_cb_options(cb_father_line, cols_clean)
        _set_cb_options(cb_mother_line, cols_clean)
        _autoselect_mapping(cols_clean)

    def _set_mapping_ready(enabled: bool) -> None:
        apply_mapping_btn.configure(state="normal" if enabled else "disabled")

    def on_apply_mapping() -> None:
        df_raw = state.get("df_raw")
        if df_raw is None:
            messagebox.showinfo("Info", "Najpierw wczytaj plik.")
            return

        def _sel(v: str) -> Optional[str]:
            if not v or v == LOAD_NONE:
                return None
            return v

        column_mapping = {
            "id": _sel(map_id_var.get()),
            "sex": _sel(map_sex_var.get()),
            "line": _sel(map_line_var.get()),
            "birth_year": _sel(map_birth_year_var.get()),
            "birth_location": _sel(map_birth_location_var.get()),
            "father_id": _sel(map_father_id_var.get()),
            "mother_id": _sel(map_mother_id_var.get()),
            "name": _sel(map_name_var.get()),
            "father_line": _sel(map_father_line_var.get()),
            "mother_line": _sel(map_mother_line_var.get()),
        }

        try:
            df_std = standardize_bison_report_dataframe_with_column_mapping(df_raw, column_mapping)
            people = _build_people_map(df_std)
            _apply_dataset(
                df_std=df_std,
                people=people,
                source=f"Wczytano (mapowanie): {state.get('raw_filename')}",
            )
            notebook.select(tab_validation)
        except Exception as e:
            messagebox.showerror("Błąd mapowania", str(e))
            _set_status(f"Błąd mapowania: {e}")

    def on_load_default() -> None:
        try:
            df_std, _info = load_default_bison_report()
            people = _build_people_map(df_std)
            _apply_dataset(df_std=df_std, people=people, source="Wczytano domyślną bazę")
            notebook.select(tab_validation)
            raw_file_var.set("Wczytano domyślną bazę (bez mapowania).")
            mapping_note_var.set("Domyślna baza została wczytana automatycznie.")
            _set_mapping_ready(False)
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się wczytać domyślnej bazy: {e}")
            _set_status(f"Błąd: {e}")

    def on_download_from_ebpb() -> None:
        url = str(ebpb_url_var.get()).strip()
        if not url:
            messagebox.showinfo("Info", "Wklej URL do pliku z EBPB.")
            return

        # Krok 1: spróbuj wczytać od razu jako standardowy eksport aplikacji.
        try:
            df_std, _info = load_dataset_from_url(url)
            people = _build_people_map(df_std)
            _apply_dataset(df_std=df_std, people=people, source=f"Pobrano z internetu: {url}")
            notebook.select(tab_validation)
            state["df_raw"] = None
            state["raw_filename"] = "ebpb_download"
            raw_file_var.set("Pobrano z EBPB i wczytano automatycznie.")
            mapping_note_var.set("Mapowanie kolumn nie było potrzebne.")
            _set_mapping_ready(False)
            return
        except Exception:
            # OK: przechodzimy na tryb 'mapowanie', jeśli eksport ma inne nagłówki.
            pass

        # Krok 2: wczytaj surową tabelę i przełącz UI na mapowanie.
        try:
            df_raw = load_raw_dataframe_from_url(url)
            state["df_raw"] = df_raw
            state["raw_filename"] = url

            set_controls_enabled(False)
            raw_file_var.set("Pobrano z EBPB: mapowanie wymagane.")
            mapping_note_var.set("Automatyczne mapowanie po nazwach nie zadziałało — przypisz kolumny w formularzu.")

            try:
                for item in tree.get_children():
                    tree.delete(item)
            except Exception:
                pass

            dataset_range_var.set("Zakres ID: brak")
            _clear_mapping_vars()
            _refresh_mapping_form_for_raw_df(df_raw)
            _set_mapping_ready(True)
            notebook.select(tab_loading)
            root.update_idletasks()
        except Exception as e:
            messagebox.showerror("Błąd pobierania", str(e))
            _set_status(f"Błąd pobierania: {e}")

    ebpb_download_btn.configure(command=on_download_from_ebpb)

    def on_choose_file() -> None:
        path = filedialog.askopenfilename(
            title="Wybierz plik",
            filetypes=[("CSV", "*.csv"), ("Excel", "*.xlsx *.xls"), ("All", "*.*")],
        )
        if not path:
            return

        # Krok 1: spróbujmy automatycznego mapowania po nazwach.
        try:
            df_std, _info = load_dataset_from_path(path)
            people = _build_people_map(df_std)
            _apply_dataset(df_std=df_std, people=people, source=f"Wczytano: {Path(path).name}")
            notebook.select(tab_validation)
            state["df_raw"] = None
            state["raw_filename"] = Path(path).name
            raw_file_var.set(f"Wczytano automatycznie: {Path(path).name}")
            mapping_note_var.set("Mapowanie kolumn nie jest potrzebne (dopasowane do formatu aplikacji).")
            _set_mapping_ready(False)
            return
        except Exception:
            # OK: przejdziemy do mapowania ręcznego.
            pass

        # Krok 2: wczytujemy surowy plik i prosimy o mapowanie kolumn.
        try:
            import pandas as pd

            ext = Path(path).suffix.lower()
            if ext == ".csv":
                df_raw = pd.read_csv(path)
            elif ext in {".xlsx", ".xls"}:
                df_raw = pd.read_excel(path, sheet_name=0)
            else:
                raise ValueError(f"Nieobsługiwany typ pliku: {ext}")

            state["df_raw"] = df_raw
            state["raw_filename"] = Path(path).name

            # Przestaw UI w tryb "oczekuje na mapowanie".
            set_controls_enabled(False)
            raw_file_var.set(f"Wczytano plik: {Path(path).name} (mapowanie wymagane)")
            mapping_note_var.set("Wybierz kolumny dla pól aplikacji, a potem kliknij „Zastosuj mapowanie i wczytaj”.")

            # Wyczyść podgląd osób (żeby nie mylić danych).
            try:
                for item in tree.get_children():
                    tree.delete(item)
            except Exception:
                pass

            dataset_range_var.set("Zakres ID: brak")

            _clear_mapping_vars()
            _refresh_mapping_form_for_raw_df(df_raw)
            _set_mapping_ready(True)
            root.update_idletasks()
        except Exception as e:
            messagebox.showerror("Błąd wczytywania", str(e))
            _set_status(f"Błąd wczytywania: {e}")

    # Przycisk wymaga zmapowania
    apply_mapping_btn = ttk.Button(loading_frame, text="Zastosuj mapowanie i wczytaj", command=on_apply_mapping, state="disabled")
    apply_mapping_btn.pack(anchor="w", pady=(10, 0))

    loading_btns_left = ttk.Frame(loading_top_btns)
    loading_btns_left.pack(side=tk.LEFT)
    ttk.Button(loading_top_btns, text="Wybierz plik bazy", command=on_choose_file).pack(side=tk.LEFT, padx=(0, 10))
    ttk.Button(loading_top_btns, text="Wczytaj domyślną bazę", command=on_load_default).pack(side=tk.LEFT)

    dataset_range_var = tk.StringVar(value="")
    ttk.Label(persons_left, textvariable=dataset_range_var, foreground=colors.MUTED).pack(side=tk.LEFT, padx=(20, 0))

    persons_tree_frame = ttk.Frame(tab_persons)
    persons_tree_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True, pady=(14, 0))

    tree_columns = (
        "id",
        "name",
        "sex",
        "birth_year",
        "birth_location",
        "father_id",
        "mother_id",
        "SireFounder",
        "SireSteps",
        "DamFounder",
        "DamSteps",
        "line",
        "father_line",
        "mother_line",
    )
    tree = ttk.Treeview(persons_tree_frame, columns=tree_columns, show="headings", height=18)
    tree_sort_state: dict[str, bool] = {"id": True}  # True = A->Z (ascending)

    def _update_heading(col: str) -> None:
        asc = tree_sort_state.get(col, True)
        arrow = "▲" if asc else "▼"
        if col == "id":
            base = "id (A->Z)"
        elif col == "birth_location":
            base = "Miejsce ur. (birth_location)"
        else:
            base = col
        # command utrzymujemy przez ponowne ustawienie, żeby nie gubić sortowania.
        tree.heading(col, text=f"{base} {arrow}", command=lambda c=col: on_sort_click(c))

    def _sort_value_key(col: str, val: object) -> tuple[int, object]:
        # key: (is_missing, coerced_value)
        if val is None:
            return (1, "")
        s = str(val).strip()
        if s == "":
            return (1, "")

        if col == "id":
            # _id_sort_key jest zdefiniowane niżej w kodzie; w praktyce callback wywołuje się dopiero później.
            try:
                return (0, _id_sort_key(s))
            except Exception:
                return (0, s)

        # Spróbuj traktować jako liczbę dla pól typu steps / birth_year.
        if col in {"birth_year", "SireSteps", "DamSteps"}:
            try:
                return (0, float(s))
            except Exception:
                return (0, s.lower())

        # Domyślnie: leksykograficznie (A->Z).
        return (0, s.lower())

    def _sort_tree_by(col: str) -> None:
        # Toggle kierunku.
        prev_asc = tree_sort_state.get(col, True)
        next_asc = not prev_asc
        tree_sort_state[col] = next_asc

        items = list(tree.get_children())
        items_with_keys = []
        for item in items:
            val = tree.set(item, col)
            key = _sort_value_key(col, val)
            items_with_keys.append((key, item))

        # Ascending sort uses reverse=False, descending uses reverse=True.
        reverse = not next_asc
        items_with_keys.sort(key=lambda t: t[0], reverse=reverse)

        for _, item in items_with_keys:
            tree.move(item, "", "end")

        _update_heading(col)

    def on_sort_click(col: str) -> None:
        if not tree.get_children():
            return
        _sort_tree_by(col)

    for col in tree_columns:
        _update_heading(col)
        tree.column(col, width=120, anchor="w")
    vsb = ttk.Scrollbar(persons_tree_frame, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=vsb.set)
    tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    vsb.pack(side=tk.RIGHT, fill=tk.Y)

    persons_detail_lf = ttk.LabelFrame(
        tab_persons,
        text="Szczegóły osobnika — udział założycieli w genach (founder-stop, jak przy F)",
    )
    persons_detail_lf.pack(fill=tk.BOTH, expand=False, pady=(10, 0))
    persons_detail_text = tk.Text(
        persons_detail_lf,
        height=11,
        wrap="word",
        state="disabled",
        bg=colors.ENTRY_BG,
        fg=colors.TEXT,
        insertbackground=colors.TEXT,
    )
    persons_detail_scroll = ttk.Scrollbar(persons_detail_lf, orient="vertical", command=persons_detail_text.yview)
    persons_detail_text.configure(yscrollcommand=persons_detail_scroll.set)
    persons_detail_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(8, 0), pady=8)
    persons_detail_scroll.pack(side=tk.RIGHT, fill=tk.Y, pady=8, padx=(0, 8))

    PERSONS_FOUNDER_TOP_N = 30

    def _set_persons_detail_text(content: str) -> None:
        persons_detail_text.configure(state="normal")
        persons_detail_text.delete("1.0", tk.END)
        persons_detail_text.insert("1.0", content)
        persons_detail_text.configure(state="disabled")

    def _fill_person_detail(person_id: str) -> None:
        pid = str(person_id).strip()
        people = state.get("people") or {}
        if not pid:
            _set_persons_detail_text("Zaznacz wiersz w tabeli lub użyj pola Szukaj ID.")
            return
        if pid not in people:
            _set_persons_detail_text(f"Brak osobnika o ID „{pid}” w wczytanej bazie.")
            return
        p = people[pid]
        lines: list[str] = [
            f"ID: {pid}",
            f"Imię: {p.name or '—'}",
            f"Płeć: {getattr(p, 'sex', None) or '—'}  •  Linia (kolumna): {getattr(p, 'line', None) or '—'}",
            f"Rok ur.: {getattr(p, 'birth_year', None) or '—'}",
            f"Miejsce ur.: {getattr(p, 'birth_location', None) or '—'}",
            f"Ojciec: {p.father_id or '—'}  •  Matka: {p.mother_id or '—'}",
            "",
            "Udział genów od założycieli (suma 100 % przy pełnym drzewie w modelu founder-stop):",
            "",
        ]
        comp = FounderContributionComputer(people)
        contribs = comp.contributions_for(pid)
        if not contribs:
            lines.append("— brak rozkładu —")
        else:
            items = sorted(contribs.items(), key=lambda kv: kv[1], reverse=True)
            for fid, share in items[:PERSONS_FOUNDER_TOP_N]:
                fp = people.get(fid)
                fn = (fp.name if fp and fp.name else None) or "—"
                lines.append(f"  {fid}  ({fn}):  {share * 100.0:.4f} %")
            if len(items) > PERSONS_FOUNDER_TOP_N:
                lines.append(f"  … (+{len(items) - PERSONS_FOUNDER_TOP_N} kolejnych założycieli z mniejszym udziałem)")
        lm = (state.get("line_memberships") or {}).get(pid)
        if lm is not None:
            lines.extend(
                [
                    "",
                    "Linie hodowlane (sire / dam):",
                    f"  Sire: założyciel {lm.sire_founder_id or 'NA'}, kroki {lm.sire_steps}",
                    f"  Dam:  założyciel {lm.dam_founder_id or 'NA'}, kroki {lm.dam_steps}",
                ]
            )
        _set_persons_detail_text("\n".join(lines))

    def on_person_tree_select(_event: tk.Event | None = None) -> None:
        sel = tree.selection()
        if not sel:
            return
        vals = tree.item(sel[0], "values")
        if not vals:
            return
        _fill_person_detail(str(vals[0]))

    def on_find_person_in_tree() -> None:
        q = str(persons_search_var.get()).strip()
        if not q:
            messagebox.showinfo("Szukaj ID", "Wpisz pełne ID lub fragment (np. początek numeru).")
            return
        items = tree.get_children()
        if not items:
            messagebox.showinfo("Szukaj ID", "Najpierw wczytaj bazę.")
            return
        exact: str | None = None
        starts: list[str] = []
        contains: list[str] = []
        for it in items:
            iid = str(tree.set(it, "id"))
            if iid == q:
                exact = it
                break
            if iid.startswith(q):
                starts.append(it)
            elif q.lower() in iid.lower():
                contains.append(it)
        chosen = exact or (starts[0] if starts else None) or (contains[0] if contains else None)
        if not chosen:
            messagebox.showinfo("Szukaj ID", f"Brak wiersza pasującego do „{q}”.")
            return
        tree.selection_set(chosen)
        tree.focus(chosen)
        tree.see(chosen)
        _fill_person_detail(str(tree.set(chosen, "id")))

    find_person_btn.configure(command=on_find_person_in_tree)
    persons_search_entry.bind("<Return>", lambda _e: on_find_person_in_tree())
    tree.bind("<<TreeviewSelect>>", on_person_tree_select, add="+")

    # -------------------------
    # Rodowód tab controls
    # -------------------------
    rod_header = ttk.Frame(tab_pedigree)
    rod_header.pack(side=tk.TOP, fill=tk.X)
    ttk.Label(rod_header, text="ID (Number):").grid(row=0, column=0, sticky="w")
    id_anc_var = tk.StringVar(value="")
    id_anc_entry = ttk.Entry(rod_header, textvariable=id_anc_var, width=24, state="disabled")
    id_anc_entry.grid(row=0, column=1, sticky="w", padx=(8, 16))

    ttk.Label(rod_header, text="Max pokoleń:").grid(row=0, column=2, sticky="w")
    depth_anc_strvar = tk.StringVar(value="4")

    depth_anc_controls = ttk.Frame(rod_header)
    depth_anc_controls.grid(row=0, column=3, sticky="w", padx=(8, 16))
    # Spinbox zamiast suwaka: na macOS suwaki Tk/ttk bywają martwe; Spinbox jest stabilny.
    depth_anc_spin = ttk.Spinbox(
        depth_anc_controls,
        from_=0,
        to=30,
        increment=1,
        textvariable=depth_anc_strvar,
        width=6,
        wrap=False,
    )
    depth_anc_spin.pack(side=tk.LEFT)

    readable_anc_var = tk.BooleanVar(value=True)
    readable_anc_cb = ttk.Checkbutton(
        rod_header,
        text="Tryb czytelny (mniej etykiet)",
        variable=readable_anc_var,
        state="disabled",
    )
    readable_anc_cb.grid(row=1, column=0, columnspan=2, sticky="w", pady=(6, 0))

    # Filtr linii przodków (sire/dam) + zakres (bez limitu pokoleń).
    anc_line_mode_var = tk.StringVar(value="both")
    anc_line_radio_sire = ttk.Radiobutton(
        rod_header, text="Tylko sire (ojciec)", value="sire", variable=anc_line_mode_var, state="disabled"
    )
    anc_line_radio_dam = ttk.Radiobutton(
        rod_header, text="Tylko dam (matka)", value="dam", variable=anc_line_mode_var, state="disabled"
    )
    anc_line_radio_both = ttk.Radiobutton(
        rod_header, text="Sire+Dam (obie linie)", value="both", variable=anc_line_mode_var, state="disabled"
    )
    anc_line_radio_both.grid(row=2, column=0, sticky="w")
    anc_line_radio_sire.grid(row=2, column=1, sticky="w")
    anc_line_radio_dam.grid(row=2, column=2, sticky="w")

    # Domyślnie: limit pokoleń aktywny (pole max pokoleń ma być dostępne).
    unbounded_anc_var = tk.BooleanVar(value=False)
    unbounded_anc_cb = ttk.Checkbutton(
        rod_header,
        text="Bez limitu (do founderów)",
        variable=unbounded_anc_var,
        state="disabled",
    )
    unbounded_anc_cb.grid(row=1, column=2, sticky="w", padx=(8, 0), pady=(6, 0))

    def _sync_anc_depth_state() -> None:
        # Gdy bez limitu: blokujemy max pokoleń; w przeciwnym razie zależnie od tego, czy UI jest odblokowane.
        anc_enabled = bool(controls_enabled_ref[0])
        st = "normal" if anc_enabled else "disabled"
        depth_state = "disabled" if bool(unbounded_anc_var.get()) else st
        depth_anc_spin.configure(state=depth_state)

    _sync_anc_depth_state()
    unbounded_anc_var.trace_add("write", lambda *_args: _sync_anc_depth_state())

    anc_btn = ttk.Button(rod_header, text="Generuj przodków", state="disabled")
    anc_btn.grid(row=0, column=4, sticky="w", padx=(14, 0))
    ttk.Button(rod_header, text="Pomoc", command=lambda: show_help_window(root, "Graf pedigree", hc.SECTION_PEDIGREE)).grid(
        row=0, column=5, sticky="w", padx=(8, 0)
    )

    rod_canvas_container = ttk.Frame(tab_pedigree)
    rod_canvas_container.pack(fill=tk.BOTH, expand=True, pady=(14, 0))
    rod_title_var = tk.StringVar(value="Wykres przodków")
    ttk.Label(rod_canvas_container, textvariable=rod_title_var, foreground=colors.MUTED).pack(anchor="w")
    ttk.Label(rod_canvas_container, text="Linie (sire/dam):", foreground=colors.MUTED, font=tk_font(11, bold=True)).pack(
        anchor="w", pady=(4, 0)
    )

    rod_line_subj_var = tk.StringVar(value="")
    rod_line_father_var = tk.StringVar(value="")
    rod_line_mother_var = tk.StringVar(value="")

    rod_line_cols = ttk.Frame(rod_canvas_container)
    rod_line_cols.pack(fill=tk.X, expand=False, pady=(0, 0))

    rod_line_cols.columnconfigure(0, weight=1)
    rod_line_cols.columnconfigure(1, weight=1)
    rod_line_cols.columnconfigure(2, weight=1)

    ttk.Label(
        rod_line_cols,
        textvariable=rod_line_subj_var,
        foreground=colors.MUTED,
        justify="left",
        wraplength=360,
    ).grid(row=0, column=0, sticky="nsew", padx=(0, 10))
    ttk.Label(
        rod_line_cols,
        textvariable=rod_line_father_var,
        foreground=colors.MUTED,
        justify="left",
        wraplength=360,
    ).grid(row=0, column=1, sticky="nsew", padx=(0, 10))
    ttk.Label(
        rod_line_cols,
        textvariable=rod_line_mother_var,
        foreground=colors.MUTED,
        justify="left",
        wraplength=360,
    ).grid(row=0, column=2, sticky="nsew")
    rod_plot_frame = ttk.Frame(rod_canvas_container)
    rod_plot_frame.pack(fill=tk.BOTH, expand=True, pady=(8, 0))

    # -------------------------
    # Zakładka Analityka hodowlana — panel Inbred (analyses_nb utworzony wyżej).
    # -------------------------
    # Ogólne parametry populacyjne są w zakładce Metryki populacji (nie tutaj).

    ana_header = ttk.Frame(tab_inb)
    ana_header.pack(side=tk.TOP, fill=tk.X)

    ttk.Label(ana_header, text="ID (Number):").grid(row=0, column=0, sticky="w")
    id_inb_var = tk.StringVar(value="")
    id_inb_entry = ttk.Entry(ana_header, textvariable=id_inb_var, width=24, state="disabled")
    id_inb_entry.grid(row=0, column=1, sticky="w", padx=(8, 16))

    ttk.Label(ana_header, text="Max pokoleń:").grid(row=0, column=2, sticky="w")
    depth_inb_var = tk.StringVar(value="4")
    depth_inb_entry = ttk.Entry(ana_header, textvariable=depth_inb_var, width=10, state="disabled")
    depth_inb_entry.grid(row=0, column=3, sticky="w", padx=(8, 16))

    unbounded_inb_var = tk.BooleanVar(value=True)
    unbounded_inb_cb = ttk.Checkbutton(
        ana_header,
        text="Bez ograniczenia (do founderów)",
        variable=unbounded_inb_var,
        state="disabled",
    )
    unbounded_inb_cb.grid(row=1, column=0, columnspan=4, sticky="w", pady=(6, 0))

    inb_btn = ttk.Button(ana_header, text="Policz F (Wright)", state="disabled")
    inb_btn.grid(row=0, column=4, sticky="w", padx=(14, 0))
    ttk.Button(
        ana_header, text="Pomoc", command=lambda: show_help_window(root, "Inbred — współczynnik F", hc.SECTION_INBRED)
    ).grid(
        row=0, column=5, sticky="w", padx=(8, 0)
    )

    inb_result_var = tk.StringVar(value="F = -")
    ttk.Label(tab_inb, textvariable=inb_result_var, font=tk_font(16, bold=True)).pack(anchor="w", pady=(14, 0))
    inb_note_var = tk.StringVar(value="")
    # Zamiast jednej długiej etykiety używamy siatki (3-4 kolumny),
    # żeby treść była czytelniejsza i nie zajmowała całej szerokości.
    inb_note_frame = ttk.Frame(tab_inb)
    inb_note_frame.pack(side=tk.TOP, fill=tk.X, expand=False, pady=(6, 0))
    inb_plot_frame = ttk.Frame(tab_inb)
    inb_plot_frame.pack(fill=tk.BOTH, expand=True, pady=(12, 0))

    # -------------------------
    # Optymalizacja kojarzeń (ranking par)
    # -------------------------
    ttk.Button(
        tab_mating,
        text="Pomoc: optymalizacja kojarzeń",
        command=lambda: show_help_window(root, "Optymalizacja kojarzeń", hc.SECTION_MATING),
    ).pack(anchor="w", pady=(0, 4))

    mating_current_year_note_var = tk.StringVar(value="")
    mating_note_lbl = ttk.Label(
        tab_mating,
        textvariable=mating_current_year_note_var,
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    )
    mating_note_lbl.pack(anchor="w", pady=(2, 8))

    mating_age_limit_years = int(APP_CFG.mating_age_limit_years)
    mating_ranking_top_n = int(APP_CFG.mating_ranking_top_n)
    mating_ranking_max_uses_per_id = 3  # sire i dam liczone osobno w liście wynikowej

    mating_unbounded_var = tk.BooleanVar(value=False)
    mating_unbounded_cb = ttk.Checkbutton(
        tab_mating,
        text="Bez ograniczenia (do founderów) - może być wolne",
        variable=mating_unbounded_var,
        state="disabled",
    )
    mating_unbounded_cb.pack(anchor="w", pady=(0, 4))

    mating_depth_row = ttk.Frame(tab_mating)
    mating_depth_row.pack(anchor="w", pady=(0, 8))
    ttk.Label(mating_depth_row, text="Max pokoleń (gdy limit):").pack(side=tk.LEFT)
    mating_depth_var = tk.StringVar(value=str(settings_inb_depth_var.get()).strip() or "4")
    mating_depth_entry = ttk.Entry(mating_depth_row, textvariable=mating_depth_var, width=8, state="disabled")
    mating_depth_entry.pack(side=tk.LEFT, padx=(8, 0))

    # Ograniczenie puli kandydatów, aby uniknąć eksplozji par.
    mating_limits_frame = ttk.LabelFrame(tab_mating, text="Limit kandydatów (dla wydajności)")
    mating_limits_frame.pack(anchor="w", pady=(0, 8), fill=tk.X)
    ttk.Label(mating_limits_frame, text="Samce (M) max:").grid(row=0, column=0, sticky="w", padx=(10, 8), pady=(10, 6))
    mating_male_limit_var = tk.StringVar(value=str(APP_CFG.mating_default_male_limit))
    mating_male_limit_entry = ttk.Entry(mating_limits_frame, textvariable=mating_male_limit_var, width=8, state="disabled")
    mating_male_limit_entry.grid(row=0, column=1, sticky="w", pady=(10, 6))
    ttk.Label(mating_limits_frame, text="Samice (F) max:").grid(row=1, column=0, sticky="w", padx=(10, 8), pady=(0, 10))
    mating_female_limit_var = tk.StringVar(value=str(APP_CFG.mating_default_female_limit))
    mating_female_limit_entry = ttk.Entry(mating_limits_frame, textvariable=mating_female_limit_var, width=8, state="disabled")
    mating_female_limit_entry.grid(row=1, column=1, sticky="w", pady=(0, 10))

    mating_btn_row = ttk.Frame(tab_mating)
    mating_btn_row.pack(anchor="w", pady=(0, 8))
    mating_calc_btn = ttk.Button(
        mating_btn_row,
        text=f"Oblicz ranking kojarzeń (TOP {mating_ranking_top_n})",
        state="disabled",
    )
    mating_calc_btn.pack(side=tk.LEFT)
    mating_export_phi_btn = ttk.Button(
        mating_btn_row,
        text="Eksport macierzy Φ do CSV (ostatnie liczenie)",
        state="disabled",
    )
    mating_export_phi_btn.pack(side=tk.LEFT, padx=(10, 0))

    mating_pair_stats_var = tk.StringVar(value="")
    ttk.Label(tab_mating, textvariable=mating_pair_stats_var, foreground=colors.MUTED, wraplength=900).pack(
        anchor="w", pady=(6, 0)
    )

    mating_output = tk.Text(tab_mating, height=16, wrap="word")
    mating_output.pack(fill=tk.BOTH, expand=True, pady=(8, 0))
    mating_output.configure(state="disabled", bg=colors.ENTRY_BG, fg=colors.TEXT, insertbackground=colors.TEXT)

    ttk.Label(
        tab_kpair,
        text="Ustawienia max pokoleń / bez limitu są wspólne z zakładką „Ranking kojarzeń”.",
        foreground=colors.MUTED,
        wraplength=900,
        justify="left",
    ).pack(anchor="w", pady=(0, 6))

    kinship_lf = ttk.LabelFrame(
        tab_kpair,
        text="Kinship — para rodziców (ojciec × matka): Φ, R, F potomka, wyjaśnienie",
    )
    kinship_lf.pack(anchor="w", fill=tk.X, pady=(10, 0))
    ttk.Label(
        kinship_lf,
        text=(
            "Φ (Malecot / Wright) między hipotetycznym ojcem a matką jest symetryczne; R = 2Φ. "
            "F potomka = Φ(ojciec, matka). Poniżej: dlaczego taki wynik — wspólni przodkowie, wkład do Φ, pary ścieżek."
        ),
        foreground=colors.MUTED,
        wraplength=880,
    ).grid(row=0, column=0, columnspan=4, sticky="w", padx=10, pady=(8, 4))
    ttk.Label(
        kinship_lf,
        text="Listy zawierają tylko ID o płci M / F (z kolumny sex). Możesz też wpisać numer ID ręcznie w polu.",
        foreground=colors.MUTED,
        font=tk_font(9),
        wraplength=880,
    ).grid(row=1, column=0, columnspan=4, sticky="w", padx=10, pady=(0, 4))
    mating_kin_sire_var = tk.StringVar(value="")
    mating_kin_dam_var = tk.StringVar(value="")
    ttk.Label(kinship_lf, text="Ojciec (M):").grid(row=2, column=0, sticky="w", padx=(10, 6), pady=(4, 8))
    mating_kin_sire_cb = ttk.Combobox(
        kinship_lf, width=28, state="normal", textvariable=mating_kin_sire_var
    )
    mating_kin_sire_cb.grid(row=2, column=1, sticky="w", pady=(4, 8))
    ttk.Label(kinship_lf, text="Matka (F):").grid(row=2, column=2, sticky="w", padx=(12, 6), pady=(4, 8))
    mating_kin_dam_cb = ttk.Combobox(
        kinship_lf, width=28, state="normal", textvariable=mating_kin_dam_var
    )
    mating_kin_dam_cb.grid(row=2, column=3, sticky="w", padx=(0, 10), pady=(4, 8))
    mating_kinship_btn = ttk.Button(kinship_lf, text="Oblicz Φ, R i wyjaśnienie", state="disabled")
    mating_kinship_btn.grid(row=3, column=0, sticky="w", padx=(10, 0), pady=(0, 10))
    mating_kin_result_var = tk.StringVar(value="")
    ttk.Label(kinship_lf, textvariable=mating_kin_result_var, foreground=colors.TEXT, wraplength=860).grid(
        row=3, column=1, columnspan=3, sticky="w", pady=(0, 10)
    )

    kpair_explain_frame = ttk.LabelFrame(tab_kpair, text="Dlaczego ta para ma taki wynik? (wspólni przodkowie, ścieżki)")
    kpair_explain_frame.pack(fill=tk.BOTH, expand=True, pady=(10, 0))
    kpair_explain_text = tk.Text(
        kpair_explain_frame,
        height=14,
        wrap="word",
        state="disabled",
        bg=colors.ENTRY_BG,
        fg=colors.TEXT,
        insertbackground=colors.TEXT,
    )
    kpair_explain_scroll = ttk.Scrollbar(kpair_explain_frame, orient="vertical", command=kpair_explain_text.yview)
    kpair_explain_text.configure(yscrollcommand=kpair_explain_scroll.set)
    kpair_explain_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(8, 0), pady=8)
    kpair_explain_scroll.pack(side=tk.RIGHT, fill=tk.Y, pady=8, padx=(0, 8))

    def _sync_mating_depth_state() -> None:
        base = "normal" if controls_enabled_ref[0] else "disabled"
        depth_state = "disabled" if bool(mating_unbounded_var.get()) else base
        mating_depth_entry.configure(state=depth_state)

    mating_unbounded_var.trace_add("write", lambda *_args: _sync_mating_depth_state())

    def on_calc_mating_ranking() -> None:
        people = state.get("people")
        df_std = state.get("df_std")
        if not people or df_std is None:
            messagebox.showinfo("Info", "Najpierw wczytaj bazę.")
            return

        from datetime import datetime
        import pandas as pd

        try:
            male_limit = int(str(mating_male_limit_var.get()).strip())
        except Exception:
            male_limit = 200
        try:
            female_limit = int(str(mating_female_limit_var.get()).strip())
        except Exception:
            female_limit = 200
        male_limit = max(1, male_limit)
        female_limit = max(1, female_limit)

        current_year = datetime.now().year
        cutoff_birth_year = current_year - mating_age_limit_years

        _loc_part = ""
        _loc_sel_raw_for_note = str(birth_location_filter_var.get()).strip()
        if _loc_sel_raw_for_note and _loc_sel_raw_for_note.lower() != "bez filtra":
            _loc_part = f" • Filtr miejsca ur.: {_norm_birth_location_val(_loc_sel_raw_for_note)}"
        mating_current_year_note_var.set(
            f"Filtr wieku: wiek <= {mating_age_limit_years} lat => birth_year >= {cutoff_birth_year} (rok {current_year}).{_loc_part} "
            f"Ranking: do {mating_ranking_top_n} par (najmniejsze Φ=F potomka); R=2Φ; w liście każdy osobnik (jako sire lub dam) "
            f"maks. {mating_ranking_max_uses_per_id} razy."
        )

        # Standardowy DataFrame po loaderze ma już id/sex/father_id/mother_id, więc filtrujemy po birth_year + sex.
        df_tmp = df_std.copy()
        df_tmp["id"] = df_tmp["id"].astype(str)
        df_tmp["birth_year_num"] = pd.to_numeric(df_tmp.get("birth_year"), errors="coerce")

        # Dla bezpieczeństwa: trzymamy tylko rekordy istniejące w `people`.
        ids_set = set(people.keys())
        df_tmp = df_tmp[df_tmp["id"].isin(ids_set)]

        df_tmp = df_tmp[df_tmp["sex"].isin(["M", "F"])]
        df_tmp = df_tmp[df_tmp["birth_year_num"].notna()]
        df_tmp = df_tmp[df_tmp["birth_year_num"] >= float(cutoff_birth_year)]

        # Filtr miejsca urodzenia (wspólny z rejestrem osobników).
        loc_sel_raw = str(birth_location_filter_var.get()).strip()
        if loc_sel_raw and loc_sel_raw.lower() != "bez filtra":
            loc_sel_norm = _norm_birth_location_val(loc_sel_raw)
            try:
                # Normalizacja do 'NA' dla braków (NaN/None/None/'nan').
                loc_norm = df_tmp["birth_location"].astype(str).str.strip()
                loc_norm = loc_norm.replace({"": "NA", "nan": "NA", "None": "NA", "NaN": "NA"})
                df_tmp = df_tmp[loc_norm == loc_sel_norm]
            except Exception:
                # Jeśli brakuje kolumny / typ jest nietypowy — nie filtrujemy (bez ryzyka wywalenia UI).
                pass

        males_df = df_tmp[df_tmp["sex"] == "M"].sort_values("birth_year_num", ascending=False)
        females_df = df_tmp[df_tmp["sex"] == "F"].sort_values("birth_year_num", ascending=False)

        males_all = males_df["id"].tolist()
        females_all = females_df["id"].tolist()

        males = males_all[:male_limit]
        females = females_all[:female_limit]

        pair_count = len(males) * len(females)
        mating_pair_stats_var.set(
            f"Kandydaci po filtrze: samce={len(males_all)} (użyte {len(males)}), samice={len(females_all)} (użyte {len(females)}). "
            f"Par do policzenia: {pair_count}."
        )

        if pair_count <= 0:
            messagebox.showinfo("Info", "Brak kandydatów do policzenia (po filtrach).")
            return

        # UI ma domyślnie limitowane liczenie; tryb unbounded zostawiamy jako opcję świadomą.
        if bool(mating_unbounded_var.get()):
            max_generations_back: int | None = None
        else:
            try:
                depth = int(str(mating_depth_var.get()).strip())
            except Exception:
                depth = 4
            depth = max(0, depth)
            depth = min(depth, 30)
            max_generations_back = depth

        mating_calc_btn.configure(state="disabled")
        state.pop("mating_phi_matrix", None)
        mating_export_phi_btn.configure(state="disabled")
        try:
            mating_output.configure(state="normal")
            mating_output.delete("1.0", tk.END)
            mating_output.insert("1.0", "Liczenie ranking (może chwilę potrwać)…\n")
            mating_output.configure(state="disabled")
            root.update_idletasks()

            parent_pairs: list[tuple[str, str]] = []
            for sire_id in males:
                for dam_id in females:
                    parent_pairs.append((sire_id, dam_id))

            Fs = batch_offspring_inbreeding_F_from_parent_pairs(
                parent_pairs=parent_pairs,
                people=people,
                max_generations_back=max_generations_back,
            )

            state["mating_phi_matrix"] = {
                "males": list(males),
                "females": list(females),
                "phi_row_major": [float(x) for x in Fs],
                "max_generations_back": max_generations_back,
            }
            mating_export_phi_btn.configure(state="normal")

            ranked = [
                (float(Fs[i]), parent_pairs[i][0], parent_pairs[i][1])
                for i in range(len(parent_pairs))
            ]
            ranked.sort(key=lambda t: t[0])  # smallest Phi first
            use_count: dict[str, int] = {}
            top: list[tuple[float, str, str]] = []
            for F_val, sire_id, dam_id in ranked:
                if len(top) >= mating_ranking_top_n:
                    break
                if use_count.get(sire_id, 0) >= mating_ranking_max_uses_per_id:
                    continue
                if use_count.get(dam_id, 0) >= mating_ranking_max_uses_per_id:
                    continue
                top.append((F_val, sire_id, dam_id))
                use_count[sire_id] = use_count.get(sire_id, 0) + 1
                use_count[dam_id] = use_count.get(dam_id, 0) + 1

            lines: list[str] = [
                f"Wybrane {len(top)} par (cel do {mating_ranking_top_n}, najmniejsze Φ = F potomka; "
                f"każdy ID max {mating_ranking_max_uses_per_id}× w tej liście).\n\n",
                "Φ — współczynnik współzgodności (kinship) między sire a dam; F hipotetycznego potomka = Φ. "
                "R = 2Φ — współczynnik relacji Wrighta.\n\n",
            ]
            for idx, (phi_val, sire_id, dam_id) in enumerate(top, start=1):
                R_val = 2.0 * phi_val
                sire_nm = getattr(people.get(sire_id), "name", None) if people.get(sire_id) else None
                dam_nm = getattr(people.get(dam_id), "name", None) if people.get(dam_id) else None
                sire_lbl = f"{sire_id}" + (f" ({sire_nm})" if sire_nm else "")
                dam_lbl = f"{dam_id}" + (f" ({dam_nm})" if dam_nm else "")
                lines.append(
                    f"{idx}. Φ=F={phi_val:.6f}  R={R_val:.6f}  |  sire={sire_lbl}  x  dam={dam_lbl}\n"
                )

            mating_output.configure(state="normal")
            mating_output.delete("1.0", tk.END)
            mating_output.insert("1.0", "".join(lines))
            mating_output.configure(state="disabled")

            _set_status("Gotowe: ranking mating policzony.")
        except Exception as e:
            state.pop("mating_phi_matrix", None)
            mating_export_phi_btn.configure(state="disabled")
            messagebox.showerror("Błąd", f"Nie udało się policzyć rankingu mating: {e}")
        finally:
            mating_calc_btn.configure(state="normal" if bool(state.get("people")) else "disabled")

    def on_export_mating_phi_matrix() -> None:
        m = state.get("mating_phi_matrix")
        if not m:
            messagebox.showinfo("Info", "Najpierw oblicz ranking kojarzeń — macierz Φ pochodzi z ostatniego liczenia.")
            return
        filename = filedialog.asksaveasfilename(
            title="Zapis macierzy Φ (CSV)",
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv"), ("Wszystkie pliki", "*.*")],
        )
        if not filename:
            return
        males_m = m["males"]
        females_m = m["females"]
        phis = m["phi_row_major"]
        nf = len(females_m)
        try:
            with open(filename, "w", newline="", encoding="utf-8") as f:
                w = csv.writer(f, delimiter=";")
                w.writerow(["sire_id"] + [str(x) for x in females_m])
                for i, mid in enumerate(males_m):
                    row = [str(mid)] + [f"{phis[i * nf + j]:.10g}" for j in range(nf)]
                    w.writerow(row)
        except Exception as ex:
            messagebox.showerror("Błąd", f"Nie udało się zapisać pliku: {ex}")
            return
        _set_status(f"Zapisano macierz Φ: {filename}")
        _open_exported_file(filename)

    def _norm_kin_combo_id(raw: str) -> str:
        s = (raw or "").strip()
        if not s:
            return ""
        return s.split(None, 1)[0].strip()

    def on_kinship_pair() -> None:
        people = state.get("people")
        if not people:
            messagebox.showinfo("Info", "Najpierw wczytaj bazę.")
            return

        sire_id = _norm_kin_combo_id(str(mating_kin_sire_var.get() or ""))
        dam_id = _norm_kin_combo_id(str(mating_kin_dam_var.get() or ""))
        if not sire_id or not dam_id:
            messagebox.showinfo("Info", "Podaj ID ojca (M) i matki (F) — z listy lub wpisane ręcznie.")
            return
        if sire_id not in people or dam_id not in people:
            messagebox.showinfo("Info", "Oba ID muszą istnieć w wczytanej bazie.")
            return

        def _declared_sex(pid: str) -> str | None:
            p = people.get(pid)
            if not p:
                return None
            sx = (p.sex or "").strip().upper()
            return sx if sx else None

        sx_s = _declared_sex(sire_id)
        if sx_s and sx_s != "M":
            messagebox.showinfo("Info", "W polu „Ojciec” wybierz lub wpisz osobnika zapisanego jako samiec (M) w bazie.")
            return
        sx_d = _declared_sex(dam_id)
        if sx_d and sx_d != "F":
            messagebox.showinfo("Info", "W polu „Matka” wybierz lub wpisz osobnika zapisaną jako samica (F) w bazie.")
            return

        if bool(mating_unbounded_var.get()):
            max_generations_back: int | None = None
        else:
            try:
                depth = int(str(mating_depth_var.get()).strip())
            except Exception:
                depth = 4
            max_generations_back = max(0, min(30, depth))
        try:
            phi_v, r_v = wright_kinship_phi_and_relationship_R(
                sire_id, dam_id, people, max_generations_back=max_generations_back
            )
            f_pot = wright_offspring_inbreeding_F_from_parents(
                sire_id, dam_id, people, max_generations_back=max_generations_back
            )
            mating_kin_result_var.set(
                f"Φ = {phi_v:.6f}  |  R = 2Φ = {r_v:.6f}  |  F potomka (ojciec×matka) = {f_pot:.6f} "
                f"({'bez limitu' if max_generations_back is None else f'max {max_generations_back} pok.'})"
            )
            ex = explain_pair_kinship(sire_id, dam_id, people, max_generations_back=max_generations_back)
            warn = close_kinship_note(phi_v)
            lines: list[str] = []
            if warn:
                lines.append("Ostrzeżenie (hodowlane): " + warn)
                lines.append("")
            lines.extend(
                [
                    f"Φ z rekurencji (wzorzec Wrighta / Malecot, symetryczne): {ex.phi_recursive:.8f}",
                    f"Suma surowych wyrazów ścieżkowych: {ex.phi_path_sum_raw:.8f}  |  skala do Φ: {ex.path_scale:.6f}",
                    f"Liczba par ścieżek (kombinacje ścieżek do wspólnych przodków): {ex.n_path_pairs}",
                    f"Wspólnych węzłów z niezerowym wkładem: {ex.n_distinct_common_ancestors}",
                    "",
                    "Top wspólni przodkowie (wkład do Φ, % z Φ, liczba par ścieżek, suma surowa):",
                ]
            )
            for aid, to_phi, raw_v, npair in ex.by_ancestor[:28]:
                ap = people.get(aid)
                nm = (ap.name if ap and ap.name else "-")[:40]
                pct = (100.0 * to_phi / ex.phi_recursive) if ex.phi_recursive > 1e-15 else 0.0
                lines.append(f"  • {aid}  ({nm})  Φ={to_phi:.6f}  ({pct:.2f}%)  |  pary={npair}  |  surowe={raw_v:.6f}")
            lines.append("")
            lines.append("Najsilniejsze pary ścieżek (skrót; pierwsza ścieżka = od ojca, druga = od matki):")
            for d in ex.path_pairs[:12]:
                pa = "→".join(d.path_a) if d.path_a else "·"
                pb = "→".join(d.path_b) if d.path_b else "·"
                lines.append(
                    f"  {d.ancestor_id}: krawędzie {d.n_edges_a}+{d.n_edges_b}, wkład Φ={d.contribution_to_phi:.6f}  |  {pa}  ||  {pb}"
                )
            kpair_explain_text.configure(state="normal")
            kpair_explain_text.delete("1.0", tk.END)
            kpair_explain_text.insert("1.0", "\n".join(lines))
            kpair_explain_text.configure(state="disabled")
        except Exception as ex:
            messagebox.showerror("Błąd", f"Nie udało się policzyć kinship: {ex}")

    mating_calc_btn.configure(command=on_calc_mating_ranking)
    mating_export_phi_btn.configure(command=on_export_mating_phi_matrix)
    mating_kinship_btn.configure(command=on_kinship_pair)

    def set_controls_enabled(enabled: bool) -> None:
        controls_enabled_ref[0] = bool(enabled)
        st = "normal" if enabled else "disabled"
        find_person_btn.configure(state=st)
        persons_search_entry.configure(state=st)
        persons_birth_location_filter_cb.configure(state="readonly" if enabled else "disabled")
        if not enabled:
            _set_persons_detail_text("Wczytaj bazę, aby zobaczyć rejestr i szczegóły osobnika.")
        id_anc_entry.configure(state=st)
        readable_anc_cb.configure(state=st)
        anc_line_radio_sire.configure(state=st)
        anc_line_radio_dam.configure(state=st)
        anc_line_radio_both.configure(state=st)
        unbounded_anc_cb.configure(state=st)
        anc_btn.configure(state=st)

        id_inb_entry.configure(state=st)
        depth_inb_entry.configure(state=st)
        unbounded_inb_cb.configure(state=st)
        inb_btn.configure(state=st)
        mating_unbounded_cb.configure(state=st)
        mating_male_limit_entry.configure(state=st)
        mating_female_limit_entry.configure(state=st)
        mating_calc_btn.configure(state=st)
        mating_kinship_btn.configure(state=st)
        mating_output.configure(state="disabled")
        if enabled and state.get("mating_phi_matrix"):
            mating_export_phi_btn.configure(state="normal")
        else:
            mating_export_phi_btn.configure(state="disabled")

        mating_depth_state = "disabled" if bool(mating_unbounded_var.get()) else st
        mating_depth_entry.configure(state=mating_depth_state)
        mating_kin_sire_cb.configure(state=st if st == "normal" else "disabled")
        mating_kin_dam_cb.configure(state=st if st == "normal" else "disabled")

        # Plan hodowlany (dobór par)
        plan_risk_cb.configure(state=st)
        plan_dam_id_entry.configure(state=st)
        plan_sire_id_entry.configure(state=st)
        plan_calc_pair_btn.configure(state=st)
        plan_min_age_entry.configure(state=st)
        plan_max_age_entry.configure(state=st)
        plan_candidate_limit_entry.configure(state=st)
        plan_top_n_entry.configure(state=st)
        plan_suggest_btn.configure(state=st)
        if not enabled:
            plan_sel_pair["dam"] = None
            plan_sel_pair["sire"] = None
            plan_pedigree_win_btn.configure(state="disabled")
        plan_goal_mean_enabled_cb.configure(state=st)
        plan_goal_mean_F_entry.configure(state=st)
        plan_goal_max_enabled_cb.configure(state=st)
        plan_goal_max_F_entry.configure(state=st)
        plan_max_uses_dam_entry.configure(state=st)
        plan_max_uses_sire_entry.configure(state=st)

        # Combobox: dla enabled używamy "readonly", dla disabled: "disabled"
        plan_line_filter_combo.configure(state="readonly" if enabled else "disabled")

        plan_depth_state = "disabled" if bool(plan_risk_unbounded_var.get()) else st
        plan_risk_depth_entry.configure(state=plan_depth_state)

        if enabled:
            # Tk potrafi nie zaktualizować stanu wejść po trace; wymuszamy sync (tylko gdy włączone).
            _sync_plan_risk_depth_state()
        if enabled:
            # Ustaw wartości domyślne z zakładki "Ustawienia".
            readable_anc_var.set(bool(settings_anc_readable_var.get()))
            unbounded_anc_var.set(bool(settings_anc_unbounded_var.get()))
            try:
                _danc = int(str(settings_anc_depth_var.get()).strip() or "4")
            except Exception:
                _danc = 4
            depth_anc_strvar.set(str(max(0, min(30, _danc))))

            unbounded_inb_var.set(bool(settings_inb_unbounded_var.get()))
            depth_inb_var.set(str(settings_inb_depth_var.get()).strip() or "4")

            # (Klik w grafie przodków jest sterowany osobno w `on_generate_pedigree`.)

            # Domyślne parametry rankingu kojarzeń (Mating).
            mating_unbounded_var.set(False)
            mating_depth_var.set(str(settings_inb_depth_var.get()).strip() or "4")
            mating_male_limit_var.set(str(APP_CFG.mating_default_male_limit))
            mating_female_limit_var.set(str(APP_CFG.mating_default_female_limit))
            mating_current_year_note_var.set("")
            mating_pair_stats_var.set("")

            # Domyślne parametry Planu hodowlanego.
            plan_risk_unbounded_var.set(bool(settings_plan_unbounded_var.get()))
            plan_risk_depth_var.set(str(settings_plan_depth_var.get()).strip() or "4")
            plan_min_age_var.set(str(settings_plan_min_age_var.get()).strip() or "0")
            plan_max_age_var.set(str(settings_plan_max_age_var.get()).strip() or "80")
            plan_candidate_limit_var.set(str(settings_plan_candidate_limit_var.get()).strip() or "25")
            plan_top_n_var.set(str(settings_plan_top_n_var.get()).strip() or "20")
            plan_max_uses_dam_var.set(str(settings_plan_max_dam_uses_var.get()).strip() or "3")
            plan_max_uses_sire_var.set(str(settings_plan_max_sire_uses_var.get()).strip() or "3")
            plan_goal_mean_enabled_var.set(bool(settings_plan_goal_mean_enabled_var.get()))
            plan_goal_mean_F_var.set(str(settings_plan_goal_mean_f_var.get()).strip() or "0.05")
            plan_goal_max_enabled_var.set(bool(settings_plan_goal_max_enabled_var.get()))
            plan_goal_max_F_var.set(str(settings_plan_goal_max_f_var.get()).strip() or "0.10")
            rep_include_plots_export_var.set(bool(rep_default_include_plots_var.get()))

        validation_export_btn.configure(
            state="normal" if enabled and state.get("validation_report") is not None else "disabled"
        )

        # Po zmianie unbounded/depth z konfiguracji trace bywa opóźniony — wymuszamy stan pól głębokości.
        _sync_anc_depth_state()
        _sync_mating_depth_state()

    # Double-click w Osobniki ustawia ID do rodowodu i analiz
    def on_tree_double_click(_event: tk.Event) -> None:
        selection = tree.selection()
        if not selection:
            return
        values = tree.item(selection[0], "values")
        if not values:
            return
        picked_id = str(values[0])
        id_anc_var.set(picked_id)
        id_inb_var.set(picked_id)

    tree.bind("<Double-1>", on_tree_double_click)

    # -------------------------
    # Dataset apply
    # -------------------------
    def _id_sort_key(s: str) -> tuple[int, str]:
        import re

        m = re.match(r"^(\d+)([A-Za-z]*)$", s)
        if not m:
            return (10**30, s)
        return (int(m.group(1)), m.group(2) or "")

    def _refresh_persons_tree_rows(df_std_in) -> None:
        """Rebuild tabelę Osobników według aktualnego filtra 'birth_location'."""
        try:
            for item in tree.get_children():
                tree.delete(item)
        except Exception:
            pass

        df_std_sorted = df_std_in
        try:
            df_std_sorted = df_std_in.sort_values(
                by="id",
                key=lambda s: s.astype(str).map(_id_sort_key),
            ).reset_index(drop=True)
        except Exception:
            df_std_sorted = df_std_in

        loc_sel_raw = str(birth_location_filter_var.get()).strip()
        loc_sel_norm: str | None = None
        if loc_sel_raw and loc_sel_raw.lower() != "bez filtra":
            loc_sel_norm = _norm_birth_location_val(loc_sel_raw)

        if loc_sel_norm and "birth_location" in df_std_sorted.columns:
            try:
                df_std_sorted = df_std_sorted.copy()
                df_std_sorted["_birth_loc_norm"] = df_std_sorted["birth_location"].apply(_norm_birth_location_val)
                df_std_sorted = df_std_sorted[df_std_sorted["_birth_loc_norm"] == loc_sel_norm].reset_index(drop=True)
            except Exception:
                # Jeżeli filtr nie zadziała (np. nietypowy typ kolumny), pokazujemy pełną listę.
                pass

        # Insert all rows in A->Z order by ID.
        for _, row in df_std_sorted.iterrows():
            pid = str(row.get("id"))
            lm = state.get("line_memberships", {}).get(pid, None)

            def _cell(v: object) -> str:
                if v is None:
                    return ""
                # NaN check (bez numpy).
                if isinstance(v, float) and v != v:
                    return ""
                return str(v)

            sire_founder = ""
            sire_steps = ""
            dam_founder = ""
            dam_steps = ""
            if lm is not None:
                sire_founder = _cell(lm.sire_founder_id) + (f" ({_cell(lm.sire_founder_name)})" if lm.sire_founder_name else "")
                sire_steps = _cell(lm.sire_steps)
                dam_founder = _cell(lm.dam_founder_id) + (f" ({_cell(lm.dam_founder_name)})" if lm.dam_founder_name else "")
                dam_steps = _cell(lm.dam_steps)

            def _norm_line(line_val: object) -> str:
                if line_val is None:
                    return "NA"
                if isinstance(line_val, float) and line_val != line_val:
                    return "NA"
                s = str(line_val).strip().upper()
                if s in {"LB", "LC"}:
                    return s
                return "NA"

            line = _norm_line(row.get("line"))

            # Rodzice mogą być poza bazą jako rekordy - wtedy "NA".
            # Wartości pochodzą bezpośrednio z kolumn Excela (E/J/M):
            father_line = _norm_line(row.get("father_line"))
            mother_line = _norm_line(row.get("mother_line"))

            tree.insert(
                "",
                "end",
                values=(
                    _cell(row.get("id")),
                    _cell(row.get("name")),
                    _cell(row.get("sex")),
                    _cell(row.get("birth_year")),
                    _cell(row.get("birth_location")),
                    _cell(row.get("father_id")),
                    _cell(row.get("mother_id")),
                    sire_founder,
                    sire_steps,
                    dam_founder,
                    dam_steps,
                    line,
                    father_line,
                    mother_line,
                ),
            )

    def _on_birth_location_filter_change(_event: tk.Event | None = None) -> None:
        if not controls_enabled_ref[0]:
            return
        df_cur = state.get("df_std")
        if df_cur is None or getattr(df_cur, "empty", True):
            return
        _refresh_persons_tree_rows(df_cur)

    persons_birth_location_filter_cb.bind("<<ComboboxSelected>>", _on_birth_location_filter_change)

    def _apply_dataset(df_std, people, source: str) -> None:
        state["df_std"] = df_std
        state["people"] = people
        state.pop("mating_phi_matrix", None)
        mating_export_phi_btn.configure(state="disabled")
        mating_kin_result_var.set("")
        try:
            kpair_explain_text.configure(state="normal")
            kpair_explain_text.delete("1.0", tk.END)
            kpair_explain_text.configure(state="disabled")
        except Exception:
            pass
        try:
            df_k = df_std.copy()
            df_k["id"] = df_k["id"].astype(str)
            if "sex" in df_k.columns:
                sx = df_k["sex"].astype(str).str.strip().str.upper()
                males = sorted(df_k.loc[sx == "M", "id"].drop_duplicates().tolist(), key=_id_sort_key)
                females = sorted(df_k.loc[sx == "F", "id"].drop_duplicates().tolist(), key=_id_sort_key)
            else:
                males, females = [], []
            mating_kin_sire_cb.configure(values=males)
            mating_kin_dam_cb.configure(values=females)
            mating_kin_sire_var.set("")
            mating_kin_dam_var.set("")
        except Exception:
            pass
        # Precompute line memberships for quick display in "Osobniki" and "Rodowód".
        ids = [str(x) for x in df_std["id"].tolist()]
        try:
            state["line_memberships"] = compute_all_line_memberships(people, person_ids=ids)
        except Exception:
            state["line_memberships"] = {}

        # Cross-walidacja bazy (spójność rodowodu i danych pomocniczych).
        try:
            from datetime import datetime

            report = validate_loaded_dataset(df_std=df_std, people=people, current_year=datetime.now().year)
            state["validation_report"] = report
            validation_var.set(report.ui_summary())
            if controls_enabled_ref[0]:
                validation_export_btn.configure(state="normal")
        except Exception:
            state["validation_report"] = None
            validation_var.set("Walidacja bazy: nie udało się przeprowadzić walidacji.")
            validation_export_btn.configure(state="disabled")

        ids = df_std["id"].dropna().astype(str)
        if len(ids) > 0:
            min_id = min(ids.tolist(), key=_id_sort_key)
            max_id = max(ids.tolist(), key=_id_sort_key)
            dataset_range_var.set(f"Zakres ID (Number): min {min_id} / max {max_id}")
        else:
            dataset_range_var.set("Zakres ID: brak")

        _set_status(source + f" • {len(df_std)} wierszy")

        # Wczytujemy CAŁĄ bazę (bez limitu podglądu).
        # Aktualizacja listy opcji filtra miejsca urodzenia.
        try:
            if "birth_location" in df_std.columns:
                uniq: set[str] = set()
                has_na = False
                for v in df_std["birth_location"].tolist():
                    nv = _norm_birth_location_val(v)
                    if nv == "NA":
                        has_na = True
                    else:
                        uniq.add(nv)
                opts = ["Bez filtra"] + sorted(uniq, key=lambda s: str(s).lower())
                if has_na:
                    opts.append("NA")
                persons_birth_location_filter_cb.configure(values=opts)
            else:
                persons_birth_location_filter_cb.configure(values=["Bez filtra"])
        except Exception:
            persons_birth_location_filter_cb.configure(values=["Bez filtra"])
        birth_location_filter_var.set("Bez filtra")

        _refresh_persons_tree_rows(df_std)

        # Default IDs: first person with at least one parent.
        if not df_std.empty:
            with_parents = df_std[df_std["father_id"].notna() | df_std["mother_id"].notna()]
            first_row = with_parents.iloc[0] if not with_parents.empty else df_std.iloc[0]
            picked = str(first_row["id"])
            id_anc_var.set(picked)
            id_inb_var.set(picked)

        _refresh_validation_missing_heatmap()
        _update_population_metrics(df_std)
        set_controls_enabled(True)

    # -------------------------
    # Rodowód generate
    # -------------------------
    def on_generate_pedigree() -> None:
        people = state.get("people")
        if not people:
            messagebox.showinfo("Info", "Najpierw wczytaj bazę.")
            return

        person_id = str(id_anc_var.get()).strip()
        if not person_id:
            messagebox.showerror("Błąd", "Podaj ID (Number).")
            return
        if person_id not in people:
            messagebox.showerror("Błąd", "Nie ma takiego ID w wczytanych danych.")
            return

        # Callback używany przez plot: klik w węzeł aktualizuje informacje o linii.
        lm_map = state.get("line_memberships", {}) or {}

        def _pid_to_line(pid: object, people_map: dict[str, object]) -> str:
            if pid is None:
                return "NA"
            pid_s = str(pid)
            p = people_map.get(pid_s)
            line_val = getattr(p, "line", None) if p else None
            if not line_val:
                return "NA"
            s = str(line_val).strip().upper()
            if s in {"LB", "LC"}:
                return s
            return "NA"

        def _norm_line_val(v: object) -> str:
            if v is None:
                return "NA"
            if isinstance(v, float) and v != v:
                return "NA"
            s = str(v).strip().upper()
            if s in {"LB", "LC"}:
                return s
            return "NA"

        def _fmt_pid(pid: object) -> str:
            if not pid:
                return "brak danych"
            pid_s = str(pid)
            mem = lm_map.get(pid_s)
            if mem is None:
                return (
                    f"Sireline: {pid_s} (NA) [steps=0]\n"
                    f"Damline: {pid_s} (NA) [steps=0]"
                )
            return (
                f"Sireline: {mem.sire_founder_id} ({mem.sire_founder_name or 'NA'}) [steps={mem.sire_steps}]\n"
                f"Damline: {mem.dam_founder_id} ({mem.dam_founder_name or 'NA'}) [steps={mem.dam_steps}]"
            )

        unbounded = bool(unbounded_anc_var.get())
        line_mode = str(anc_line_mode_var.get()).strip().lower()
        if line_mode not in {"both", "sire", "dam"}:
            line_mode = "both"

        depth: int | None = None
        if not unbounded:
            try:
                depth = int(str(depth_anc_strvar.get()).strip())
            except Exception:
                depth = 4
            depth = max(0, min(30, depth))

        # Budujemy `levels` i `edges` w zależności od trybu linii.
        def _build_levels_edges_line(follow: str) -> tuple[dict[str, int], list[tuple[str, str]]]:
            from collections import deque

            levels_local: dict[str, int] = {person_id: 0}
            edges_local: list[tuple[str, str]] = []
            q = deque([person_id])

            while q:
                cur = q.popleft()
                cur_lvl = levels_local[cur]
                if depth is not None and cur_lvl >= depth:
                    continue

                p = people.get(cur)
                if p is None:
                    continue

                parent_ids: list[str] = []
                if follow == "sire" and p.father_id:
                    parent_ids.append(p.father_id)
                if follow == "dam" and p.mother_id:
                    parent_ids.append(p.mother_id)

                for parent_id in parent_ids:
                    edges_local.append((parent_id, cur))
                    nxt_lvl = cur_lvl + 1
                    prev = levels_local.get(parent_id)
                    if prev is None or nxt_lvl < prev:
                        levels_local[parent_id] = nxt_lvl
                        q.append(parent_id)
            return levels_local, edges_local

        if line_mode == "both":
            if unbounded:
                from app.pedigree.ancestor_pedigree import get_ancestor_levels_unbounded

                levels = get_ancestor_levels_unbounded(person_id=person_id, people=people)  # type: ignore[arg-type]
                # krawędzie wyznaczamy z mapy rodziców (rodzice będą w `levels` nawet jako placeholdery).
                people_all_tmp = ensure_people_for_nodes(levels=levels, people=people)
                edges = []
                for child_id in levels.keys():
                    p = people_all_tmp.get(child_id)
                    if not p:
                        continue
                    if p.father_id and p.father_id in levels:
                        edges.append((p.father_id, child_id))
                    if p.mother_id and p.mother_id in levels:
                        edges.append((p.mother_id, child_id))
            else:
                levels, edges = get_ancestor_levels_and_edges(
                    person_id=person_id, depth=int(depth) if depth is not None else 0, people=people
                )
        else:
            levels, edges = _build_levels_edges_line(follow=line_mode)

        if not levels:
            messagebox.showerror("Błąd", "Nie znaleziono przodków w podanym zakresie.")
            return

        people_all = ensure_people_for_nodes(levels=levels, people=people)

        father_id_ref = people.get(person_id).father_id if people.get(person_id) else None
        mother_id_ref = people.get(person_id).mother_id if people.get(person_id) else None

        def _on_node_click(nid: str, dbl: bool) -> None:
            # Podmieniamy opis w panelu linii.
            rod_title_var.set(f"Zaznaczono osobnika: {nid}")
            clicked_line = _pid_to_line(nid, people_all)
            clicked_text = f"Zaznaczono ({nid}) | line={clicked_line}:\n{_fmt_pid(nid)}"
            if nid == person_id:
                rod_line_subj_var.set(clicked_text.replace(f"Kliknięty ({nid})", f"Oceniany ({nid})"))
            elif nid == father_id_ref:
                rod_line_father_var.set(clicked_text.replace(f"Kliknięty ({nid})", f"Ojciec ({nid})"))
            elif nid == mother_id_ref:
                rod_line_mother_var.set(clicked_text.replace(f"Kliknięty ({nid})", f"Matka ({nid})"))
            else:
                # Nie jest to bezpośredni rodzic/badany osobnik - zostawiamy bez zmiany pozostałych kolumn.
                rod_line_subj_var.set(clicked_text)

            # Double-click: ustaw ID w Analityka hodowlana → Inbred i policz F.
            if dbl:
                id_inb_var.set(nid)
                try:
                    notebook.select(tab_analysis)
                except Exception:
                    pass
                try:
                    analyses_nb.select(tab_inb)
                except Exception:
                    pass
                try:
                    on_calc_inbreeding()
                except Exception:
                    pass

        fig = plot_ancestor_pedigree(
            person_id=person_id,
            levels=levels,
            edges=edges,
            people=people_all,
            readable_mode=bool(readable_anc_var.get()),
            enable_click_highlight=bool(settings_anc_click_var.get()),
            on_node_click=_on_node_click,
        )
        rod_title_var.set(f"Przodkowie: {person_id}")

        # Linie dla ocenianego osobnika i jego rodziców bierzemy z kolumn Excela:
        # - osobnik: `line` (E)
        # - ojciec: `father_line` (J)
        # - matka: `mother_line` (M)
        df_std = state.get("df_std")
        df_row = None
        try:
            if df_std is not None and not df_std.empty:
                matches = df_std[df_std["id"] == person_id]
                if not matches.empty:
                    df_row = matches.iloc[0]
        except Exception:
            df_row = None

        own_line = _norm_line_val(df_row.get("line") if df_row is not None else None)
        father_line = _norm_line_val(df_row.get("father_line") if df_row is not None else None)
        mother_line = _norm_line_val(df_row.get("mother_line") if df_row is not None else None)

        # Fallback, jeśli dla jakiegoś powodu nie uda się znaleźć wiersza w df_std.
        if own_line == "NA":
            own_line = _pid_to_line(person_id, people_all)
        tmp_father_id = people.get(person_id).father_id if people.get(person_id) else None
        if father_line == "NA" and tmp_father_id:
            father_line = _pid_to_line(tmp_father_id, people_all)
        tmp_mother_id = people.get(person_id).mother_id if people.get(person_id) else None
        if mother_line == "NA" and tmp_mother_id:
            mother_line = _pid_to_line(tmp_mother_id, people_all)

        def _fmt_pid(pid: object) -> str:
            if not pid:
                return "brak danych"
            pid_s = str(pid)
            mem = lm_map.get(pid_s)
            if mem is None:
                # Brak rekordu osoby w bazie (założyciel/nieznany rodzic).
                return (
                    f"Sireline: {pid_s} (NA) [steps=0]\n"
                    f"Damline: {pid_s} (NA) [steps=0]"
                )
            return (
                f"Sireline: {mem.sire_founder_id} ({mem.sire_founder_name or 'NA'}) [steps={mem.sire_steps}]\n"
                f"Damline: {mem.dam_founder_id} ({mem.dam_founder_name or 'NA'}) [steps={mem.dam_steps}]"
            )

        father_id = people.get(person_id).father_id if people.get(person_id) else None
        mother_id = people.get(person_id).mother_id if people.get(person_id) else None

        rod_line_subj_var.set(f"Oceniany ({person_id}) | line={own_line}:\n{_fmt_pid(person_id)}")
        rod_line_father_var.set(
            f"Ojciec ({father_id or 'NA'}) | line={father_line}:\n{_fmt_pid(father_id)}"
        )
        rod_line_mother_var.set(
            f"Matka ({mother_id or 'NA'}) | line={mother_line}:\n{_fmt_pid(mother_id)}"
        )

        _clear_frame(rod_plot_frame)
        try:
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie mogę osadzić wykresu: {e}")
            return

        ttk.Button(
            rod_plot_frame,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig: _save_figure_as_jpeg(f, default_basename=f"rodowod_{person_id}"),
        ).pack(anchor="w", pady=(0, 6))

        canvas = FigureCanvasTkAgg(fig, master=rod_plot_frame)
        canvas.draw()
        try:
            toolbar = NavigationToolbar2Tk(canvas, rod_plot_frame)  # type: ignore[arg-type]
            toolbar.update()
            toolbar.pack(side=tk.TOP, fill=tk.X)
        except Exception:
            pass
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    anc_btn.configure(command=on_generate_pedigree)

    # -------------------------
    # Analyses: inbred
    # -------------------------
    def on_calc_inbreeding() -> None:
        people = state.get("people")
        if not people:
            messagebox.showinfo("Info", "Najpierw wczytaj bazę.")
            return

        person_id = str(id_inb_var.get()).strip()
        if not person_id:
            messagebox.showerror("Błąd", "Podaj ID (Number).")
            return
        if person_id not in people:
            messagebox.showerror("Błąd", "Nie ma takiego ID w wczytanych danych.")
            return

        if bool(unbounded_inb_var.get()):
            f_res = wright_inbreeding_F(person_id=person_id, people=people, max_generations_back=None)
        else:
            try:
                depth = int(str(depth_inb_var.get()).strip())
            except Exception:
                messagebox.showerror("Błąd", "Max pokoleń musi być liczbą całkowitą.")
                return
            if depth < 0:
                messagebox.showerror("Błąd", "Max pokoleń nie może być ujemne.")
                return
            depth = min(depth, 30)
            f_res = wright_inbreeding_F(person_id=person_id, people=people, max_generations_back=depth)

        inb_result_var.set(f"F = {f_res.F:.6f}")

        # --- Kompletność rodowodu (MG/EG/PCI) dla wskazanego osobnika (ANC, bez limitu) ---
        MG = 0
        EG = 0.0
        PCI = 0.0
        by_gen: dict[int, int] = {}
        try:
            levels = get_ancestor_levels_unbounded(person_id=person_id, people=people)
            for _aid, lvl in levels.items():
                if lvl is None:
                    continue
                try:
                    g = int(lvl)
                except Exception:
                    continue
                if g <= 0:
                    continue
                by_gen[g] = by_gen.get(g, 0) + 1

            if by_gen:
                MG = int(max(by_gen.keys()))
                pci_sum = 0.0
                for g in range(1, MG + 1):
                    a_g = int(by_gen.get(g, 0))
                    pcl_g = float(a_g) / float(2**g)
                    EG += pcl_g
                    pci_sum += pcl_g
                PCI = pci_sum / float(MG) if MG > 0 else 0.0
        except Exception:
            pass

        # Linia (sire/dam) dla ocenianego osobnika i jego rodziców.
        subj = get_line_membership(person_id, people)

        father_id = f_res.father_id
        mother_id = f_res.mother_id
        father_mem = get_line_membership(father_id, people) if father_id else None
        mother_mem = get_line_membership(mother_id, people) if mother_id else None

        def _pid_to_line(pid: object) -> str:
            if pid is None:
                return "NA"
            pid_s = str(pid)
            p = people.get(pid_s)
            line_val = getattr(p, "line", None) if p else None
            if not line_val:
                return "NA"
            s = str(line_val).strip().upper()
            if s in {"LB", "LC"}:
                return s
            return "NA"

        def _norm_line_val(v: object) -> str:
            if v is None:
                return "NA"
            if isinstance(v, float) and v != v:
                return "NA"
            s = str(v).strip().upper()
            if s in {"LB", "LC"}:
                return s
            return "NA"

        # Linie z kolumn Excela (E/J/M) dla wiersza ocenianego osobnika.
        df_std = state.get("df_std")
        df_row = None
        try:
            if df_std is not None and not df_std.empty:
                matches = df_std[df_std["id"] == person_id]
                if not matches.empty:
                    df_row = matches.iloc[0]
        except Exception:
            df_row = None

        own_line = _norm_line_val(df_row.get("line") if df_row is not None else None)
        father_line = _norm_line_val(df_row.get("father_line") if df_row is not None else None)
        mother_line = _norm_line_val(df_row.get("mother_line") if df_row is not None else None)

        # Fallback: jeśli wiersz nie istnieje w df_std (rzadkie), bierz z rekordów w `people`.
        if own_line == "NA":
            own_line = _pid_to_line(subj.person_id)
        if father_line == "NA" and father_id:
            father_line = _pid_to_line(father_id)
        if mother_line == "NA" and mother_id:
            mother_line = _pid_to_line(mother_id)

        def _fmt(mem) -> str:
            if mem is None or (getattr(mem, "sire_founder_id", None) is None and getattr(mem, "dam_founder_id", None) is None):
                return "brak danych"
            sire = f"{mem.sire_founder_id} ({mem.sire_founder_name})"
            dam = f"{mem.dam_founder_id} ({mem.dam_founder_name})"
            return (
                f"Sireline: {sire} [steps={mem.sire_steps}]"
                f"\nDamline: {dam} [steps={mem.dam_steps}]"
            )

        note_text = (
            "Inbred (Wright F) + przynależność do linii:\n"
            f"- max pokoleń (ścieżki n1+n2) w Phi: {f_res.used_generations}\n"
            f"- Oceniany osobnik: {subj.person_id} ({subj.person_name}) | line={own_line}\n"
            f"{_fmt(subj)}\n"
            f"- Ojciec: {father_id} ({f_res.father_name}) | line={father_line}\n"
            f"{_fmt(father_mem)}\n"
            f"- Matka: {mother_id} ({f_res.mother_name}) | line={mother_line}\n"
            f"{_fmt(mother_mem)}\n\n"
            f"Kompletność rodowodu (ANC, bez limitu pokoleń): MG={MG}, EG={EG:.4f}, PCI={PCI:.4f}\n"
            f"Inbred F (Wright) liczone jako F(i)=Phi(sire(i), dam(i)); Phi rekurencyjnie wykorzystuje wspólnych przodków,\n"
            f"a brakujący rodzic traktowany jest jak „founder-stop” (wkład do Phi dla tej ścieżki wynosi 0).\n"
        )

        # Render w 3 kolumnach (czytelniej niż jedna długie etykieta).
        # Każda linia tekstu trafia kolejno do kolumn row-wise.
        for w in inb_note_frame.winfo_children():
            w.destroy()

        raw_lines = [ln.rstrip() for ln in note_text.splitlines()]
        # Pomijamy puste wiersze (kontrolę odstępów robimy pady na etykietach).
        note_items = [ln for ln in raw_lines if ln.strip()]

        n_cols = 3
        wraplength = 320
        for i, ln in enumerate(note_items):
            col = i % n_cols
            row = i // n_cols
            ttk.Label(
                inb_note_frame,
                text=ln,
                foreground=colors.MUTED,
                wraplength=wraplength,
                justify="left",
            ).grid(row=row, column=col, sticky="w", padx=(6, 10), pady=(2, 2))

        # Diagnostic plot (F vs max generations)
        _clear_frame(inb_plot_frame)
        try:
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie mogę narysować wykresu: {e}")
            return

        diag_frame = ttk.Frame(inb_plot_frame)
        diag_frame.pack(fill=tk.BOTH, expand=True)
        comp_frame = ttk.Frame(inb_plot_frame)
        comp_frame.pack(fill=tk.BOTH, expand=True, pady=(12, 0))

        max_trace_depth = min(20, int(f_res.used_generations) if f_res.used_generations else 0)
        depths = list(range(0, max_trace_depth + 1))
        Fs = [
            wright_inbreeding_F(person_id=person_id, people=people, max_generations_back=int(d)).F
            for d in depths
        ]

        fig, ax = plt.subplots(figsize=(7.5, 3.6))
        ax.plot(depths, Fs, marker="o", linewidth=2, color=colors.EDGE_PLOT)
        ax.set_title(f"Inbred (Wright F) - diagnostyka (ID {person_id})")
        ax.set_xlabel("max pokoleń")
        ax.set_ylabel("F")
        ax.grid(True, alpha=0.25)

        ttk.Button(
            diag_frame,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig: _save_figure_as_jpeg(f, default_basename=f"analiza_inbred_diag_{person_id}"),
        ).pack(anchor="w", pady=(0, 6))

        canvas = FigureCanvasTkAgg(fig, master=diag_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # --- Completeness per generation (ANC, PCL) ---
        fig2 = plt.Figure(figsize=(7.5, 3.6))
        ax_a = fig2.add_subplot(1, 2, 1)
        ax_pcl = fig2.add_subplot(1, 2, 2)

        if by_gen:
            gens = sorted(by_gen.keys())
            a_vals = [by_gen[g] for g in gens]
            pcl_vals = [float(by_gen[g]) / float(2**g) for g in gens]

            ax_a.bar([str(g) for g in gens], a_vals, color=colors.BUTTON_BG2, edgecolor=colors.ACCENT)
            ax_a.set_title("Ankiety przodków (a_g)")
            ax_a.set_xlabel("pokolenie (g)")
            ax_a.set_ylabel("a_g")
            ax_a.tick_params(axis="x", labelsize=8)

            ax_pcl.plot(gens, pcl_vals, marker="o", linewidth=2, color=colors.BUTTON_BG)
            ax_pcl.axhline(1.0, color=colors.ACCENT, linewidth=1, alpha=0.6)
            ax_pcl.set_title("Kompletność per pokolenie (PCL)")
            ax_pcl.set_xlabel("pokolenie (g)")
            ax_pcl.set_ylabel("PCL = a_g / 2^g")
            ax_pcl.set_ylim(0, 1.05)
            ax_pcl.grid(True, alpha=0.25)
        else:
            ax_a.text(0.5, 0.5, "Brak danych ANC", ha="center", va="center")
            ax_a.axis("off")
            ax_pcl.axis("off")

        fig2.tight_layout()
        ttk.Button(
            comp_frame,
            text="Zapis wykresu (jpeg)",
            command=lambda f=fig2: _save_figure_as_jpeg(f, default_basename=f"analiza_inbred_anc_{person_id}"),
        ).pack(anchor="w", pady=(0, 6))

        canvas2 = FigureCanvasTkAgg(fig2, master=comp_frame)
        canvas2.draw()
        canvas2.get_tk_widget().pack(fill=tk.BOTH, expand=True)
        ttk.Label(
            comp_frame,
            text="PCL (Pedigree Completeness Level) pokazuje, w którym pokoleniu wstecz mamy ilu przodków względem maksimum (2^g).",
            font=tk_font(8),
            foreground=colors.MUTED,
            wraplength=900,
            justify="left",
        ).pack(anchor="w", pady=(6, 0))

    inb_btn.configure(command=on_calc_inbreeding)

    # -------------------------
    # Main menu (top bar)
    # -------------------------
    def _go_to_tab(tab_widget: ttk.Frame) -> None:
        try:
            notebook.select(tab_widget)
        except Exception:
            pass

    def _go_to_plan_hodowlany() -> None:
        _go_to_tab(tab_analysis)
        try:
            analyses_nb.select(tab_breeding)
        except Exception:
            pass

    def _show_about_app() -> None:
        about_text = (
            "WisentPedigree Pro+\n\n"
            "Aplikacja do analizy rodowodów żubrów i wsparcia zarządzania stadem.\n\n"
            "Najważniejsze funkcjonalności:\n"
            "- Import danych (plik/URL) i mapowanie kolumn,\n"
            "- Walidacja i kontrola spójności danych,\n"
            "- Analityka osobnika (Wright F, kompletność rodowodu, linie),\n"
            "- Metryki populacji i wizualizacje,\n"
            "- Interaktywny graf pedigree,\n"
            "- Plan hodowli w module Analityka hodowlana (dobór par i ryzyko inbredu potomstwa),\n"
            "- Raporty DOCX/PDF (opcjonalnie z wykresami).\n\n"
            "Autor: Magdalena Perlinska-Teresiak\n"
            "Rok: 2026"
        )
        messagebox.showinfo("O aplikacji", about_text)

    menubar = tk.Menu(root)

    menu_file = tk.Menu(menubar, tearoff=0)
    menu_file.add_command(label="Import danych", command=lambda: _go_to_tab(tab_loading))
    menu_file.add_command(label="Wczytaj domyslna baze", command=on_load_default)
    menu_file.add_command(label="Wybierz plik...", command=on_choose_file)
    menu_file.add_separator()
    menu_file.add_command(label="Zakoncz", command=root.quit)
    menubar.add_cascade(label="Plik", menu=menu_file)

    menu_data = tk.Menu(menubar, tearoff=0)
    menu_data.add_command(label="Import danych", command=lambda: _go_to_tab(tab_loading))
    menu_data.add_command(label="Walidacja bazy", command=lambda: _go_to_tab(tab_validation))
    menu_data.add_command(label="Rejestr osobników", command=lambda: _go_to_tab(tab_persons))
    menu_data.add_command(label="Graf pedigree", command=lambda: _go_to_tab(tab_pedigree))
    menu_data.add_command(label="Metryki populacji", command=lambda: _go_to_tab(tab_population))
    menubar.add_cascade(label="Dane", menu=menu_data)

    menu_analysis = tk.Menu(menubar, tearoff=0)
    menu_analysis.add_command(
        label="Analiza: osobnik / para / ranking",
        command=lambda: _go_to_tab(tab_analysis),
    )
    menu_analysis.add_command(label="Plan hodowli", command=_go_to_plan_hodowlany)
    menu_analysis.add_command(label="Metryki populacji", command=lambda: _go_to_tab(tab_population))
    menubar.add_cascade(label="Analiza", menu=menu_analysis)

    menu_viz = tk.Menu(menubar, tearoff=0)
    menu_viz.add_command(label="Graf pedigree", command=lambda: _go_to_tab(tab_pedigree))
    menu_viz.add_command(
        label="Analiza: osobnik / para / ranking",
        command=lambda: _go_to_tab(tab_analysis),
    )
    menu_viz.add_command(label="Metryki populacji", command=lambda: _go_to_tab(tab_population))
    menubar.add_cascade(label="Wizualizacja", menu=menu_viz)

    menu_export = tk.Menu(menubar, tearoff=0)
    menu_export.add_command(label="Raportowanie", command=lambda: _go_to_tab(tab_reports))
    menu_export.add_command(label="Konfiguracja", command=lambda: _go_to_tab(tab_settings))
    menubar.add_cascade(label="Eksport", menu=menu_export)

    menu_help = tk.Menu(menubar, tearoff=0)
    menu_help.add_command(
        label="Słownik parametrów (F, GI, N_e…)",
        command=lambda: show_help_window(root, "Słownik parametrów", hc.GLOSSARY),
    )
    menu_help.add_command(
        label="Interpretacja wykresów populacji",
        command=lambda: show_help_window(root, "Wykresy populacji", hc.all_charts_text()),
    )
    menu_help.add_command(
        label="Pełna dokumentacja (tekst)",
        command=lambda: show_help_window(root, "WisentPedigree Pro+ — pomoc", hc.FULL_HELP_DOCUMENT),
    )
    menu_help.add_command(
        label="Przewodnik metod (PDF — cytowania)",
        command=on_save_methods_guide_pdf,
    )
    menu_help.add_separator()
    menu_help.add_command(label="O aplikacji", command=_show_about_app)
    menu_help.add_command(label="Konfiguracja", command=lambda: _go_to_tab(tab_settings))
    menubar.add_cascade(label="Pomoc", menu=menu_help)

    root.config(menu=menubar)

    # Disable controls until dataset is loaded.
    set_controls_enabled(False)
    _refresh_validation_missing_heatmap()

    # Auto-load default dataset for convenience.
    try:
        on_load_default()
    except Exception:
        _set_status("Nie udało się wczytać domyślnej bazy.")

    root.mainloop()

