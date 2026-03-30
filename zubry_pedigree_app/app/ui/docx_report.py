"""Proste raporty DOCX z płaskiego tekstu (linia → akapit)."""

from __future__ import annotations

from io import BytesIO


def report_plain_text_to_docx_bytes(
    text: str,
    *,
    title: str = "WisentPedigree Pro+ — Raport",
    footer_note: str | None = None,
) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph("")
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.add_paragraph("")
    if footer_note:
        doc.add_paragraph(footer_note)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
